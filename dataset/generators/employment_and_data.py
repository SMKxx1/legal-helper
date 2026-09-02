"""Generate synthetic employment and data processing documents for legal review practice.

Three documents designed to exercise specific reviewer coverage scenarios:
an employment agreement drafted heavily in the employer's favour (non-compete, IP assignment,
termination asymmetry, perpetual confidentiality), a data processing agreement with processor-
favourable terms and missing GDPR controls, and a sparse consultancy agreement missing key
protective clauses entirely (no liability cap, no governing law, no IP clause, vague payment).
Run with ``python ../dataset/generators/employment_and_data.py`` from ``backend/``.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

_REPO_ROOT = Path(__file__).resolve().parents[2]
_DATASET_DIR = _REPO_ROOT / "dataset"


def _heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def _clause(doc: Document, heading: str, body: str) -> None:
    doc.add_paragraph(heading, style="Heading 2")
    doc.add_paragraph(body)


def build_employment_agreement_senior_engineer() -> Document:
    """An employment agreement for a senior engineer, drafted hard in the employer's favour.
    Plants four major issues: (a) 24-month post-termination non-compete with worldwide scope
    and zero compensation; (b) uncapped IP assignment covering personal-time inventions
    unrelated to the business, with no carve-out for prior inventions; (c) termination for
    convenience by the employer with no notice or severance, asymmetric to the employee's
    3-month notice requirement; (d) perpetual confidentiality surviving indefinitely."""
    doc = Document()
    _heading(doc, "EMPLOYMENT AGREEMENT")
    doc.add_paragraph(
        "This Employment Agreement (this \"Agreement\") is entered into on this 1st day of "
        "September, 2026, by and between Vertex Innovations LLC, a Delaware limited liability "
        "company (\"Company\"), and Harrison J. Mitchell (\"Employee\"). This Agreement sets "
        "forth the terms and conditions of Employee's employment as Senior Software Engineer "
        "reporting to the Vice President of Engineering."
    )
    _clause(
        doc,
        "1. Position and Duties",
        "Employee shall serve as Senior Software Engineer in the Company's Platform Engineering "
        "group, with responsibilities including architecting scalable backend systems, leading "
        "technical design reviews, mentoring junior engineers, and contributing to the Company's "
        "core product development roadmap. Employee shall report to the Vice President of "
        "Engineering and shall perform such duties as are customary for this role or as may be "
        "reasonably assigned by the Company from time to time.",
    )
    _clause(
        doc,
        "2. Compensation and Benefits",
        "The Company shall pay Employee an annual base salary of Four Hundred Fifty Thousand "
        "Dollars ($450,000), payable in accordance with the Company's standard payroll practices. "
        "Employee shall be eligible for an annual discretionary performance bonus of up to thirty "
        "percent (30%) of base salary, contingent upon achievement of objectives determined by "
        "the Company in its sole discretion. Employee shall receive standard employee benefits "
        "including health insurance, retirement plan eligibility, and paid time off in accordance "
        "with Company policy.",
    )
    _clause(
        doc,
        "3. Non-Compete and Non-Solicitation",
        "During the term of this Agreement and for a period of twenty-four (24) months following "
        "the cessation of employment for any reason, Employee shall not, directly or indirectly, "
        "(a) engage in any business, employment, or other activity that is competitive with the "
        "Company's business in any industry or geography served by the Company; (b) provide "
        "services or products to any entity that competes with the Company; or (c) solicit or "
        "accept business from any customer or prospective customer of the Company. The geographic "
        "scope of this restriction encompasses all territories worldwide where the Company conducts "
        "or plans to conduct business. Employee acknowledges that this restriction is reasonable "
        "and necessary to protect the Company's legitimate business interests. During the "
        "24-month restricted period, Employee shall not be entitled to any compensation, severance, "
        "or other consideration from the Company, though Employee remains prohibited from competing "
        "regardless.",
    )
    _clause(
        doc,
        "4. Intellectual Property Assignment",
        "Employee hereby assigns to the Company all right, title, and interest in and to all "
        "inventions, discoveries, works of authorship, software code, designs, improvements, and "
        "other intellectual property, whether or not patentable or copyrightable, that Employee "
        "conceives, develops, creates, or reduces to practice at any time during the employment "
        "relationship, including all such work created outside of working hours, on personal "
        "time, and regardless of whether such work relates to the Company's business or uses "
        "Company equipment or resources. This assignment includes all inventions created before "
        "the date of this Agreement and any pre-existing inventions or intellectual property "
        "owned by Employee. Employee further waives any moral rights or rights of attribution "
        "with respect to such assigned intellectual property and shall execute any documents "
        "reasonably necessary to effectuate such assignment.",
    )
    _clause(
        doc,
        "5. Termination",
        "The Company may terminate Employee's employment at any time and for any reason or no "
        "reason, without cause, without advance notice, and without any severance payment or "
        "consideration. Termination shall be effective immediately upon notice to Employee. "
        "In contrast, Employee may terminate employment voluntarily only upon providing the "
        "Company with at least ninety (90) days' prior written notice and a detailed handover "
        "plan. If Employee fails to provide such notice, the Company may pursue any remedies "
        "available at law or in equity, and may withhold any accrued but unpaid compensation.",
    )
    _clause(
        doc,
        "6. Confidentiality",
        "Employee acknowledges that during employment, Employee will have access to and become "
        "acquainted with numerous trade secrets and confidential information of the Company, "
        "including but not limited to source code, algorithms, business plans, customer lists, "
        "financial information, and proprietary processes. Employee agrees to maintain the "
        "confidentiality of all such information both during and after employment, indefinitely "
        "and without limitation in time. This confidentiality obligation survives termination "
        "of employment forever and shall apply to all information regardless of whether it is "
        "marked confidential, and regardless of whether it qualifies as a trade secret under "
        "applicable law. Employee shall not disclose any confidential information to any third "
        "party without the Company's prior written consent.",
    )
    _clause(
        doc,
        "7. Return of Property",
        "Upon termination of employment, Employee shall immediately return to the Company all "
        "documents, equipment, source code, notes, and other property or information belonging "
        "to the Company, without retention of any copies in any form or medium.",
    )
    _clause(
        doc,
        "8. Governing Law and Dispute Resolution",
        "This Agreement shall be governed by the laws of the State of California, without regard "
        "to conflict-of-law principles. The parties submit to the exclusive jurisdiction of the "
        "state and federal courts located in San Francisco County, California for resolution of "
        "any dispute arising from this Agreement.",
    )
    return doc


def build_data_processing_agreement() -> Document:
    """A data processing agreement appended to a vendor services contract, with processor-
    favourable terms and missing GDPR controls. Plants five major issues: (a) processor may
    appoint sub-processors at will with no notice or objection right; (b) no deadline for
    breach notification (only vague \"without undue delay\"); (c) no audit or inspection right
    for the controller; (d) data destruction at processor's discretion on termination; (e) no
    transfer mechanism for international data flows."""
    doc = Document()
    _heading(doc, "DATA PROCESSING AGREEMENT")
    doc.add_paragraph(
        "This Data Processing Agreement (this \"DPA\") is incorporated by reference into that "
        "certain Master Services Agreement dated September 1, 2026, by and between Prism "
        "Analytics Inc., a Delaware corporation (\"Controller\"), and Cascade Data Services Ltd., "
        "a corporation organized under the laws of the Cayman Islands (\"Processor\"). This DPA "
        "governs the processing of personal data by Processor on behalf of Controller."
    )
    _clause(
        doc,
        "1. Definitions",
        "As used in this DPA, terms shall have the meanings ascribed to them under the General "
        "Data Protection Regulation (GDPR) and, where applicable, other applicable data "
        "protection laws. \"Personal Data\" means any information relating to an identified or "
        "identifiable natural person. \"Processing\" means any operation performed on personal "
        "data. \"Processor\" means Cascade Data Services Ltd., which processes personal data on "
        "behalf of Controller. \"Sub-processor\" means any entity engaged by Processor to process "
        "personal data on behalf of Controller.",
    )
    _clause(
        doc,
        "2. Scope of Processing",
        "Processor shall process personal data only in accordance with the written instructions "
        "of Controller and only for the purpose of providing the services described in the "
        "Master Services Agreement. Processor shall not process personal data for any other purpose. "
        "Processing shall be limited to the categories of personal data, types of data subjects, "
        "types of processing, and duration as specified in the Data Processing Appendix attached "
        "as Exhibit A to this DPA.",
    )
    _clause(
        doc,
        "3. Sub-processor Appointment",
        "Processor may engage sub-processors at any time and at its sole discretion to perform "
        "processing activities, without prior notice to or consent from Controller. Processor "
        "shall not be required to inform Controller of the identity of any sub-processor, nor "
        "shall Controller have any right to object to or veto the appointment of sub-processors. "
        "Processor remains fully liable to Controller for performance of any sub-processor engaged.",
    )
    _clause(
        doc,
        "4. Security and Confidentiality",
        "Processor shall implement and maintain appropriate technical and organizational measures "
        "to ensure a level of security appropriate to the risk, including protection against "
        "unauthorized or unlawful processing and against accidental loss, destruction, or damage. "
        "Processor shall ensure that any persons authorized to process personal data on behalf "
        "of Processor have committed to confidentiality or are under an appropriate statutory "
        "obligation of confidentiality.",
    )
    _clause(
        doc,
        "5. Data Subject Rights",
        "Processor shall, taking into account the nature of processing, assist Controller by "
        "appropriate technical and organizational measures in fulfilling Controller's obligation "
        "to respond to data subject requests for access, correction, erasure, and data portability. "
        "Processor shall promptly inform Controller of any data subject request received and shall "
        "not respond directly to the data subject without Controller's prior written authorization.",
    )
    _clause(
        doc,
        "6. Data Breach Notification",
        "Processor shall notify Controller of any personal data breach without undue delay. "
        "No specific timeframe or deadline is specified for such notification. Processor shall "
        "provide such information as is available to Processor at the time of notification, and "
        "shall comply with any further instructions from Controller regarding the incident response. "
        "Processor bears no liability for delays in notification or for incomplete information.",
    )
    _clause(
        doc,
        "7. Audit and Inspection Rights",
        "Processor acknowledges that Controller is responsible for ensuring compliance with data "
        "protection laws. Processor shall make available to Controller only such information as "
        "Processor determines to be relevant in Processor's sole discretion. Controller shall have "
        "no independent right to audit, inspect, or examine Processor's facilities, systems, "
        "processes, or records. Any inquiries from Controller regarding Processor's processing "
        "activities shall be answered at Processor's option and in Processor's sole discretion.",
    )
    _clause(
        doc,
        "8. Data Return and Deletion",
        "Upon termination of the Master Services Agreement or upon Controller's request, Processor "
        "shall, at Processor's sole discretion, either return all personal data to Controller or "
        "delete all personal data in Processor's possession. Processor is not required to certify "
        "completion of deletion or return. Processor may retain personal data as required by "
        "applicable law or for its own business purposes, and Processor may continue to use "
        "anonymized or aggregated data derived from the personal data without restriction.",
    )
    _clause(
        doc,
        "9. International Transfers",
        "Processor may transfer personal data to any jurisdiction, including jurisdictions that "
        "have not been determined to provide an adequate level of data protection, at any time "
        "and for any reason. No adequacy determination, Standard Contractual Clause, Binding "
        "Corporate Rule, or other transfer mechanism is required. Processor shall not be obligated "
        "to inform Controller of transfers or to implement any safeguard for such transfers. "
        "Controller consents to such transfers by entering into this DPA.",
    )
    _clause(
        doc,
        "10. Governing Law",
        "This DPA shall be governed by and construed in accordance with the laws of the Cayman "
        "Islands, without regard to conflict-of-law principles. Any dispute shall be resolved "
        "through binding arbitration administered by the Cayman Islands Arbitration Association.",
    )
    return doc


def build_consultancy_agreement_short() -> Document:
    """A deliberately thin two-page consultancy agreement, intentionally missing key protective
    clauses to exercise the \"required clause absent\" detection. No limitation of liability,
    no governing law, no dispute resolution, no IP ownership clause, and vague payment terms.
    The document is sparse and genuine rather than padded."""
    doc = Document()
    _heading(doc, "CONSULTANCY SERVICES AGREEMENT")
    doc.add_paragraph(
        "This Consultancy Services Agreement (this \"Agreement\") is entered into between Nexus "
        "Consulting Partners, a partnership organized in New York (\"Client\"), and Dr. Elena "
        "Rothschild, an independent consultant (\"Consultant\")."
    )
    _clause(
        doc,
        "1. Services",
        "Consultant shall provide strategic business consulting services as requested by Client "
        "from time to time. The scope and nature of services shall be determined by mutual "
        "agreement between the parties on a project basis.",
    )
    _clause(
        doc,
        "2. Engagement Term",
        "This Agreement shall commence on the date first written above and shall continue until "
        "either party terminates the engagement by providing thirty (30) days' written notice to "
        "the other party.",
    )
    _clause(
        doc,
        "3. Compensation",
        "Client shall pay Consultant fees for services upon satisfactory completion of the "
        "consulting work. The amount and schedule of fees shall be mutually agreed upon in writing "
        "prior to commencement of each engagement.",
    )
    _clause(
        doc,
        "4. Confidentiality",
        "Consultant shall maintain the confidentiality of all Client information disclosed during "
        "the engagement and shall not disclose such information to any third party.",
    )
    _clause(
        doc,
        "5. Independent Contractor",
        "Consultant is an independent contractor and is not an employee of Client. Consultant is "
        "responsible for all taxes, insurance, and other statutory obligations related to the "
        "compensation received.",
    )
    _clause(
        doc,
        "6. Termination",
        "Either party may terminate this Agreement upon thirty (30) days' written notice to the "
        "other party.",
    )
    return doc


def main() -> None:
    _DATASET_DIR.mkdir(parents=True, exist_ok=True)
    targets = {
        "employment_agreement_senior_engineer.docx": build_employment_agreement_senior_engineer,
        "data_processing_agreement.docx": build_data_processing_agreement,
        "consultancy_agreement_short.docx": build_consultancy_agreement_short,
    }
    for filename, builder in targets.items():
        path = _DATASET_DIR / filename
        builder().save(str(path))
        print(f"wrote {path.relative_to(_REPO_ROOT)}")


if __name__ == "__main__":
    main()
