"""Generate synthetic NDA samples for the Legal Helper teaching workshop (§6 Phase 2).

Three documents, each built to exercise a specific reviewer behavior: a market-standard
mutual NDA with no issues (GREEN baseline), a one-way receiving-party-hostile NDA with
aggressive carve-outs and no caps (RED case), and a vendor-evaluation mutual with moderate
7-year confidentiality and a unilateral residuals clause (YELLOW case). Run with
``python dataset/generators/ndas.py`` from the repo root.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

_REPO_ROOT = Path(__file__).resolve().parents[2]
_DATASET_DIR = _REPO_ROOT / "dataset"


def _heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def _clause(doc: Document, heading: str, body: str) -> None:
    doc.add_paragraph(heading, style="Heading 2")
    doc.add_paragraph(body)


def build_nda_mutual_balanced() -> Document:
    """A clean, market-standard mutual NDA between two invented companies with standard
    protections, carve-outs, reasonable terms, and mutual limitations of liability.
    This is the baseline GREEN case — nothing to flag."""
    doc = Document()
    _heading(doc, "MUTUAL NON-DISCLOSURE AGREEMENT")
    doc.add_paragraph(
        "This Mutual Non-Disclosure Agreement (this \"Agreement\") is entered into as of the "
        "date first written below (the \"Effective Date\"), between Pinnacle Analytics Inc., "
        "a Delaware corporation (\"Company A\"), and CloudSync Solutions Ltd., a British Columbia "
        "corporation (\"Company B\"), collectively referred to as the \"Parties.\""
    )
    doc.add_paragraph(
        "WHEREAS, the Parties desire to explore a potential business opportunity and to exchange "
        "certain confidential, proprietary, and sensitive information for the purpose of evaluating "
        "such opportunity (the \"Purpose\"); and"
    )
    doc.add_paragraph(
        "WHEREAS, the Parties wish to protect such information and establish procedures for the "
        "use and disclosure of the same;"
    )
    doc.add_paragraph("NOW, THEREFORE, the Parties agree as follows:")

    _clause(
        doc,
        "1. Definition of Confidential Information",
        "For purposes of this Agreement, \"Confidential Information\" means any non-public, "
        "proprietary, or trade secret information (whether technical, commercial, financial, or "
        "strategic in nature) disclosed by one Party (the \"Disclosing Party\") to the other "
        "(the \"Receiving Party\"), regardless of the form or media of disclosure (written, oral, "
        "visual inspection, electronic, or otherwise), and whether or not marked as confidential. "
        "Confidential Information includes but is not limited to product plans, software code, "
        "algorithms, business strategies, customer lists, financial information, technical data, "
        "pricing information, and business processes. Confidential Information does NOT include "
        "any information that: (a) is or becomes publicly available through no breach of this "
        "Agreement by the Receiving Party; (b) was independently known to the Receiving Party "
        "without obligation of confidentiality, as evidenced by prior written records; (c) is "
        "independently developed by the Receiving Party without access to or use of the "
        "Confidential Information, as evidenced by written records; or (d) is rightfully received "
        "by the Receiving Party from a third party without restriction on disclosure, provided that "
        "the Receiving Party makes reasonable inquiry to determine whether the third party obtained "
        "it lawfully."
    )

    _clause(
        doc,
        "2. Obligations of the Receiving Party",
        "Each Receiving Party shall: (a) maintain the confidentiality of the Disclosing Party's "
        "Confidential Information using the same degree of care it uses to protect its own "
        "confidential information, but in no event less than reasonable care; (b) use the "
        "Confidential Information solely for the Purpose and not for any other business purpose; "
        "(c) not disclose the Confidential Information to any third party except to employees, "
        "contractors, and professional advisors (including legal, financial, and technical "
        "advisors) on a need-to-know basis, provided that such persons are bound by written "
        "confidentiality obligations at least as protective as those contained herein and the "
        "Receiving Party remains responsible for their compliance; (d) not reverse engineer, "
        "disassemble, or attempt to derive the underlying concepts or source code of any technical "
        "Confidential Information without the prior written consent of the Disclosing Party; and "
        "(e) promptly notify the Disclosing Party in writing of any unauthorized access to, use of, "
        "or disclosure of the Confidential Information."
    )

    _clause(
        doc,
        "3. Permitted Disclosures",
        "Notwithstanding Section 2, if the Receiving Party is required by applicable law, regulation, "
        "court order, or regulatory authority to disclose any Confidential Information, the Receiving "
        "Party shall, to the extent permitted by law: (a) promptly notify the Disclosing Party of "
        "such requirement in writing; (b) cooperate with the Disclosing Party in seeking a protective "
        "order or other remedy; and (c) disclose only the minimum Confidential Information legally "
        "required to be disclosed."
    )

    _clause(
        doc,
        "4. Term and Survival",
        "This Agreement shall remain in effect for three (3) years from the Effective Date (the "
        "\"Term\"), unless earlier terminated by either Party upon thirty (30) days' written notice. "
        "The confidentiality obligations set forth in Section 2 shall survive termination of this "
        "Agreement for an additional three (3) years, except that obligations with respect to "
        "Confidential Information that constitutes a trade secret shall continue for so long as such "
        "information remains a trade secret under applicable law."
    )

    _clause(
        doc,
        "5. Return or Destruction of Confidential Information",
        "Upon the Disclosing Party's written request or upon termination of this Agreement, the "
        "Receiving Party shall, at the Disclosing Party's election: (a) return all Confidential "
        "Information in tangible form (and all copies thereof) in the Receiving Party's possession "
        "or control; or (b) certify in writing that all such Confidential Information has been "
        "securely destroyed or deleted. Notwithstanding the foregoing, the Receiving Party may "
        "retain one archival copy of the Confidential Information solely for the purpose of "
        "demonstrating compliance with this Agreement, to be stored securely and subject to continued "
        "confidentiality obligations."
    )

    _clause(
        doc,
        "6. No License or Obligation to Proceed",
        "Nothing in this Agreement grants either Party any license, right, or obligation with respect "
        "to the other Party's Confidential Information, patents, copyrights, trademarks, trade secrets, "
        "or other intellectual property. Neither Party is obligated by this Agreement to enter into any "
        "business transaction or relationship with the other Party. Either Party may terminate its "
        "participation in discussions regarding the Purpose at any time without penalty or liability."
    )

    _clause(
        doc,
        "7. No Warranty",
        "EACH PARTY DISCLAIMS ANY WARRANTY, EXPRESS OR IMPLIED, WITH RESPECT TO THE CONFIDENTIAL "
        "INFORMATION, INCLUDING WITHOUT LIMITATION ANY WARRANTY OF MERCHANTABILITY, FITNESS FOR A "
        "PARTICULAR PURPOSE, OR NON-INFRINGEMENT. EACH PARTY PROVIDES ITS CONFIDENTIAL INFORMATION "
        "\"AS IS\" WITHOUT WARRANTY OF ANY KIND."
    )

    _clause(
        doc,
        "8. Limitation of Liability",
        "IN NO EVENT SHALL EITHER PARTY BE LIABLE TO THE OTHER FOR ANY INDIRECT, INCIDENTAL, SPECIAL, "
        "CONSEQUENTIAL, OR PUNITIVE DAMAGES ARISING OUT OF OR RELATED TO THIS AGREEMENT, INCLUDING "
        "WITHOUT LIMITATION LOST PROFITS, LOSS OF DATA, OR BUSINESS INTERRUPTION, EVEN IF ADVISED OF "
        "THE POSSIBILITY OF SUCH DAMAGES. EACH PARTY'S TOTAL LIABILITY ARISING OUT OF OR RELATED TO "
        "THIS AGREEMENT SHALL NOT EXCEED ONE THOUSAND DOLLARS ($1,000.00)."
    )

    _clause(
        doc,
        "9. No Proprietary Relationship",
        "Nothing in this Agreement creates a partnership, joint venture, agency, or employment "
        "relationship between the Parties. Neither Party has authority to bind, represent, or act on "
        "behalf of the other Party except as expressly authorized in writing."
    )

    _clause(
        doc,
        "10. Assignment",
        "Neither Party may assign this Agreement or any rights or obligations hereunder without the "
        "prior written consent of the other Party, which consent shall not be unreasonably withheld. "
        "Notwithstanding the foregoing, either Party may assign this Agreement without consent to an "
        "affiliate, or in connection with a merger, acquisition, or sale of substantially all of its "
        "assets, provided that the assigning Party remains liable for all its obligations hereunder "
        "and the assignee agrees in writing to be bound by the terms of this Agreement."
    )

    _clause(
        doc,
        "11. Governing Law and Jurisdiction",
        "This Agreement shall be governed by and construed in accordance with the laws of the State of "
        "California, without regard to its conflicts of law principles. Each Party irrevocably consents "
        "to the exclusive jurisdiction of the state and federal courts located in San Francisco County, "
        "California, and waives any objection to venue or inconvenient forum in those courts."
    )

    _clause(
        doc,
        "12. Equitable Relief",
        "Each Party acknowledges that a breach of this Agreement may cause irreparable harm for which "
        "monetary damages would be an inadequate remedy. Accordingly, in addition to any other remedies "
        "available at law or in equity, either Party shall be entitled to seek injunctive or other "
        "equitable relief to prevent or remedy any breach of this Agreement."
    )

    _clause(
        doc,
        "13. Entire Agreement",
        "This Agreement constitutes the entire agreement between the Parties with respect to the "
        "subject matter hereof and supersedes all prior negotiations, understandings, and agreements, "
        "whether written or oral. This Agreement may not be amended or modified except by a written "
        "instrument signed by authorized representatives of both Parties."
    )

    _clause(
        doc,
        "14. Severability",
        "If any provision of this Agreement is held by a court of competent jurisdiction to be invalid, "
        "illegal, or unenforceable, such provision shall be modified to the minimum extent necessary to "
        "make it enforceable, or if such modification is not possible, such provision shall be severed, "
        "and the remaining provisions shall remain in full force and effect."
    )

    _clause(
        doc,
        "15. Counterparts and Electronic Signatures",
        "This Agreement may be executed in multiple counterparts, each of which shall be deemed an "
        "original and all of which together shall constitute one and the same instrument. Execution "
        "and delivery of this Agreement by electronic facsimile, PDF, or other electronic means shall "
        "have the same force and effect as delivery of manually executed originals."
    )

    return doc


def build_nda_oneway_receiving_party_hostile() -> Document:
    """A one-way NDA aggressively structured against the Receiving Party. Plants specific
    problems: perpetual confidentiality on all info (no just-trade-secrets), zero carve-outs,
    24-month non-solicitation of employees and customers, one-way injunctive relief with
    receiver pays legal fees, and no liability cap. This is the RED case."""
    doc = Document()
    _heading(doc, "NON-DISCLOSURE AGREEMENT")
    doc.add_paragraph(
        "This Non-Disclosure Agreement (this \"Agreement\") is entered into as of the date "
        "first written below, between Nexus Capital Ventures LLC, a Delaware limited liability "
        "company (\"Disclosing Party\"), and TechVenture Partners, Inc., a Delaware corporation "
        "(\"Receiving Party\")."
    )
    doc.add_paragraph(
        "WHEREAS, the Disclosing Party is willing to disclose certain Confidential Information "
        "to the Receiving Party for the purpose of exploring a potential investment or business "
        "transaction (the \"Purpose\"); and"
    )
    doc.add_paragraph("NOW, THEREFORE, the Parties agree as follows:")

    _clause(
        doc,
        "1. Definition of Confidential Information",
        "\"Confidential Information\" means ALL information and materials disclosed by the "
        "Disclosing Party to the Receiving Party, whether in written, oral, visual, electronic, "
        "or any other form whatsoever, including but not limited to business plans, financial "
        "projections, customer lists, supplier lists, pricing strategies, software, source code, "
        "algorithms, manufacturing processes, research and development activities, market data, "
        "product specifications, business models, and any other information relating to the "
        "Disclosing Party's business. All Confidential Information shall be deemed confidential "
        "and proprietary regardless of whether marked as such. Confidential Information includes "
        "information disclosed orally if the Disclosing Party identifies it as confidential "
        "at the time of oral disclosure."
    )

    _clause(
        doc,
        "2. Obligations of Receiving Party",
        "The Receiving Party shall: (a) treat all Confidential Information as strictly confidential "
        "and proprietary; (b) use the Confidential Information solely for the Purpose; (c) not "
        "disclose any Confidential Information to any third party whatsoever without the prior "
        "written consent of the Disclosing Party, which consent may be withheld in the Disclosing "
        "Party's sole and absolute discretion; (d) ensure that any employees or agents who access "
        "the Confidential Information are bound by written confidentiality agreements at least as "
        "stringent as the terms of this Agreement; (e) not reverse engineer, decompile, or attempt "
        "to derive any trade secrets or proprietary methods from the Confidential Information; (f) "
        "implement and maintain security measures adequate to protect the Confidential Information "
        "from unauthorized access or theft; and (g) keep detailed written records of all persons who "
        "access the Confidential Information, which records shall be made available to the Disclosing "
        "Party upon request."
    )

    _clause(
        doc,
        "3. No Exceptions or Carve-Outs",
        "The Receiving Party's confidentiality obligations shall apply to all Confidential Information "
        "without exception. There are no exceptions to the Receiving Party's obligations based on the "
        "information becoming public, being independently developed, or being received from a third "
        "party. The Receiving Party remains bound by this Agreement regardless of any circumstances."
    )

    _clause(
        doc,
        "4. Non-Solicitation Covenant",
        "During the term of this Agreement and for a period of twenty-four (24) months following "
        "termination or completion of the discussions contemplated by the Purpose (whichever is later), "
        "the Receiving Party shall not, directly or indirectly: (a) solicit, recruit, or encourage any "
        "current or former employee, contractor, or agent of the Disclosing Party to leave their "
        "employment or engagement with the Disclosing Party; (b) solicit, contact, or do business with "
        "any customer, client, or prospective customer of the Disclosing Party or any entity in which "
        "the Disclosing Party has a business interest; (c) encourage any supplier or vendor of the "
        "Disclosing Party to cease supplying goods or services to the Disclosing Party; or (d) enter "
        "into any business arrangement with any employee, customer, or vendor of the Disclosing Party "
        "for any competitive or overlapping purpose."
    )

    _clause(
        doc,
        "5. Term",
        "This Agreement shall commence on the Effective Date and shall continue in perpetuity unless "
        "terminated by mutual written agreement of the Parties. The Receiving Party's obligations with "
        "respect to the Confidential Information shall be perpetual and shall survive any termination "
        "or completion of discussions with respect to the Purpose. There is no expiration date on the "
        "Receiving Party's confidentiality obligations."
    )

    _clause(
        doc,
        "6. Return and Certification",
        "Upon demand by the Disclosing Party or upon termination of discussions, the Receiving Party "
        "shall immediately return or certify the destruction of all Confidential Information. The "
        "Receiving Party shall provide the Disclosing Party with a written certification of destruction "
        "within five (5) business days. Failure to provide such certification may be deemed evidence of "
        "retention."
    )

    _clause(
        doc,
        "7. Injunctive Relief and Legal Fees",
        "The Receiving Party acknowledges that any breach of this Agreement may cause irreparable harm "
        "that cannot be remedied by monetary damages alone. Accordingly, in addition to any other remedies "
        "available to the Disclosing Party, the Disclosing Party shall be entitled to seek and obtain "
        "injunctive relief, temporary restraining orders, and specific performance to prevent breaches of "
        "this Agreement. Furthermore, in the event of any breach of this Agreement by the Receiving Party, "
        "the Receiving Party shall be responsible for and shall reimburse the Disclosing Party for all "
        "costs and expenses incurred by the Disclosing Party in enforcing its rights hereunder, including "
        "without limitation reasonable attorneys' fees, expert fees, court costs, and investigative costs, "
        "regardless of whether legal proceedings are commenced or any judgment is obtained."
    )

    _clause(
        doc,
        "8. No Limitation of Liability",
        "THERE IS NO CAP, LIMIT, OR LIMITATION ON THE LIABILITY OF THE RECEIVING PARTY FOR BREACHES OF "
        "THIS AGREEMENT. THE RECEIVING PARTY SHALL BE LIABLE FOR ALL DIRECT, INDIRECT, INCIDENTAL, "
        "SPECIAL, CONSEQUENTIAL, AND PUNITIVE DAMAGES ARISING FROM ANY BREACH OF THIS AGREEMENT, WITHOUT "
        "LIMITATION. THIS INCLUDES BUT IS NOT LIMITED TO LOST PROFITS, LOSS OF BUSINESS OPPORTUNITY, "
        "DIMINUTION OF VALUE, BUSINESS INTERRUPTION, AND ANY OTHER DAMAGES WHATSOEVER."
    )

    _clause(
        doc,
        "9. No License; No Obligation",
        "Nothing in this Agreement grants the Receiving Party any license, right, or permission with "
        "respect to the Confidential Information or any intellectual property of the Disclosing Party. "
        "The Receiving Party's access to the Confidential Information shall not be deemed a basis for "
        "any claim of entitlement or future business relationship. The Disclosing Party retains all "
        "rights to the Confidential Information and may use it in any manner it sees fit."
    )

    _clause(
        doc,
        "10. Entire Agreement",
        "This Agreement constitutes the entire agreement between the Parties with respect to the "
        "subject matter and supersedes all prior understandings and agreements. No amendment or "
        "modification of this Agreement shall be effective unless made in writing and signed by an "
        "authorized representative of the Disclosing Party."
    )

    return doc


def build_nda_vendor_evaluation_moderate() -> Document:
    """A mutual NDA with moderate, arguable issues: seven-year confidentiality term, a broad
    residuals clause, no data-protection/personal-data handling clause, and unilateral assignment
    rights for one party only. This is the YELLOW case."""
    doc = Document()
    _heading(doc, "MUTUAL CONFIDENTIALITY AND RESIDUALS AGREEMENT")
    doc.add_paragraph(
        "This Mutual Confidentiality and Residuals Agreement (this \"Agreement\") is entered into "
        "as of the date first written below (the \"Effective Date\"), between InnovateCorp Solutions, "
        "a Delaware corporation (\"Company A\"), and VendorPro Consulting, a state corporation "
        "(\"Company B\"), collectively referred to as the \"Parties.\""
    )
    doc.add_paragraph(
        "WHEREAS, the Parties wish to share certain confidential business and technical information "
        "to evaluate a potential vendor relationship and integration of services (the \"Purpose\"); and"
    )
    doc.add_paragraph("NOW, THEREFORE, the Parties agree as follows:")

    _clause(
        doc,
        "1. Confidential Information",
        "\"Confidential Information\" means any proprietary or non-public information disclosed by "
        "one Party to the other in connection with the Purpose, including but not limited to business "
        "plans, technical specifications, product roadmaps, customer information, pricing models, "
        "software code, and business strategies. Confidential Information does not include information "
        "that is or becomes publicly available through no breach of this Agreement, was already known "
        "to the Receiving Party without restriction, or is independently developed without use of the "
        "Confidential Information."
    )

    _clause(
        doc,
        "2. Use and Protection",
        "Each Party shall protect the other Party's Confidential Information using reasonable care and "
        "shall use it solely for the Purpose. Neither Party shall disclose the Confidential Information "
        "to third parties except to employees and advisors on a need-to-know basis who are bound by "
        "confidentiality obligations. Each Party shall be liable for breaches by its employees and "
        "advisors."
    )

    _clause(
        doc,
        "3. Residuals",
        "Notwithstanding any other provision in this Agreement, the Parties acknowledge and agree that "
        "employees and advisors of each Party may retain in unaided memory general ideas, concepts, "
        "know-how, and techniques concerning the Confidential Information (\"Residuals\"). Either Party "
        "may use Residuals without restriction or obligation to the other Party, provided that such use "
        "shall not violate any third-party intellectual property rights or constitute an unauthorized "
        "use of the other Party's trade secrets. For clarity, Residuals include any concepts, ideas, or "
        "mental impressions retained in the unaided memories of employees or agents, and neither Party "
        "has any obligation to account for or restrict such residual knowledge."
    )

    _clause(
        doc,
        "4. Term",
        "This Agreement remains in effect for seven (7) years from the Effective Date, unless earlier "
        "terminated by mutual written agreement. The confidentiality obligations shall survive "
        "termination for the duration of the seven-year term only."
    )

    _clause(
        doc,
        "5. Return of Information",
        "Upon written request or termination, each Party shall return or destroy Confidential "
        "Information at its discretion, except that each Party may retain one copy for archival "
        "compliance purposes."
    )

    _clause(
        doc,
        "6. Governing Law",
        "This Agreement is governed by the laws of New York, without regard to conflicts of law. "
        "The Parties consent to the jurisdiction of the courts of New York County."
    )

    _clause(
        doc,
        "7. Assignment",
        "Company B may assign this Agreement to any affiliate or in connection with a sale or merger "
        "without the consent of Company A, provided that Company B remains liable for performance. "
        "Company A may not assign this Agreement without the prior written consent of Company B."
    )

    _clause(
        doc,
        "8. Warranty Disclaimer",
        "Each Party provides its Confidential Information \"AS IS\" without warranty of any kind, "
        "express or implied."
    )

    _clause(
        doc,
        "9. Limitation of Liability",
        "Neither Party shall be liable for indirect, consequential, or special damages arising from "
        "this Agreement, except in cases of willful breach or violation of law. Each Party's liability "
        "shall not exceed the value of any fees paid or payable under this Agreement."
    )

    _clause(
        doc,
        "10. Entire Agreement",
        "This Agreement constitutes the entire understanding of the Parties regarding confidentiality "
        "and supersedes all prior agreements on this subject matter."
    )

    return doc


def main() -> None:
    _DATASET_DIR.mkdir(parents=True, exist_ok=True)
    (_DATASET_DIR / "generators").mkdir(parents=True, exist_ok=True)

    targets = {
        "nda_mutual_balanced.docx": build_nda_mutual_balanced,
        "nda_oneway_receiving_party_hostile.docx": build_nda_oneway_receiving_party_hostile,
        "nda_vendor_evaluation_moderate.docx": build_nda_vendor_evaluation_moderate,
    }
    for filename, builder in targets.items():
        path = _DATASET_DIR / filename
        builder().save(str(path))
        print(f"wrote {path.relative_to(_REPO_ROOT)}")


if __name__ == "__main__":
    main()
