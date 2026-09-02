"""Generate synthetic demo commercial contracts for Legal Helper (teaching workshop material).

Three documents with deliberate planted issues:
  - msa_professional_services.docx: pro-customer bias (supplier red flags)
  - saas_subscription_agreement.docx: pro-vendor bias (customer red flags)
  - reseller_distribution_agreement.docx: moderate issues (yellow flags)

Run with: cd backend && .venv/bin/python ../dataset/generators/commercial.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

_REPO_ROOT = Path(__file__).resolve().parents[2]
_DATASET_DIR = _REPO_ROOT / "dataset"


def _heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def _clause(doc: Document, heading: str, body: str) -> None:
    doc.add_paragraph(heading, style="Heading 2")
    doc.add_paragraph(body)


def build_msa_professional_services() -> Document:
    """A Master Services Agreement for professional/consulting services, drafted heavily in the
    CUSTOMER's favour so the supplier side has multiple red flags to catch: one-sided indemnity
    covering the customer's negligence, IP assignment of pre-existing supplier background IP,
    90-day payment terms with unilateral customer set-off rights, and termination for convenience
    on 5 days with no payment for work in progress."""
    doc = Document()
    _heading(doc, "MASTER SERVICES AGREEMENT FOR PROFESSIONAL SERVICES")
    doc.add_paragraph(
        "This Master Services Agreement (\"Agreement\") is entered into effective as of the date "
        "of final execution by and between Pinnacle Consulting Group, LLC, a limited liability "
        "company (\"Customer\"), and Velocity Solutions, Inc., a Delaware corporation "
        "(\"Supplier\"), for the provision of professional consulting and implementation services."
    )
    _clause(
        doc,
        "1. Services",
        "Supplier shall provide the professional services, staff augmentation, consulting advice, "
        "and implementation support described in one or more Statements of Work (each, an \"SOW\") "
        "executed under this Agreement (collectively, the \"Services\"). Each SOW shall specify "
        "the scope of Services, deliverables, timeline, and resource allocation. Services shall "
        "be performed in a professional manner consistent with generally accepted industry "
        "standards for similar engagements. Supplier shall dedicate the personnel specified in "
        "the applicable SOW and shall ensure such personnel possess the skills, qualifications, "
        "and certifications necessary to perform the Services competently.",
    )
    _clause(
        doc,
        "2. Fees and Payment Terms",
        "Customer shall pay Supplier for Services in accordance with the fees and expenses set "
        "forth in the applicable SOW. Unless otherwise specified in an SOW, invoices are due and "
        "payable within ninety (90) days of receipt by Customer. Notwithstanding any other "
        "provision, Customer shall have the unilateral right to set off against any amounts due "
        "to Supplier any amounts Customer claims Supplier owes Customer or any affiliate of "
        "Customer for any reason whatsoever, including disputed claims or anticipated damages, "
        "without limiting Customer's other rights or remedies. Supplier waives any right to "
        "object to or challenge any set-off applied by Customer.",
    )
    _clause(
        doc,
        "3. Intellectual Property",
        "All intellectual property, including but not limited to inventions, works of authorship, "
        "software code, methodologies, frameworks, templates, pre-existing materials, background "
        "IP, tools, utilities, and any other intellectual property developed, conceived, created, "
        "or used by Supplier in connection with the Services or any SOW shall be the sole and "
        "exclusive property of Customer, and Supplier hereby assigns all right, title, and "
        "interest in and to any such intellectual property to Customer. This includes, expressly "
        "and without limitation, any pre-existing intellectual property owned or controlled by "
        "Supplier before this Agreement or any SOW, reusable tools, frameworks, or methodologies "
        "developed by Supplier, and any background or foundational IP that Supplier may have "
        "developed for use across multiple engagements. Supplier retains no rights to use, "
        "sublicense, or exploit any such intellectual property except as necessary to perform the "
        "Services during the Term of this Agreement and agrees not to use any such intellectual "
        "property in any other engagement or for any other purpose.",
    )
    _clause(
        doc,
        "4. Indemnification",
        "Supplier shall indemnify, defend with counsel reasonably acceptable to Customer (at "
        "Supplier's sole expense), and hold harmless Customer and its affiliates, officers, "
        "directors, employees, and agents from and against any and all claims, damages, losses, "
        "liabilities, costs, and expenses (including reasonable attorneys' fees and court costs) "
        "arising out of or related to: (a) Supplier's performance or non-performance of the "
        "Services; (b) any breach of this Agreement by Supplier; (c) Supplier's negligence, "
        "gross negligence, or willful misconduct; (d) any injury to persons or damage to property "
        "caused by Supplier; (e) Supplier's infringement or alleged infringement of any third-party "
        "intellectual property rights; (f) the use of the Services by Customer; (g) Customer's "
        "reliance on advice, recommendations, or work product provided by Supplier; and (h) any "
        "act or omission by Supplier, including negligence or actions taken by Customer relying "
        "on Supplier's professional judgment, regardless of whether such act or omission was "
        "foreseeable or whether Supplier was advised of the possibility of such damages.",
    )
    _clause(
        doc,
        "5. Limitation of Liability",
        "In no event shall Supplier's total liability arising out of or related to this Agreement "
        "or any SOW exceed the lesser of: (i) the total fees paid by Customer to Supplier in the "
        "twelve (12) months preceding the claim, or (ii) fifty thousand dollars ($50,000). "
        "Notwithstanding the foregoing, Supplier's liability shall NOT be limited with respect "
        "to claims brought by Customer, indemnification obligations to Customer, breaches of "
        "Customer's intellectual property rights, or any claims arising from Supplier's gross "
        "negligence or willful misconduct.",
    )
    _clause(
        doc,
        "6. Term and Termination for Convenience",
        "This Agreement shall commence on the Effective Date and continue until terminated by "
        "either party. Customer may terminate this Agreement and any outstanding SOW at any time, "
        "for any reason or no reason, upon five (5) business days' written notice to Supplier. "
        "Upon such termination by Customer, Supplier shall immediately cease work and shall "
        "receive payment only for Services actually completed and accepted by Customer as of the "
        "termination date. Supplier shall receive NO compensation for work in progress, partially "
        "completed deliverables, mobilization costs, demobilization costs, or any other expenses "
        "or anticipated profits related to work not fully completed and accepted. Supplier may "
        "not terminate this Agreement except for Customer's uncured material breach, which must "
        "be provided in writing with at least sixty (60) days' notice and opportunity to cure.",
    )
    _clause(
        doc,
        "7. Confidentiality",
        "Each party shall maintain the confidentiality of the other party's Confidential "
        "Information using at least the same degree of care it uses for its own confidential "
        "information and shall not disclose such information to any third party without the other "
        "party's prior written consent, except to employees and advisors with a legitimate need "
        "to know who are bound by confidentiality obligations. Confidential Information shall not "
        "include information that is or becomes publicly available through no breach of this "
        "Agreement.",
    )
    _clause(
        doc,
        "8. Data Protection and Privacy",
        "If Supplier has access to any personal data or information subject to data protection "
        "laws, Supplier shall comply with all applicable laws and regulations. Supplier shall "
        "implement and maintain appropriate technical and organizational security measures to "
        "protect such data from unauthorized access or processing.",
    )
    _clause(
        doc,
        "9. Independent Contractor",
        "Supplier is an independent contractor and is not an employee, agent, or representative "
        "of Customer. Nothing in this Agreement creates a partnership, joint venture, or agency "
        "relationship between the parties. Supplier is solely responsible for all employment-related "
        "taxes, benefits, workers' compensation, and other statutory obligations.",
    )
    _clause(
        doc,
        "10. Governing Law and Dispute Resolution",
        "This Agreement shall be governed by and construed in accordance with the laws of the "
        "State of California, without regard to its conflict-of-laws principles. Any dispute "
        "arising out of or related to this Agreement shall be subject to the exclusive jurisdiction "
        "of the state and federal courts located in San Francisco County, California, and both "
        "parties irrevocably consent to such jurisdiction.",
    )
    return doc


def build_saas_subscription_agreement() -> Document:
    """A SaaS subscription agreement from the vendor's perspective with multiple red flags for
    the customer: auto-renewal without opt-out reminder, uncapped fee increases at renewal,
    liability cap at one month's fees with data loss excluded, minimal data protection with no
    sub-processor notice, and unilateral modification rights by the vendor."""
    doc = Document()
    _heading(doc, "SAAS SUBSCRIPTION AGREEMENT")
    doc.add_paragraph(
        "This Software-as-a-Service Subscription Agreement (\"Agreement\") is entered into between "
        "CloudVault Analytics, Inc., a California corporation (\"Provider\"), and the entity "
        "executing this Agreement (\"Customer\"), effective as of the date of acceptance "
        "(\"Effective Date\")."
    )
    _clause(
        doc,
        "1. Subscription Services",
        "Provider shall make available to Customer a cloud-based software service for data "
        "analytics and business intelligence (the \"Service\") accessible via the internet at the "
        "Service URL specified by Provider. The Service includes all features, functionality, "
        "updates, bug fixes, and patches provided by Provider at its discretion. Customer "
        "acknowledges that the Service is provided on an 'as-is' and 'as-available' basis and "
        "that Provider makes no warranty regarding uptime, performance, or suitability for any "
        "particular purpose.",
    )
    _clause(
        doc,
        "2. Subscription Term and Renewal",
        "The initial subscription term shall be twelve (12) months from the Effective Date "
        "(\"Initial Term\"). Unless Customer provides written notice of non-renewal at least "
        "ninety (90) days prior to the end of the Initial Term, this Agreement shall automatically "
        "renew for successive twelve (12) month periods (each, a \"Renewal Term\") on the same "
        "terms and conditions unless modified by Provider in accordance with Section 9. "
        "Termination of this Agreement, if permitted, shall be effective only at the end of the "
        "then-current Term; there is no mid-term termination right.",
    )
    _clause(
        doc,
        "3. Fees and Payment",
        "Customer shall pay the subscription fees set forth on the Order Form or invoice, due "
        "within thirty (30) days of invoice. In addition to subscription fees, Customer shall "
        "pay all applicable taxes, duties, and governmental charges. At each renewal, Provider "
        "may increase the subscription fees by any amount at Provider's sole discretion with no "
        "cap or limitation. Customer will be notified of the new fees, and by continuing to use "
        "the Service after the renewal date, Customer accepts the new fees. Failure to accept "
        "new fees does not suspend the Service or limit Customer's obligation to pay.",
    )
    _clause(
        doc,
        "4. Limitation of Liability",
        "EXCEPT FOR BREACHES OF CONFIDENTIALITY OBLIGATIONS, INFRINGEMENT OF INTELLECTUAL "
        "PROPERTY RIGHTS, OR INDEMNIFICATION OBLIGATIONS, PROVIDER'S TOTAL LIABILITY TO CUSTOMER "
        "FOR ANY CLAIM ARISING OUT OF OR RELATED TO THIS AGREEMENT SHALL NOT EXCEED THE FEES PAID "
        "BY CUSTOMER IN THE ONE (1) MONTH PRECEDING THE CLAIM. IN NO EVENT SHALL PROVIDER BE "
        "LIABLE FOR: (A) LOSS, CORRUPTION, OR UNAVAILABILITY OF DATA, REGARDLESS OF THE CAUSE OR "
        "THEORY OF LIABILITY; (B) LOSS OF PROFITS, REVENUE, BUSINESS OPPORTUNITY, OR ANTICIPATED "
        "SAVINGS; (C) INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES; OR "
        "(D) ANY DAMAGES THAT COULD HAVE BEEN AVOIDED WITH REASONABLE CARE. THESE LIMITATIONS "
        "APPLY REGARDLESS OF WHETHER PROVIDER HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.",
    )
    _clause(
        doc,
        "5. Data Protection and Security",
        "Provider shall maintain reasonable security measures to protect Customer data from "
        "unauthorized access. Provider may engage third-party service providers to process or "
        "store Customer data on its behalf. Customer waives any requirement for Provider to notify "
        "Customer of sub-processors or to obtain prior written consent before engaging "
        "sub-processors. Customer data may be processed or stored in any jurisdiction Provider "
        "determines appropriate. Provider is not responsible for any data loss, corruption, or "
        "unauthorized access resulting from Customer's use of the Service or failure to comply "
        "with Provider's instructions. In the event of a data breach or security incident, "
        "Provider shall provide notice to Customer in its discretion but is not obligated to "
        "provide any specific notice within any particular timeframe.",
    )
    _clause(
        doc,
        "6. Service Availability and Modifications",
        "Provider makes no guarantee regarding Service availability or uptime. Provider may modify, "
        "suspend, or discontinue any feature or function of the Service at any time without notice "
        "or liability. Provider may perform maintenance at any time, including during business "
        "hours, without advance notice. Provider may remove or restrict Customer data if Provider "
        "deems it necessary for security, legal, or operational reasons. Customer's sole remedy "
        "for unavailability or modification of the Service is termination of this Agreement, and "
        "termination shall not entitle Customer to any refund or compensation.",
    )
    _clause(
        doc,
        "7. Intellectual Property",
        "Provider retains all right, title, and interest in the Service, including all software, "
        "content, interfaces, and intellectual property. Customer is granted a limited, "
        "non-exclusive, non-transferable license to access and use the Service solely for "
        "Customer's internal business purposes during the Term and in accordance with this "
        "Agreement. Any feedback, suggestions, or ideas Customer provides to Provider may be used "
        "by Provider without restriction and without compensation to Customer.",
    )
    _clause(
        doc,
        "8. Fees and Renewal Terms",
        "At the renewal of this Agreement for any Renewal Term, Provider may increase subscription "
        "fees at Provider's sole discretion. Renewal fees are effective immediately upon the start "
        "of the Renewal Term. If Customer does not agree to increased fees, Customer's sole remedy "
        "is to terminate this Agreement effective at the end of the then-current Term, which "
        "requires at least ninety (90) days' notice. Failure to provide such notice shall result "
        "in automatic renewal at the increased fees.",
    )
    _clause(
        doc,
        "9. Amendment and Modification Rights",
        "Provider may modify the terms and conditions of this Agreement at any time by posting "
        "updated terms on the Service website or notifying Customer by email. Modified terms shall "
        "be effective immediately upon posting or upon the date specified by Provider in the "
        "notice, which may be no more than thirty (30) days after notice. By continuing to use the "
        "Service after such notice, Customer accepts the modified terms. Provider may also modify "
        "the features, functionality, performance characteristics, or pricing of the Service at any "
        "time without notice or obligation to Customer.",
    )
    _clause(
        doc,
        "10. Confidentiality",
        "Each party shall maintain the confidentiality of the other party's Confidential "
        "Information and shall not use such information except as necessary to perform this "
        "Agreement. Confidential Information shall not include information that is or becomes "
        "publicly available or is rightfully received from a third party.",
    )
    _clause(
        doc,
        "11. Term and Termination",
        "This Agreement shall commence on the Effective Date and shall continue for the Initial "
        "Term, automatically renewing for Renewal Terms as set forth in Section 2. Customer may "
        "terminate this Agreement only by providing ninety (90) days' prior written notice of "
        "non-renewal before the end of the then-current Term; there is no other termination right. "
        "Provider may terminate this Agreement for Customer's material breach on thirty (30) days' "
        "notice. Upon any termination, Customer's access to the Service shall be immediately "
        "revoked and Customer shall have no right to retrieve Customer data.",
    )
    _clause(
        doc,
        "12. Governing Law",
        "This Agreement shall be governed by the laws of the State of Delaware without regard to "
        "conflict-of-laws principles. Any legal action or dispute shall be subject to the exclusive "
        "jurisdiction of the federal and state courts located in Wilmington, Delaware.",
    )
    return doc


def build_reseller_distribution_agreement() -> Document:
    """A reseller/distribution agreement with moderate yellow-flag issues: one-sided exclusivity "
    that binds the reseller but not the supplier, annual minimum purchase commitments with no force
    majeure relief, asymmetric assignment rights favoring the supplier, and no force majeure clause
    anywhere."""
    doc = Document()
    _heading(doc, "RESELLER AND DISTRIBUTION AGREEMENT")
    doc.add_paragraph(
        "This Reseller and Distribution Agreement (\"Agreement\") is entered into between "
        "Nexus Technology Partners, Inc., a Delaware corporation (\"Supplier\"), and Elite "
        "Distribution Solutions, LLC (\"Reseller\"), effective as of the date of last execution "
        "(\"Effective Date\")."
    )
    _clause(
        doc,
        "1. Grant of Rights",
        "Supplier grants Reseller a non-exclusive right to purchase the products and services "
        "described in the Product Schedule (\"Products\") from Supplier and to resell such Products "
        "to end customers within the Territory specified in Appendix A (\"Territory\"). Reseller "
        "shall market, sell, distribute, and support the Products in accordance with the terms of "
        "this Agreement and any applicable policies established by Supplier from time to time.",
    )
    _clause(
        doc,
        "2. Exclusivity and Territorial Rights",
        "Reseller shall not sell or distribute any products or services that are competitive with "
        "or similar to the Products in the Territory without Supplier's prior written consent. "
        "Reseller agrees not to operate or maintain offices, distribution centers, or sales "
        "operations outside the Territory without Supplier's approval. This exclusivity obligation "
        "is binding on Reseller and shall continue in effect as long as Reseller remains in "
        "compliance with this Agreement. However, Supplier retains the right to sell or distribute "
        "Products directly to end customers in the Territory, appoint other resellers in the "
        "Territory or overlapping territories, and to establish distribution channels that compete "
        "with Reseller. Supplier is not restricted by exclusivity and may sell Products to any "
        "customer or reseller at Supplier's discretion.",
    )
    _clause(
        doc,
        "3. Minimum Purchase Commitments",
        "Reseller shall purchase and resell a minimum volume of Products as specified in the "
        "annual Reseller Performance Plan (\"Minimum Purchase Commitment\"). The Minimum Purchase "
        "Commitment shall be set by Supplier annually and shall be binding on Reseller. If Reseller "
        "fails to meet the Minimum Purchase Commitment in any calendar year, Supplier may, at its "
        "option: (a) terminate this Agreement without liability; (b) impose additional fees or "
        "penalties; or (c) reduce any rebates or promotional allowances due to Reseller. Reseller "
        "shall be obligated to meet the Minimum Purchase Commitment regardless of market demand, "
        "competitive pressures, economic conditions, or Supplier's failure to deliver Products, "
        "services, or support as promised. There shall be no relief from the Minimum Purchase "
        "Commitment for any reason, including force majeure events, supply chain disruptions, "
        "natural disasters, pandemics, or any other unforeseen circumstances affecting Reseller's "
        "ability to perform.",
    )
    _clause(
        doc,
        "4. Pricing and Distributor Terms",
        "Supplier shall provide Reseller with volume discount pricing and reseller margin as set "
        "forth in the Price Schedule attached hereto. All pricing is in U.S. dollars and excludes "
        "shipping, insurance, and applicable taxes. Supplier may modify pricing and distributor "
        "margins at any time upon thirty (30) days' notice to Reseller. Reseller shall be "
        "responsible for all freight, shipping, insurance, tariffs, and customs duties associated "
        "with the shipment of Products. Reseller shall not advertise, disclose, or use Supplier "
        "pricing for any purpose other than internal calculations.",
    )
    _clause(
        doc,
        "5. Assignment and Restrictions",
        "Supplier may freely assign this Agreement, its rights, and its obligations to any "
        "successor entity, acquirer, or affiliate without notice to Reseller, and such assignment "
        "shall be binding on Reseller. Reseller may not assign or transfer this Agreement, or any "
        "of Reseller's rights or obligations hereunder, without Supplier's prior written consent, "
        "which may be withheld in Supplier's sole discretion. Any attempted assignment by Reseller "
        "without consent shall be void. This restriction applies to any change of control of "
        "Reseller, including the sale of Reseller's business, merger, acquisition, or transfer of "
        "ownership, whether by operation of law or otherwise. In the event of any such change of "
        "control, Reseller shall immediately notify Supplier and shall provide Supplier the right "
        "to terminate this Agreement.",
    )
    _clause(
        doc,
        "6. Product Support and Training",
        "Supplier shall provide Reseller with product documentation, technical data sheets, and "
        "access to Supplier's online support portal. Reseller shall be responsible for providing "
        "first-level support to end customers. Supplier may provide advanced technical support at "
        "Supplier's discretion and at Reseller's cost. Supplier is not obligated to provide "
        "training, certification programs, or marketing support materials.",
    )
    _clause(
        doc,
        "7. Intellectual Property and Trademarks",
        "Supplier retains all right, title, and interest in all intellectual property, including "
        "trademarks, logos, patents, and copyrights related to the Products and Supplier's brand. "
        "Reseller is granted a limited license to use Supplier's trademarks solely in connection "
        "with the sale and resale of the Products. Reseller shall not modify, alter, or remove "
        "Supplier branding or intellectual property markings from any Products or marketing "
        "materials.",
    )
    _clause(
        doc,
        "8. Confidentiality",
        "Each party shall maintain the confidentiality of the other party's confidential "
        "information and shall not disclose such information without the other party's prior "
        "written consent. Confidential information shall not include information that is or becomes "
        "publicly available through no breach of this Agreement or is rightfully received from a "
        "third party without confidentiality restrictions.",
    )
    _clause(
        doc,
        "9. Limitation of Liability",
        "Neither party shall be liable for indirect, incidental, special, or consequential damages "
        "arising out of this Agreement. Each party's total liability shall not exceed the fees or "
        "compensation paid or due under this Agreement in the twelve (12) months preceding the "
        "claim.",
    )
    _clause(
        doc,
        "10. Term and Termination",
        "This Agreement shall commence on the Effective Date and shall continue for an initial "
        "term of three (3) years, automatically renewing for successive one-year periods unless "
        "either party provides at least ninety (90) days' written notice of non-renewal before the "
        "end of the then-current term. Either party may terminate this Agreement immediately for "
        "the other party's uncured material breach with thirty (30) days' written notice. Supplier "
        "may terminate this Agreement for convenience at any time upon thirty (30) days' notice "
        "without cause or liability.",
    )
    _clause(
        doc,
        "11. Governing Law",
        "This Agreement shall be governed by the laws of the State of New York without regard to "
        "conflict-of-laws principles. Any dispute shall be subject to the exclusive jurisdiction "
        "of the state and federal courts located in New York County.",
    )
    return doc


def main() -> None:
    _DATASET_DIR.mkdir(parents=True, exist_ok=True)
    targets = {
        "msa_professional_services.docx": build_msa_professional_services,
        "saas_subscription_agreement.docx": build_saas_subscription_agreement,
        "reseller_distribution_agreement.docx": build_reseller_distribution_agreement,
    }
    for filename, builder in targets.items():
        path = _DATASET_DIR / filename
        builder().save(str(path))
        print(f"wrote {path.relative_to(_REPO_ROOT)}")


if __name__ == "__main__":
    main()
