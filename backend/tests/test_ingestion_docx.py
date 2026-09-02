"""Real-document parsing tests for app.ingestion.docx.

Uploads are always ``.docx`` (plan §2.1: no PDF/OCR/LibreOffice), so this only exercises the
production ingestion path (``parse_document`` -> ``extract_docx``) against ACTUAL .docx files
built with python-docx (no mocks).
"""

from __future__ import annotations

from pathlib import Path

import pytest

from app.ingestion.docx import (
    _detect_format,
    _looks_like_heading,
    extract_docx,
    parse_document,
)

# --------------------------------------------------------------------------- #
# Format detection + heading heuristic (pure)
# --------------------------------------------------------------------------- #


def test_detect_format_from_extension_and_explicit_override():
    assert _detect_format(Path("nda.docx"), None) == "docx"
    # Explicit file_format wins over (and works without) an extension.
    assert _detect_format(Path("blob"), "docx") == "docx"
    assert _detect_format(Path("blob.bin"), ".docx") == "docx"


def test_detect_format_rejects_non_docx():
    with pytest.raises(ValueError, match="Cannot detect format"):
        _detect_format(Path("mystery.pdf"), None)
    with pytest.raises(ValueError, match="Unsupported file_format"):
        _detect_format(Path("x.docx"), "pdf")


@pytest.mark.parametrize(
    "line",
    [
        "1. Confidentiality",
        "1.2.3 Sub-obligation",
        "ARTICLE 4",
        "Section 2. Term",
        "Mutual Non-Disclosure Agreement",  # short Title-Case, no trailing punctuation
    ],
)
def test_looks_like_heading_accepts_headings(line):
    assert _looks_like_heading(line) is True


@pytest.mark.parametrize(
    "line",
    [
        "",
        "   ",
        "The receiving party shall keep the information secret.",  # sentence w/ period
        "this is lower-case running prose without any caps",
    ],
)
def test_looks_like_heading_rejects_body_text(line):
    assert _looks_like_heading(line) is False


# --------------------------------------------------------------------------- #
# parse_document guards
# --------------------------------------------------------------------------- #


def test_parse_document_missing_file_raises(tmp_path):
    with pytest.raises(ValueError, match="File not found"):
        parse_document(tmp_path / "nope.docx")


def test_parse_document_directory_is_not_a_file(tmp_path):
    with pytest.raises(ValueError, match="Not a file"):
        parse_document(tmp_path)


# --------------------------------------------------------------------------- #
# .docx extraction against a REAL generated document
# --------------------------------------------------------------------------- #


def _build_docx(path: Path) -> None:
    """Write a .docx exercising the tricky paths: a styled heading, a clause paragraph, a table with a
    NESTED table, and a horizontally-MERGED cell (the two correctness traps _cell_text/_distinct_cells
    guard)."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Confidentiality Obligations", style="Heading 1")
    doc.add_paragraph(
        "Section 2. Term. This Agreement remains in effect for three (3) years."
    )

    # Outer table whose second cell contains a nested table — _cell_text must recurse into it.
    outer = doc.add_table(rows=1, cols=2)
    outer.rows[0].cells[0].text = "Term"
    nested = outer.rows[0].cells[1].add_table(rows=1, cols=1)
    nested.rows[0].cells[0].text = "NESTEDMARK"

    # A horizontally-merged cell — row.cells repeats the same _tc per spanned column, so without the
    # _distinct_cells collapse this single cell's text would be emitted twice.
    merged_table = doc.add_table(rows=1, cols=2)
    merged = merged_table.cell(0, 0).merge(merged_table.cell(0, 1))
    merged.text = "MERGEDMARK"

    doc.save(str(path))


def test_extract_docx_pulls_paragraphs_tables_nested_and_merged(tmp_path):
    path = tmp_path / "nda.docx"
    _build_docx(path)

    parsed = extract_docx(path)

    assert parsed.file_format == "docx"
    assert "three (3) years" in parsed.full_text  # clause paragraph survived verbatim
    assert "NESTEDMARK" in parsed.full_text  # nested-table recursion (_cell_text)
    # Merged cell emitted exactly once, not duplicated across the spanned columns (_distinct_cells).
    assert parsed.full_text.count("MERGEDMARK") == 1

    heading = next(b for b in parsed.blocks if b.text == "Confidentiality Obligations")
    assert heading.is_heading is True
    assert heading.style == "Heading 1"


def test_extract_docx_empty_document_raises(tmp_path):
    from docx import Document

    path = tmp_path / "empty.docx"
    Document().save(str(path))  # no paragraphs, no tables
    with pytest.raises(ValueError, match="No extractable text"):
        extract_docx(path)


def test_extract_docx_preserves_paragraph_table_order(tmp_path):
    """A mid-document table must NOT be reordered to the end of full_text — its row text
    stays interleaved with the surrounding paragraphs so clause offsets line up."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("BEFORE TABLE PARA")
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "MIDTABLE"
    doc.add_paragraph("AFTER TABLE PARA")
    path = tmp_path / "ordered.docx"
    doc.save(str(path))

    parsed = extract_docx(path)

    before = parsed.full_text.index("BEFORE TABLE PARA")
    mid = parsed.full_text.index("MIDTABLE")
    after = parsed.full_text.index("AFTER TABLE PARA")
    assert before < mid < after


def test_parse_document_dispatches_docx(tmp_path):
    docx_path = tmp_path / "a.docx"
    _build_docx(docx_path)
    assert parse_document(docx_path).file_format == "docx"
