"""Pure-logic tests for the run-aware ``{{token}}`` substitution in support_task.generator.fill_docx.

fill_docx takes .docx bytes + a {name: value} dict and returns filled .docx bytes. It needs no DB,
network, or provider — only python-docx to build/read the in-memory documents. Covers:

* simple single-run placeholder replacement, and that token-name keys ('x') and full-placeholder
  keys ('{{x}}') both work;
* a placeholder SPLIT across multiple runs in one paragraph (runs built manually) — the subtle case
  the split-path handles by collapsing into run 0;
* a missing token: stripped (with its preceding ", " separator) when strip_unfilled is the default,
  but left verbatim when strip_unfilled=False;
* empty-parens cleanup after the only token inside parens is stripped;
* a token-only line being dropped when emptied;
* None values rendering as empty string, and unreadable bytes raising EngineError(422).
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from app.api.errors import EngineError
from app.support_task.generator import fill_docx


def _bytes_from(doc) -> bytes:
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _paragraph_texts(docx_bytes: bytes) -> list[str]:
    doc = Document(BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


def _single_paragraph_doc(text: str) -> bytes:
    """A doc whose one body paragraph holds ``text`` in a single run."""
    doc = Document()
    doc.add_paragraph(text)
    return _bytes_from(doc)


def _split_runs_doc(run_texts: list[str]) -> bytes:
    """A doc with one paragraph whose runs are exactly ``run_texts`` (built manually so a
    placeholder can straddle run boundaries)."""
    doc = Document()
    para = doc.add_paragraph()
    for chunk in run_texts:
        para.add_run(chunk)
    return _bytes_from(doc)


def test_simple_single_run_replacement():
    src = _single_paragraph_doc("Hello {{counterparty_name}}, welcome.")
    out = fill_docx(src, {"counterparty_name": "Acme Ltd"})
    assert _paragraph_texts(out) == ["Hello Acme Ltd, welcome."]


def test_full_placeholder_key_also_accepted():
    # _build_replacements accepts both bare token names and full {{...}} keys.
    src = _single_paragraph_doc("Hello {{name}}.")
    out = fill_docx(src, {"{{name}}": "Bob"})
    assert _paragraph_texts(out) == ["Hello Bob."]


def test_placeholder_split_across_runs():
    # Word frequently fragments typed text; here "{{name}}" is spread over five runs.
    src = _split_runs_doc(["Dear ", "{{", "na", "me", "}}", " --"])
    out = fill_docx(src, {"name": "Acme"})
    # The split-path collapses the joined paragraph into run 0 and blanks the rest.
    assert _paragraph_texts(out) == ["Dear Acme --"]


def test_missing_token_stripped_with_preceding_separator_by_default():
    # strip_unfilled defaults True: an unfilled token AND its leading ", " separator are removed.
    src = _single_paragraph_doc("{{city}}, {{country}} office")
    out = fill_docx(src, {"city": "Shenzhen"})
    # The comma before the unfilled {{country}} is consumed -> no dangling "Shenzhen, ".
    assert _paragraph_texts(out) == ["Shenzhen office"]


def test_missing_token_left_verbatim_when_strip_disabled():
    src = _single_paragraph_doc("{{city}}, {{country}} office")
    out = fill_docx(src, {"city": "Shenzhen"}, strip_unfilled=False)
    assert _paragraph_texts(out) == ["Shenzhen, {{country}} office"]


def test_empty_parens_cleanup_after_token_stripped():
    # The only token inside the parens is unfilled -> "()" plus its leading space is removed.
    src = _single_paragraph_doc("incorporated in ({{registration_number}})")
    out = fill_docx(src, {})
    assert _paragraph_texts(out) == ["incorporated in"]


def test_empty_parens_with_filled_value_kept():
    src = _single_paragraph_doc("incorporated in ({{registration_number}})")
    out = fill_docx(src, {"registration_number": "12345"})
    assert _paragraph_texts(out) == ["incorporated in (12345)"]


def test_token_only_line_dropped_when_emptied():
    # A standalone token line that strips to empty is removed entirely (not left as a blank line),
    # guarded so it never removes the last surviving paragraph.
    doc = Document()
    doc.add_paragraph("Header stays")
    doc.add_paragraph("{{country}}")
    doc.add_paragraph("Footer stays")
    src = _bytes_from(doc)
    out = fill_docx(src, {})
    assert _paragraph_texts(out) == ["Header stays", "Footer stays"]


def test_none_value_renders_as_empty_string():
    # _build_replacements maps a None value to "" (a real, supplied, empty value — so it is
    # substituted rather than stripped as "unfilled"). Separator stripping does not apply.
    src = _single_paragraph_doc("X={{token}}Y")
    out = fill_docx(src, {"token": None})
    assert _paragraph_texts(out) == ["X=Y"]


def test_blank_and_none_keys_ignored():
    # Keys that are None or blank are skipped; the unmatched placeholder is then stripped.
    src = _single_paragraph_doc("A {{kept}} B")
    out = fill_docx(src, {None: "x", "  ": "y", "kept": "Z"})
    assert _paragraph_texts(out) == ["A Z B"]


def test_replacement_inside_table_cell():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].add_run("Party: {{counterparty_name}}")
    src = _bytes_from(doc)
    out = fill_docx(src, {"counterparty_name": "Acme"})
    rt = Document(BytesIO(out))
    assert rt.tables[0].rows[0].cells[0].paragraphs[0].text == "Party: Acme"


def test_unreadable_bytes_raise_engine_error_422():
    with pytest.raises(EngineError) as exc_info:
        fill_docx(b"this is not a zip/docx", {"x": "y"})
    assert exc_info.value.status == 422
    assert exc_info.value.code == "bad_template"
