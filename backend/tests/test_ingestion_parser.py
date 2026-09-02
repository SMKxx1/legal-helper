"""Real-document parsing tests for app.ingestion.parser.

The HTTP review tests only ever submit ``.txt``, so the docx/pdf extraction paths that REAL uploads
take (``parse_document`` -> ``extract_docx`` / ``extract_pdf``) were almost entirely uncovered — a
regression in table/nested-table/merged-cell handling or PDF text extraction would pass CI green. These
tests build ACTUAL .docx and .pdf files on disk with python-docx / fpdf2 and parse them back, so the
production ingestion path is exercised end to end (no mocks).
"""

from __future__ import annotations

from pathlib import Path

import pytest

from app.ingestion.parser import (
    NoTextLayerError,
    _detect_format,
    _looks_like_heading,
    extract_docx,
    extract_pdf,
    extract_text_file,
    parse_document,
)

# --------------------------------------------------------------------------- #
# Format detection + heading heuristic (pure)
# --------------------------------------------------------------------------- #


def test_detect_format_from_extension_and_explicit_override():
    assert _detect_format(Path("nda.docx"), None) == "docx"
    assert _detect_format(Path("scan.PDF"), None) == "pdf"  # case-insensitive
    assert _detect_format(Path("notes.markdown"), None) == "md"
    # Explicit file_format wins over (and works without) an extension.
    assert _detect_format(Path("blob"), "txt") == "txt"
    assert _detect_format(Path("blob.bin"), ".docx") == "docx"


def test_detect_format_rejects_unknown_extension_and_bad_override():
    with pytest.raises(ValueError, match="Cannot detect format"):
        _detect_format(Path("mystery.xyz"), None)
    with pytest.raises(ValueError, match="Unsupported file_format"):
        _detect_format(Path("x.docx"), "exe")


@pytest.mark.parametrize(
    "line",
    [
        "1. Confidentiality",
        "1.2.3 Sub-obligation",
        "ARTICLE 4",
        "Section 2. Term",
        "Mutual Non-Disclosure Agreement",  # short Title-Case, no trailing punctuation
    ],
)
def test_looks_like_heading_accepts_headings(line):
    assert _looks_like_heading(line) is True


@pytest.mark.parametrize(
    "line",
    [
        "",
        "   ",
        "The receiving party shall keep the information secret.",  # sentence w/ period
        "this is lower-case running prose without any caps",
    ],
)
def test_looks_like_heading_rejects_body_text(line):
    assert _looks_like_heading(line) is False


# --------------------------------------------------------------------------- #
# parse_document guards
# --------------------------------------------------------------------------- #


def test_parse_document_missing_file_raises(tmp_path):
    with pytest.raises(ValueError, match="File not found"):
        parse_document(tmp_path / "nope.txt")


def test_parse_document_directory_is_not_a_file(tmp_path):
    with pytest.raises(ValueError, match="Not a file"):
        parse_document(tmp_path)


# --------------------------------------------------------------------------- #
# .docx extraction against a REAL generated document
# --------------------------------------------------------------------------- #


def _build_docx(path: Path) -> None:
    """Write a .docx exercising the tricky paths: a styled heading, a clause paragraph, a table with a
    NESTED table, and a horizontally-MERGED cell (the two correctness traps _cell_text/_distinct_cells
    guard)."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Confidentiality Obligations", style="Heading 1")
    doc.add_paragraph(
        "Section 2. Term. This Agreement remains in effect for three (3) years."
    )

    # Outer table whose second cell contains a nested table — _cell_text must recurse into it.
    outer = doc.add_table(rows=1, cols=2)
    outer.rows[0].cells[0].text = "Term"
    nested = outer.rows[0].cells[1].add_table(rows=1, cols=1)
    nested.rows[0].cells[0].text = "NESTEDMARK"

    # A horizontally-merged cell — row.cells repeats the same _tc per spanned column, so without the
    # _distinct_cells collapse this single cell's text would be emitted twice.
    merged_table = doc.add_table(rows=1, cols=2)
    merged = merged_table.cell(0, 0).merge(merged_table.cell(0, 1))
    merged.text = "MERGEDMARK"

    doc.save(str(path))


def test_extract_docx_pulls_paragraphs_tables_nested_and_merged(tmp_path):
    path = tmp_path / "nda.docx"
    _build_docx(path)

    parsed = extract_docx(path)

    assert parsed.file_format == "docx"
    assert "three (3) years" in parsed.full_text  # clause paragraph survived verbatim
    assert "NESTEDMARK" in parsed.full_text  # nested-table recursion (_cell_text)
    # Merged cell emitted exactly once, not duplicated across the spanned columns (_distinct_cells).
    assert parsed.full_text.count("MERGEDMARK") == 1

    heading = next(b for b in parsed.blocks if b.text == "Confidentiality Obligations")
    assert heading.is_heading is True
    assert heading.style == "Heading 1"


def test_extract_docx_empty_document_raises(tmp_path):
    from docx import Document

    path = tmp_path / "empty.docx"
    Document().save(str(path))  # no paragraphs, no tables
    with pytest.raises(ValueError, match="No extractable text"):
        extract_docx(path)


def test_extract_docx_preserves_paragraph_table_order(tmp_path):
    """A mid-document table must NOT be reordered to the end of full_text — its row text
    stays interleaved with the surrounding paragraphs so clause offsets line up."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("BEFORE TABLE PARA")
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "MIDTABLE"
    doc.add_paragraph("AFTER TABLE PARA")
    path = tmp_path / "ordered.docx"
    doc.save(str(path))

    parsed = extract_docx(path)

    before = parsed.full_text.index("BEFORE TABLE PARA")
    mid = parsed.full_text.index("MIDTABLE")
    after = parsed.full_text.index("AFTER TABLE PARA")
    assert before < mid < after


# --------------------------------------------------------------------------- #
# .pdf extraction against a REAL generated PDF (text layer present)
# --------------------------------------------------------------------------- #


def _build_text_pdf(path: Path, body: str) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, body)
    pdf.output(str(path))


def test_extract_pdf_reads_text_layer_with_page_numbers(tmp_path):
    # >100 chars on the page so _looks_scanned stays False and we exercise the text path (not OCR).
    body = (
        "Section 1. Confidentiality. The Receiving Party shall keep the Confidential "
        "Information secret and shall not disclose it to any third party without prior "
        "written consent of the Disclosing Party."
    )
    path = tmp_path / "nda.pdf"
    _build_text_pdf(path, body)

    parsed = extract_pdf(path)

    assert parsed.file_format == "pdf"
    # Words survive extraction (line-wrapping may split phrases, so assert distinctive tokens).
    for token in ("Confidentiality", "Receiving", "disclose", "Disclosing"):
        assert token in parsed.full_text
    assert parsed.blocks, "expected line blocks"
    assert all(b.page == 1 for b in parsed.blocks)


def test_extract_pdf_with_no_text_layer_raises_no_text_layer_error(
    tmp_path, monkeypatch
):
    """A PDF with a truly empty text layer surfaces NoTextLayerError (the 'needs OCR' signal) — with
    OCR forced unavailable so the scanned-fallback branch is skipped deterministically."""
    import fitz

    from app.ingestion import ocr

    # Force the parser's scanned-detection + fallback to treat OCR as unavailable.
    monkeypatch.setattr(ocr, "ocr_available", lambda: False)

    path = tmp_path / "blank.pdf"
    doc = fitz.open()
    doc.new_page()  # a page with no text at all
    doc.save(str(path))
    doc.close()

    with pytest.raises(NoTextLayerError):
        extract_pdf(path)


# --------------------------------------------------------------------------- #
# .txt / .md extraction
# --------------------------------------------------------------------------- #


def test_extract_text_file_strips_bom_and_normalizes_markdown_heading(tmp_path):
    path = tmp_path / "doc.md"
    # Leading BOM + an ATX markdown heading + a body line.
    path.write_text("﻿## Confidentiality\nThe parties agree.", encoding="utf-8")

    parsed = extract_text_file(path, file_format="md")

    assert parsed.file_format == "md"
    head = parsed.blocks[0]
    assert head.text == "Confidentiality"  # '##' stripped, BOM gone
    assert head.is_heading is True
    assert not parsed.full_text.startswith("﻿")  # BOM stripped at decode


def test_extract_text_file_empty_raises(tmp_path):
    path = tmp_path / "empty.txt"
    path.write_text("   \n\n", encoding="utf-8")
    with pytest.raises(ValueError, match="No extractable text"):
        extract_text_file(path)


# --------------------------------------------------------------------------- #
# parse_document dispatch routes each format to the right extractor
# --------------------------------------------------------------------------- #


def test_parse_document_dispatches_by_format(tmp_path):
    docx_path = tmp_path / "a.docx"
    _build_docx(docx_path)
    assert parse_document(docx_path).file_format == "docx"

    pdf_path = tmp_path / "a.pdf"
    _build_text_pdf(
        pdf_path,
        "Section 1. The Receiving Party shall protect the Confidential Information "
        "and shall not disclose it to any third party at any time whatsoever.",
    )
    assert parse_document(pdf_path).file_format == "pdf"

    txt_path = tmp_path / "a.txt"
    txt_path.write_text("Plain text NDA body.", encoding="utf-8")
    assert parse_document(txt_path).file_format == "txt"
