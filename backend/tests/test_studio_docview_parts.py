"""The PART-LENGTH offset model (docview parts + the studio parts renderer).

The doc view renders formatting and friendly token labels, which makes the DOM's textContent
diverge from ``docview.paragraph_text`` — the offset basis ``tokenize_ops`` operates on. These
tests pin the model that keeps offsets sound anyway:

* parts are a pure RE-SLICING of ``paragraph_text``: the concatenation of every part's ``text``
  is exactly the segment text (golden — the offset basis itself is byte-identical to before);
* run parts carry the run's bold/italic/underline flags; a ``{{token}}`` (even one straddling
  runs) is ONE atomic token part whose ``text`` is the verbatim match;
* ``plen`` (== ``len(part.text)``) is each part's underlying length — for a token part the raw
  ``{{name}}`` length, never the rendered label length;
* the client's data-plen summation (mirrored here) computes the SAME ``(start, end)`` the old
  textContent method produced for the underlying paragraph_text;
* the parts renderer resolves registry labels (Title-Cased fallback), renders token chips showing
  the LABEL (not ``{{raw}}``), emits tables as real grids, and never emits stray text inside a seg.
"""

from __future__ import annotations

from typing import Any

from conftest_studio import doc_to_bytes, rich_doc, runs_doc, single_para_doc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.api.routes_studio import (
    _doc_blocks,
    _label_map,
    _part_ctx,
    _render_partial,
    _seg_ctx,
)
from app.studio.docview import ViewPart, extract_view, load_document, paragraph_text


# --------------------------------------------------------------------------- #
# Parts extraction
# --------------------------------------------------------------------------- #
def test_parts_reslice_paragraph_text_exactly_golden():
    """Golden: parts concatenate back to seg.text, and seg.text is STILL the filler's
    run-concatenated normalization — the offset basis did not move under the new renderer."""
    data = rich_doc()
    view = extract_view(data)
    doc = load_document(data)
    from app.studio.docview import resolve_locator

    for seg in view.segments:
        assert "".join(p.text for p in seg.parts) == seg.text, seg.locator
        assert sum(p.plen for p in seg.parts) == len(seg.text), seg.locator
        assert paragraph_text(resolve_locator(doc, seg.locator)) == seg.text, (
            seg.locator
        )


def test_run_parts_carry_formatting_flags():
    view = extract_view(
        runs_doc(
            ("plain ", {}),
            ("bold", {"bold": True}),
            ("italic", {"italic": True}),
            ("under", {"underline": True}),
        )
    )
    parts = view.segments[0].parts
    assert [(p.text, p.bold, p.italic, p.underline) for p in parts] == [
        ("plain ", False, False, False),
        ("bold", True, False, False),
        ("italic", False, True, False),
        ("under", False, False, True),
    ]
    assert not any(p.is_token for p in parts)


def test_adjacent_same_format_runs_merge_into_one_part():
    view = extract_view(runs_doc(("Hel", {}), ("lo", {})))
    parts = view.segments[0].parts
    assert len(parts) == 1
    assert parts[0].text == "Hello" and parts[0].plen == 5


def test_token_parts_are_verbatim_with_raw_plen():
    view = extract_view(
        single_para_doc("Hi {{counterparty_name}} and {{ spaced_tok }}!")
    )
    parts = view.segments[0].parts
    toks = [p for p in parts if p.is_token]
    assert [(p.name, p.text) for p in toks] == [
        ("counterparty_name", "{{counterparty_name}}"),
        ("spaced_tok", "{{ spaced_tok }}"),  # verbatim: inner spaces count toward plen
    ]
    assert toks[0].plen == len("{{counterparty_name}}")
    assert toks[1].plen == len("{{ spaced_tok }}")


def test_token_straddling_runs_is_one_atomic_part():
    view = extract_view(
        runs_doc(
            ("Dear ", {}), ("{{", {}), ("na", {"bold": True}), ("me", {}), ("}}", {})
        )
    )
    parts = view.segments[0].parts
    assert [(p.is_token, p.text) for p in parts] == [
        (False, "Dear "),
        (True, "{{name}}"),
    ]
    assert parts[1].name == "name"


def test_tables_headers_footers_have_parts():
    view = extract_view(rich_doc())
    cell = view.find("body/tbl:0:0:0/p:0")
    assert cell is not None
    assert [p.name for p in cell.parts if p.is_token] == ["existing_token"]
    hdr = view.find("hdr:0:default/p:0")
    ftr = view.find("ftr:0:default/p:0")
    assert hdr is not None and "".join(p.text for p in hdr.parts) == "Header ACME text"
    assert ftr is not None and ftr.parts and not ftr.parts[0].is_token


def test_heading_level_and_alignment_are_exposed():
    doc = Document()
    doc.add_heading("Confidentiality", level=2)
    p = doc.add_paragraph("centered text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    view = extract_view(doc_to_bytes(doc))
    heading = next(s for s in view.segments if s.text == "Confidentiality")
    assert heading.heading == 2 and heading.style == "Heading 2"
    centered = next(s for s in view.segments if s.text == "centered text")
    assert centered.align == "center" and centered.heading == 0


# --------------------------------------------------------------------------- #
# The client offset mirror: data-plen summation == the old textContent method
# --------------------------------------------------------------------------- #
def _client_offset(
    parts: tuple[ViewPart, ...], part_idx: int, intra: int, is_end: bool
) -> int:
    """Mirror of studio.js offsetInSeg: sum data-plen (== len(part.text)) of the parts before the
    boundary; inside a run add the intra-run text offset (rendered == underlying, 1:1); inside an
    ATOMIC token chip snap to its start (selection start) or end (selection end)."""
    total = sum(p.plen for p in parts[:part_idx])
    part = parts[part_idx]
    if part.is_token:
        return total + (part.plen if is_end else 0)
    return total + intra


def test_part_offsets_match_the_old_textcontent_method():
    """A selection spanning a bold run + a token computes the same (start, end) the old
    Range.toString().length method produced against the underlying paragraph_text."""
    view = extract_view(
        runs_doc(
            ("Between ", {}),
            ("ACME", {"bold": True}),
            (" and {{counterparty_name}} forever.", {}),
        )
    )
    seg = view.segments[0]
    text = seg.text
    assert text == "Between ACME and {{counterparty_name}} forever."
    parts = seg.parts
    assert [p.text for p in parts] == [
        "Between ",
        "ACME",
        " and ",
        "{{counterparty_name}}",
        " forever.",
    ]

    # Start boundary: 3 chars into the first run — the old method returned index 3.
    start = _client_offset(parts, 0, 3, is_end=False)
    assert start == 3 == text.index("ween")

    # A boundary inside the BOLD run still maps 1:1 across the part seam.
    assert _client_offset(parts, 1, 2, is_end=False) == text.index("ME")

    # End boundary inside the token chip snaps to the chip's end — exactly where the old method
    # landed for a selection ending after the raw "}}" in the underlying text.
    end = _client_offset(parts, 3, 1, is_end=True)
    assert end == text.index("}}") + 2 == len("Between ACME and {{counterparty_name}}")

    # And a start boundary inside the chip snaps to its start (never mid-token).
    assert _client_offset(parts, 3, 1, is_end=False) == text.index("{{")


# --------------------------------------------------------------------------- #
# The studio parts renderer (labels, chips, tables, whitespace discipline)
# --------------------------------------------------------------------------- #
def test_part_ctx_resolves_registry_label_with_titlecase_fallback():
    tok = ViewPart(
        text="{{counterparty_name}}", is_token=True, name="counterparty_name"
    )
    assert (
        _part_ctx(tok, {"counterparty_name": "The Other Side"})["label"]
        == "The Other Side"
    )
    fallback = _part_ctx(tok, {})
    assert fallback["label"] == "Counterparty Name"  # Title-Cased snake_case fallback
    assert fallback["plen"] == len("{{counterparty_name}}")
    assert (
        _part_ctx(ViewPart(text="{{city_zip}}", is_token=True, name="city_zip"), {})[
            "label"
        ]
        == "City Zip"
    )


def test_docview_partial_renders_label_chip_not_raw_placeholder():
    view = extract_view(single_para_doc("Hi {{counterparty_name}}."))
    html = _render_partial(
        "studio/_docview.html",
        {"blocks": _doc_blocks(view, {"counterparty_name": "CP X"})},
    )
    # The seg is rendered as EXACTLY these part nodes — no stray text inside the seg, every part
    # carrying data-plen (the underlying length: 21 for the chip, not len("CP X")).
    assert (
        '<div class="stu-seg" data-locator="body/p:0" data-kind="body">'
        '<span class="stu-run" data-plen="3">Hi </span>'
        '<span class="stu-tok" data-token="counterparty_name" data-plen="21"'
        ' contenteditable="false" title="{{counterparty_name}}">CP X</span>'
        '<span class="stu-run" data-plen="1">.</span></div>'
    ) in html
    assert (
        ">{{counterparty_name}}</span>" not in html
    )  # the raw braces are never visible text


def test_docview_partial_renders_formatting_and_heading_classes():
    doc = Document()
    doc.add_heading("Terms", level=1)
    p = doc.add_paragraph()
    p.add_run("bold").bold = True
    p.add_run("ital").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    view = extract_view(doc_to_bytes(doc))
    html = _render_partial("studio/_docview.html", {"blocks": _doc_blocks(view, {})})
    assert 'class="stu-seg stu-h1"' in html
    assert "stu-align-center" in html
    assert '<span class="stu-run stu-b" data-plen="4">bold</span>' in html
    assert '<span class="stu-run stu-i" data-plen="4">ital</span>' in html


def _tree_locators(blocks: list[dict[str, Any]]) -> list[str]:
    out: list[str] = []
    for b in blocks:
        if b["type"] == "p":
            out.append(b["seg"]["locator"])
        else:
            for row in b["rows"]:
                for cell in row:
                    out.extend(_tree_locators(cell["blocks"]))
    return out


def test_doc_blocks_builds_table_tree_preserving_every_locator_in_order():
    view = extract_view(rich_doc())
    blocks = _doc_blocks(view, {})
    # Pure regrouping: the tree walks back to the flat view, same segments, same order.
    assert _tree_locators(blocks) == [s.locator for s in view.segments]
    table = next(b for b in blocks if b["type"] == "table")
    assert len(table["rows"]) == 2 and len(table["rows"][0]) == 2
    # The (1,0) cell holds the nested table as a real nested grid.
    cell_10 = table["rows"][1][0]
    nested = [b for b in cell_10["blocks"] if b["type"] == "table"]
    assert nested and nested[0]["rows"][0][0]["blocks"][0]["seg"]["locator"].startswith(
        "body/tbl:0:1:0/tbl:0:0:0/"
    )
    html = _render_partial("studio/_docview.html", {"blocks": blocks})
    assert html.count('<table class="stu-table">') == 2  # outer + nested


def test_seg_ctx_kind_and_plen_roundtrip():
    view = extract_view(rich_doc())
    hdr = view.find("hdr:0:default/p:0")
    ctx = _seg_ctx(hdr, {})
    assert ctx["kind"] == "header"
    assert sum(p["plen"] for p in ctx["parts"]) == len(hdr.text)


def test_label_map_prefers_registry_labels():
    class _T:  # the TokenView shape _label_map consumes
        def __init__(self, name, label):
            self.name, self.label = name, label

    labels = _label_map([_T("street_address", "Street"), _T("bare", "")])
    assert labels == {"street_address": "Street", "bare": ""}
    # An empty registry label falls through to the Title-Case fallback at render time.
    assert (
        _part_ctx(ViewPart(text="{{bare}}", is_token=True, name="bare"), labels)[
            "label"
        ]
        == "Bare"
    )
