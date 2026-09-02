"""Find-and-map assistant (app.studio.findmap): detection matrix, fuzzy suggestions, batching.

* candidate matrix — square brackets, angle brackets, underscore blanks, ALL-CAPS parentheticals;
  letterless matches ("[42]") excluded; spans overlapping an existing ``{{token}}`` excluded;
  overlapping candidates resolved greedily (earliest, then longest);
* candidates carry locators + offsets directly usable by ``apply_tokenize``, across body, table,
  nested-table, header and footer segments;
* fuzzy suggestions come only from the caller-provided (name, label) list: exact-ish name match,
  label match, and a below-cutoff miss that still reports its score with ``suggested_token=None``;
* ``ordered_mappings`` applies same-paragraph mappings right-to-left so all offsets stay valid
  against the original view, and refuses overlapping mappings before anything is applied;
* ``map_all`` yields one op record per mapping, hash-chained in applied order, refuses a stale
  baseline, and aborts the whole batch on any refusal.
"""

from __future__ import annotations

import pytest
from conftest_studio import doc_to_bytes, single_para_doc
from docx import Document

from app.studio.docview import content_hash, extract_view
from app.studio.errors import (
    OverlappingMappingsError,
    StaleViewError,
    TokenOverlapError,
)
from app.studio.findmap import (
    TokenMapping,
    detect_placeholders,
    map_all,
    ordered_mappings,
)

TOKENS = [
    ("company_name", "Company name"),
    ("effective_date", "Effective date"),
    ("governing_law_state", "Governing law state"),
    ("counterparty_name", "Counterparty name"),
]


def _detect(text: str, tokens=TOKENS):
    view = extract_view(single_para_doc(text))
    return detect_placeholders(view, tokens)


# --------------------------------------------------------------------------- #
# Candidate matrix
# --------------------------------------------------------------------------- #
def test_square_bracket_candidates_upper_and_title_case():
    for text in ("Between [COMPANY NAME] and us", "Between [Company Name] and us"):
        (candidate,) = _detect(text)
        assert candidate.matched_text == text[8 : 8 + len("[COMPANY NAME]")]
        assert (candidate.start, candidate.end) == (8, 8 + len("[COMPANY NAME]"))
        assert candidate.suggested_token == "company_name"


def test_angle_bracket_candidate():
    (candidate,) = _detect("Deliver to <Company> promptly")
    assert candidate.matched_text == "<Company>"
    assert candidate.suggested_token == "company_name"


def test_underscore_blank_candidate():
    (candidate,) = _detect("Dated: ______ (sign here)")
    assert candidate.matched_text == "______"
    assert candidate.suggested_token is None  # blanks carry no text to match on
    assert candidate.score == 0.0


def test_caps_parenthetical_candidate_needs_three_words():
    (candidate,) = _detect("law of (GOVERNING LAW STATE) applies")
    assert candidate.matched_text == "(GOVERNING LAW STATE)"
    assert candidate.suggested_token == "governing_law_state"
    assert _detect("see (CLAUSE FOUR) above") == []  # two words: not a placeholder
    assert _detect("as agreed (see clause 4)") == []  # lower-case: prose


def test_letterless_matches_are_not_candidates():
    assert _detect("see clause [42] and <-> markers") == []


def test_spans_overlapping_existing_tokens_are_excluded():
    candidates = _detect("Signed {{company_name}} on [DATE]")
    assert [c.matched_text for c in candidates] == ["[DATE]"]
    # an underscore run INSIDE an existing token is not re-offered as a blank
    assert _detect("fill {{a___b}} blank") == []


def test_overlapping_candidates_resolve_earliest_then_longest():
    # the square match starts first and wins; the underscore run inside it is dropped
    candidates = _detect("x [___ COMPANY NAME] y")
    assert [c.matched_text for c in candidates] == ["[___ COMPANY NAME]"]


def test_offsets_are_directly_usable_by_apply_tokenize():
    from app.studio.tokenize_ops import apply_tokenize

    data = single_para_doc("Between [COMPANY NAME] and us")
    view = extract_view(data)
    (candidate,) = detect_placeholders(view, TOKENS)
    new, record = apply_tokenize(
        data,
        candidate.locator,
        candidate.start,
        candidate.end,
        candidate.suggested_token,
        expected_hash=view.content_hash,
    )
    assert extract_view(new).segments[0].text == "Between {{company_name}} and us"
    assert record.replaced_text == "[COMPANY NAME]"


def test_detection_covers_tables_nested_tables_headers_and_footers():
    doc = Document()
    doc.add_paragraph("Body [COMPANY NAME] here")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Cell [EFFECTIVE DATE]")
    nested = table.rows[0].cells[1].add_table(rows=1, cols=1)
    nested.rows[0].cells[0].paragraphs[0].add_run("Nested <Company>")
    doc.sections[0].header.paragraphs[0].add_run("Header ______")
    doc.sections[0].footer.paragraphs[0].add_run("Footer (GOVERNING LAW STATE)")
    view = extract_view(doc_to_bytes(doc))
    candidates = detect_placeholders(view, TOKENS)
    assert len(candidates) == 5
    locators = {c.matched_text: c.locator for c in candidates}
    assert locators["[COMPANY NAME]"] == "body/p:0"
    assert locators["[EFFECTIVE DATE]"] == "body/tbl:0:0:0/p:0"
    assert locators["<Company>"].startswith("body/tbl:0:0:1/tbl:0:0:0/")
    assert locators["______"].startswith("hdr:0:default/")
    assert locators["(GOVERNING LAW STATE)"].startswith("ftr:0:default/")


def test_multiple_candidates_in_one_paragraph_report_distinct_offsets():
    candidates = _detect("[COMPANY NAME] agrees with [COUNTERPARTY NAME] today")
    assert [c.matched_text for c in candidates] == [
        "[COMPANY NAME]",
        "[COUNTERPARTY NAME]",
    ]
    assert candidates[0].end <= candidates[1].start
    assert candidates[1].suggested_token == "counterparty_name"


# --------------------------------------------------------------------------- #
# Fuzzy suggestion matrix
# --------------------------------------------------------------------------- #
def test_exact_name_match_scores_full():
    (candidate,) = _detect("sign [COMPANY NAME] now")
    assert candidate.suggested_token == "company_name"
    assert candidate.score == 1.0


def test_label_match_suggests_the_token_name():
    (candidate,) = _detect(
        "date: [Effective date] here", tokens=[("effective_date", "Effective date")]
    )
    assert candidate.suggested_token == "effective_date"
    assert candidate.score == 1.0


def test_typo_still_matches_fuzzily():
    (candidate,) = _detect("sign [COMPNY NAME] now")
    assert candidate.suggested_token == "company_name"
    assert 0.55 <= candidate.score < 1.0


def test_below_cutoff_yields_no_suggestion_but_reports_score():
    (candidate,) = _detect("attach [EXHIBIT ZQ] here")
    assert candidate.suggested_token is None
    assert 0.0 <= candidate.score < 0.55


def test_empty_token_options_never_suggest():
    (candidate,) = _detect("sign [COMPANY NAME] now", tokens=[])
    assert candidate.suggested_token is None
    assert candidate.score == 0.0


# --------------------------------------------------------------------------- #
# Batch ordering + map_all
# --------------------------------------------------------------------------- #
def test_ordered_mappings_apply_right_to_left_within_a_paragraph():
    mappings = [
        {"locator": "body/p:0", "start": 5, "end": 10, "token_name": "a"},
        {"locator": "body/p:0", "start": 20, "end": 25, "token_name": "b"},
        TokenMapping(locator="body/p:0", start=12, end=15, token_name="c"),
    ]
    ordered = ordered_mappings(mappings)
    assert [m.start for m in ordered] == [20, 12, 5]
    assert all(isinstance(m, TokenMapping) for m in ordered)


def test_ordered_mappings_refuse_overlaps_before_applying():
    with pytest.raises(OverlappingMappingsError) as exc:
        ordered_mappings(
            [
                {"locator": "body/p:0", "start": 5, "end": 12, "token_name": "a"},
                {"locator": "body/p:0", "start": 10, "end": 15, "token_name": "b"},
            ]
        )
    assert exc.value.code == "studio_overlapping_mappings"
    # identical spans are overlaps too
    with pytest.raises(OverlappingMappingsError):
        ordered_mappings(
            [
                {"locator": "body/p:0", "start": 5, "end": 12, "token_name": "a"},
                {"locator": "body/p:0", "start": 5, "end": 12, "token_name": "b"},
            ]
        )
    # same offsets in DIFFERENT paragraphs are fine
    assert (
        len(
            ordered_mappings(
                [
                    {"locator": "body/p:0", "start": 5, "end": 12, "token_name": "a"},
                    {"locator": "body/p:1", "start": 5, "end": 12, "token_name": "b"},
                ]
            )
        )
        == 2
    )


def test_map_all_chains_records_and_survives_left_to_right_input():
    data = single_para_doc("Between [COMPANY] and [RECIPIENT] on [DATE].")
    view = extract_view(data)
    text = view.segments[0].text

    def mapping(sub: str, token: str) -> dict:
        start = text.index(sub)
        return {
            "locator": "body/p:0",
            "start": start,
            "end": start + len(sub),
            "token_name": token,
        }

    new, records = map_all(
        data,
        [
            mapping("[COMPANY]", "company"),
            mapping("[RECIPIENT]", "recipient"),
            mapping("[DATE]", "date"),
        ],
        expected_hash=view.content_hash,
    )
    assert (
        extract_view(new).segments[0].text
        == "Between {{company}} and {{recipient}} on {{date}}."
    )
    assert len(records) == 3
    assert records[0].prior_hash == view.content_hash
    for earlier, later in zip(records, records[1:], strict=False):
        assert earlier.new_hash == later.prior_hash
    assert records[-1].new_hash == content_hash(new)
    # applied right-to-left: the record order is descending start
    assert [r.replaced_text for r in records] == ["[DATE]", "[RECIPIENT]", "[COMPANY]"]


def test_map_all_refuses_stale_baseline():
    data = single_para_doc("map [THIS] please")
    with pytest.raises(StaleViewError):
        map_all(
            data,
            [{"locator": "body/p:0", "start": 4, "end": 10, "token_name": "t"}],
            expected_hash="0" * 64,
        )


def test_map_all_aborts_whole_batch_on_any_refusal():
    data = single_para_doc("Keep {{locked}} and [FREE] here.")
    text = extract_view(data).segments[0].text
    free = text.index("[FREE]")
    locked = text.index("{{locked}}")
    with pytest.raises(TokenOverlapError):
        map_all(
            data,
            [
                {
                    "locator": "body/p:0",
                    "start": free,
                    "end": free + 6,
                    "token_name": "ok",
                },
                {
                    "locator": "body/p:0",
                    "start": locked,
                    "end": locked + 4,
                    "token_name": "bad",
                },
            ],
        )


def test_map_all_with_empty_batch_returns_input_unchanged():
    data = single_para_doc("nothing here")
    out, records = map_all(data, [])
    assert out == data
    assert records == []


def test_map_all_spans_multiple_document_parts():
    doc = Document()
    doc.add_paragraph("Body [COMPANY] text")
    doc.sections[0].header.paragraphs[0].add_run("Header [DATE] text")
    data = doc_to_bytes(doc)
    view = extract_view(data)
    body_text = view.find("body/p:0").text
    hdr_text = view.find("hdr:0:default/p:0").text
    new, records = map_all(
        data,
        [
            {
                "locator": "body/p:0",
                "start": body_text.index("[COMPANY]"),
                "end": body_text.index("[COMPANY]") + 9,
                "token_name": "company",
            },
            {
                "locator": "hdr:0:default/p:0",
                "start": hdr_text.index("[DATE]"),
                "end": hdr_text.index("[DATE]") + 6,
                "token_name": "date",
            },
        ],
        expected_hash=view.content_hash,
    )
    result = extract_view(new)
    assert result.find("body/p:0").text == "Body {{company}} text"
    assert result.find("hdr:0:default/p:0").text == "Header {{date}} text"
    assert len(records) == 2
