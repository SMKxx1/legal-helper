"""Run-aware span→``{{token}}`` surgery (app.studio.tokenize_ops).

The tokenizer is the studio's only document-*writing* primitive, so these tests pin:

* formatting-inheritance goldens — mid-run / 2-run / 3+-run spans over runs with distinct
  bold/italic/size: the token run inherits the FIRST covered run's formatting, and the
  uncovered prefix/suffix keep exactly the formatting they had;
* surgery in every part — table cells, nested table cells, headers, footers;
* the full refusal matrix — stale hash, bad locator, out-of-bounds, cross-paragraph, overlap
  with an existing ``{{token}}``, empty/whitespace span, invalid token name, non-docx bytes,
  non-text (drawing) content in a covered run — all typed, all raised before any bytes exist;
* op records carry locator, range, exact replaced text, token, prior+new content hashes, and
  round-trip through dicts (the ``studio_ops.op_json`` persistence shape);
* undo restores the paragraph byte-faithfully (text AND per-run formatting, content hash equal
  to the pre-op document), redo reproduces the recorded post-op hash, and both refuse tampered
  records / wrong base bytes;
* the ``w:noBreakHyphen`` / ``w:ptab`` regression: composed run text includes them ("-" / "\\t"),
  so undo's snapshot-text reconstruction must too (this was a real severed-run defect: undo of a
  legitimate op on a paragraph containing a non-breaking hyphen refused with OpIntegrityError).
"""

from __future__ import annotations

import pytest
from conftest_studio import doc_to_bytes, rich_doc, runs_doc, single_para_doc
from docx import Document
from docx.shared import Pt
from lxml import etree

from app.studio.docview import (
    content_hash,
    extract_view,
    load_document,
    paragraph_text,
    resolve_locator,
)
from app.studio.errors import (
    BadDocxError,
    CrossParagraphSpanError,
    EmptySpanError,
    InvalidTokenNameError,
    LocatorNotFoundError,
    OpIntegrityError,
    RangeOutOfBoundsError,
    StaleViewError,
    TokenOverlapError,
    UnsupportedSpanError,
)
from app.studio.tokenize_ops import (
    OpRecord,
    apply_tokenize,
    redo_tokenize,
    undo_tokenize,
)

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _runs(docx_bytes: bytes, locator: str = "body/p:0"):
    """(text, bold, italic, size) per run of the paragraph at ``locator``."""
    paragraph = resolve_locator(load_document(docx_bytes), locator)
    return [(r.text, r.bold, r.italic, r.font.size) for r in paragraph.runs]


def _text(docx_bytes: bytes, locator: str = "body/p:0") -> str:
    return paragraph_text(resolve_locator(load_document(docx_bytes), locator))


# --------------------------------------------------------------------------- #
# Formatting-inheritance goldens
# --------------------------------------------------------------------------- #
def test_golden_mid_run_span_inherits_and_splits():
    data = runs_doc(("Hello dear world", {"bold": True}))
    new, record = apply_tokenize(data, "body/p:0", 6, 10, "who")
    assert _runs(new) == [
        ("Hello ", True, None, None),
        ("{{who}}", True, None, None),
        (" world", True, None, None),
    ]
    assert record.replaced_text == "dear"


def test_golden_two_run_span_inherits_first_covered_run():
    data = runs_doc(
        ("Hello BOLD", {"bold": True}),
        (" italic tail", {"italic": True, "size_pt": 9}),
    )
    text = "Hello BOLD italic tail"
    start = text.index("BOLD")
    end = start + len("BOLD italic")
    new, record = apply_tokenize(data, "body/p:0", start, end, "tok")
    # token takes the FIRST covered run's rPr (bold); the suffix keeps its own (italic, 9pt)
    assert _runs(new) == [
        ("Hello ", True, None, None),
        ("{{tok}}", True, None, None),
        (" tail", None, True, Pt(9)),
    ]
    assert record.replaced_text == "BOLD italic"


def test_golden_three_run_span_with_distinct_formats():
    data = runs_doc(
        ("AAAA", {"bold": True}),
        ("BBBB", {"italic": True}),
        ("CCCC", {"size_pt": 20}),
    )
    new, record = apply_tokenize(data, "body/p:0", 2, 10, "big")
    # middle run consumed entirely; boundary remainders keep their exact formatting
    assert _runs(new) == [
        ("AA", True, None, None),
        ("{{big}}", True, None, None),
        ("CC", None, None, Pt(20)),
    ]
    assert record.replaced_text == "AABBBBCC"  # text[2:10] of "AAAABBBBCCCC"


def test_golden_span_with_no_prefix_and_no_suffix():
    data = runs_doc(("AAAA", {"bold": True}), ("BBBB", {"italic": True}))
    # from paragraph start into run 2: no prefix — token replaces run 1's slot
    new, _ = apply_tokenize(data, "body/p:0", 0, 6, "t")
    assert _runs(new) == [("{{t}}", True, None, None), ("BB", None, True, None)]
    # to paragraph end: no suffix
    new, _ = apply_tokenize(data, "body/p:0", 2, 8, "t")
    assert _runs(new) == [("AA", True, None, None), ("{{t}}", True, None, None)]


def test_golden_span_exactly_covering_one_run():
    data = runs_doc(("pre ", {}), ("EXACT", {"bold": True}), (" post", {}))
    new, _ = apply_tokenize(data, "body/p:0", 4, 9, "t")
    assert _runs(new) == [
        ("pre ", None, None, None),
        ("{{t}}", True, None, None),
        (" post", None, None, None),
    ]


def test_golden_whole_paragraph_span():
    data = runs_doc(("whole", {"bold": True}))
    new, _ = apply_tokenize(data, "body/p:0", 0, 5, "all")
    assert _runs(new) == [("{{all}}", True, None, None)]


def test_offsets_over_tab_containing_text():
    doc = Document()
    doc.add_paragraph().add_run("Name:\tJohn Smith\there")
    data = doc_to_bytes(doc)
    text = extract_view(data).segments[0].text
    assert text == "Name:\tJohn Smith\there"
    start = text.index("John")
    new, record = apply_tokenize(
        data, "body/p:0", start, start + len("John Smith"), "name"
    )
    assert _text(new) == "Name:\t{{name}}\there"
    # a span COVERING the tab consumes it and records it in replaced_text
    new2, record2 = apply_tokenize(data, "body/p:0", 5, 16, "nm")
    assert _text(new2) == "Name:{{nm}}\there"
    assert record2.replaced_text == "\tJohn Smith"
    assert content_hash(undo_tokenize(new2, record2)) == content_hash(data)


def test_zero_width_drawing_run_inside_span_is_preserved():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("before ")
    image_run = paragraph.add_run("")  # zero-width run holding only a drawing
    etree.SubElement(image_run._r, f"{{{_W}}}drawing")
    paragraph.add_run("MIDDLE")
    paragraph.add_run(" after")
    data = doc_to_bytes(doc)
    text = extract_view(data).segments[0].text
    new, _ = apply_tokenize(data, "body/p:0", 0, text.index("MIDDLE") + 6, "span")
    kept = resolve_locator(load_document(new), "body/p:0")
    assert kept._p.findall(f".//{{{_W}}}drawing"), (
        "drawing run must survive the surgery"
    )
    assert paragraph_text(kept) == "{{span}} after"


# --------------------------------------------------------------------------- #
# Surgery in every document part
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize(
    ("locator", "span_text", "expected"),
    [
        ("body/tbl:0:0:1/p:0", "Plain", "{{t}} cell"),
        ("body/tbl:0:1:0/tbl:0:0:0/p:0", "COMPANY NAME", "nested {{t}} cell"),
        ("hdr:0:default/p:0", "ACME", "Header {{t}} text"),
        ("ftr:0:default/p:0", "Footer", "{{t}} fine print"),
    ],
)
def test_apply_in_tables_nested_tables_headers_footers(locator, span_text, expected):
    data = rich_doc()
    view = extract_view(data)
    segment = view.find(locator)
    start = segment.text.index(span_text)
    new, record = apply_tokenize(
        data,
        locator,
        start,
        start + len(span_text),
        "t",
        expected_hash=view.content_hash,
    )
    assert _text(new, locator) == expected
    assert record.locator == locator
    # and the op undoes cleanly wherever it happened
    assert content_hash(undo_tokenize(new, record)) == content_hash(data)


# --------------------------------------------------------------------------- #
# Refusal matrix — typed, and no bytes ever emitted
# --------------------------------------------------------------------------- #
def test_refusal_stale_view_hash():
    data = single_para_doc("Hello world")
    with pytest.raises(StaleViewError) as exc:
        apply_tokenize(data, "body/p:0", 0, 5, "t", expected_hash="0" * 64)
    assert exc.value.code == "studio_stale_view"
    assert exc.value.status == 409
    assert exc.value.details["actual_hash"] == content_hash(data)


def test_refusal_bad_locator():
    data = single_para_doc("Hello world")
    for locator in ("body/p:9", "body/tbl:0:0:0/p:0", "hdr:0:default/p:0", "junk"):
        with pytest.raises(LocatorNotFoundError) as exc:
            apply_tokenize(data, locator, 0, 5, "t")
        assert exc.value.code == "studio_locator_not_found"


@pytest.mark.parametrize(("start", "end"), [(-1, 5), (0, 999), (7, 3), (12, 15)])
def test_refusal_range_out_of_bounds(start, end):
    data = single_para_doc("Hello world")  # length 11
    with pytest.raises(RangeOutOfBoundsError) as exc:
        apply_tokenize(data, "body/p:0", start, end, "t")
    assert exc.value.code == "studio_range_out_of_bounds"
    assert exc.value.details == {"start": start, "end": end, "length": 11}


def test_refusal_cross_paragraph_selection():
    doc = Document()
    doc.add_paragraph("first paragraph")
    doc.add_paragraph("second paragraph")
    data = doc_to_bytes(doc)
    with pytest.raises(CrossParagraphSpanError) as exc:
        apply_tokenize(data, "body/p:0", 0, 5, "t", end_locator="body/p:1")
    assert exc.value.code == "studio_cross_paragraph"
    # same end locator is NOT cross-paragraph
    new, _ = apply_tokenize(data, "body/p:0", 0, 5, "t", end_locator="body/p:0")
    assert _text(new) == "{{t}} paragraph"


@pytest.mark.parametrize(
    ("start", "end"),
    [
        (4, 9),  # straddles the token's left edge
        (10, 16),  # straddles the right edge
        (6, 14),  # exactly the token
        (8, 10),  # strictly inside
        (0, 20),  # contains it
    ],
)
def test_refusal_overlap_with_existing_token(start, end):
    data = single_para_doc("Hello {{name}} world")  # token at [6, 14)
    with pytest.raises(TokenOverlapError) as exc:
        apply_tokenize(data, "body/p:0", start, end, "t")
    assert exc.value.code == "studio_token_overlap"
    assert exc.value.status == 409


def test_spans_adjacent_to_existing_token_are_allowed():
    data = single_para_doc("Hello {{name}} world")
    new, _ = apply_tokenize(data, "body/p:0", 0, 5, "greet")  # ends AT token start - 1
    assert _text(new) == "{{greet}} {{name}} world"
    new, _ = apply_tokenize(data, "body/p:0", 15, 20, "w")  # starts after the token
    assert _text(new) == "Hello {{name}} {{w}}"


def test_refusal_empty_and_whitespace_spans():
    data = single_para_doc("Hello world")
    for start, end in ((5, 5), (5, 6)):  # zero-width / lone space
        with pytest.raises(EmptySpanError) as exc:
            apply_tokenize(data, "body/p:0", start, end, "t")
        assert exc.value.code == "studio_empty_span"


@pytest.mark.parametrize("name", ["", "bad name", "{{x}}", "hy-phen", "a.b", None])
def test_refusal_invalid_token_name(name):
    data = single_para_doc("Hello world")
    with pytest.raises(InvalidTokenNameError) as exc:
        apply_tokenize(data, "body/p:0", 0, 5, name)
    assert exc.value.code == "studio_invalid_token_name"


def test_refusal_non_docx_bytes():
    with pytest.raises(BadDocxError) as exc:
        apply_tokenize(b"%PDF-1.7 not a docx", "body/p:0", 0, 5, "t")
    assert exc.value.code == "studio_bad_docx"


def test_refusal_covered_run_with_drawing():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("logo: ")
    image_run = paragraph.add_run("IMG")
    etree.SubElement(image_run._r, f"{{{_W}}}drawing")
    data = doc_to_bytes(doc)
    with pytest.raises(UnsupportedSpanError) as exc:
        apply_tokenize(data, "body/p:0", 0, 9, "t")  # span covers the IMG run
    assert exc.value.code == "studio_unsupported_span"


def test_refusals_never_alter_the_document():
    data = single_para_doc("Hello {{name}} world")
    before = content_hash(data)
    for exc_type, args in (
        (StaleViewError, dict(start=0, end=5, token_name="t", expected_hash="x")),
        (TokenOverlapError, dict(start=4, end=9, token_name="t")),
        (EmptySpanError, dict(start=5, end=5, token_name="t")),
        (RangeOutOfBoundsError, dict(start=0, end=99, token_name="t")),
        (InvalidTokenNameError, dict(start=0, end=5, token_name="")),
    ):
        with pytest.raises(exc_type):
            apply_tokenize(data, "body/p:0", **args)
    assert content_hash(data) == before
    assert extract_view(data).segments[0].text == "Hello {{name}} world"


# --------------------------------------------------------------------------- #
# Op records
# --------------------------------------------------------------------------- #
def test_op_record_carries_the_full_reversible_trail():
    data = runs_doc(("Alpha beta gamma", {"bold": True}))
    new, record = apply_tokenize(data, "body/p:0", 6, 10, "b")
    assert record.locator == "body/p:0"
    assert (record.start, record.end) == (6, 10)
    assert record.replaced_text == "beta"
    assert record.token_name == "b"
    assert record.token_text == "{{b}}"
    assert record.prior_hash == content_hash(data)
    assert record.new_hash == content_hash(new)
    assert record.prior_hash != record.new_hash
    assert record.paragraph_xml_before.startswith("<w:p")
    assert "Alpha beta gamma" in record.paragraph_xml_before


def test_op_record_round_trips_through_dict():
    _, record = apply_tokenize(single_para_doc("Alpha beta"), "body/p:0", 0, 5, "a")
    assert OpRecord.from_dict(record.to_dict()) == record


# --------------------------------------------------------------------------- #
# Undo / redo
# --------------------------------------------------------------------------- #
def test_undo_restores_text_and_formatting_byte_faithfully():
    data = runs_doc(
        ("AAAA", {"bold": True}),
        ("BBBB", {"italic": True}),
        ("CCCC", {"size_pt": 20}),
    )
    original_runs = _runs(data)
    new, record = apply_tokenize(data, "body/p:0", 2, 10, "big")
    restored = undo_tokenize(new, record)
    assert _runs(restored) == original_runs
    assert _text(restored) == "AAAABBBBCCCC"
    assert content_hash(restored) == record.prior_hash == content_hash(data)


def test_redo_reproduces_the_recorded_state_exactly():
    data = runs_doc(("Alpha beta gamma", {"bold": True}))
    new, record = apply_tokenize(data, "body/p:0", 6, 10, "b")
    restored = undo_tokenize(new, record)
    redone = redo_tokenize(restored, record)
    assert content_hash(redone) == record.new_hash
    assert _runs(redone) == _runs(new)


def test_undo_refuses_wrong_base_bytes():
    data = single_para_doc("Alpha beta")
    _, record = apply_tokenize(data, "body/p:0", 0, 5, "a")
    with pytest.raises(StaleViewError):  # undoing against the PRE-op bytes
        undo_tokenize(data, record)


def test_redo_refuses_wrong_base_bytes():
    data = single_para_doc("Alpha beta")
    new, record = apply_tokenize(data, "body/p:0", 0, 5, "a")
    with pytest.raises(StaleViewError):  # redoing against the POST-op bytes
        redo_tokenize(new, record)


def test_undo_refuses_a_tampered_record():
    data = single_para_doc("Alpha beta")
    new, record = apply_tokenize(data, "body/p:0", 0, 5, "a")
    tampered = OpRecord.from_dict({**record.to_dict(), "replaced_text": "Alphb"})
    with pytest.raises(OpIntegrityError) as exc:
        undo_tokenize(new, tampered)
    assert exc.value.code == "studio_op_integrity"


def test_redo_refuses_a_record_whose_replay_diverges():
    data = single_para_doc("Alpha beta")
    new, record = apply_tokenize(data, "body/p:0", 0, 5, "a")
    restored = undo_tokenize(new, record)
    tampered = OpRecord.from_dict({**record.to_dict(), "new_hash": "0" * 64})
    with pytest.raises(OpIntegrityError):
        redo_tokenize(restored, tampered)


def test_undo_depth_two_restores_the_original_document():
    data = single_para_doc("Alpha beta gamma delta")
    b1, r1 = apply_tokenize(data, "body/p:0", 0, 5, "a")
    text1 = extract_view(b1).segments[0].text
    start = text1.index("gamma")
    b2, r2 = apply_tokenize(b1, "body/p:0", start, start + 5, "g")
    assert extract_view(b2).segments[0].text == "{{a}} beta {{g}} delta"
    # LIFO undo: r2 then r1, hash-chained all the way back
    u1 = undo_tokenize(b2, r2)
    assert content_hash(u1) == r2.prior_hash == r1.new_hash
    u0 = undo_tokenize(u1, r1)
    assert content_hash(u0) == r1.prior_hash == content_hash(data)
    assert extract_view(u0).segments[0].text == "Alpha beta gamma delta"


# --------------------------------------------------------------------------- #
# The noBreakHyphen / ptab regression (severed-run defect, fixed)
# --------------------------------------------------------------------------- #
def _no_break_hyphen_doc() -> bytes:
    """'co-op agreement with COUNTERPARTY here' where the '-' is a w:noBreakHyphen."""
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("")
    run._r.append(etree.fromstring(f'<w:t xmlns:w="{_W}">co</w:t>'))
    run._r.append(etree.fromstring(f'<w:noBreakHyphen xmlns:w="{_W}"/>'))
    run._r.append(etree.fromstring(f'<w:t xmlns:w="{_W}">op</w:t>'))
    paragraph.add_run(" agreement with COUNTERPARTY here")
    return doc_to_bytes(doc)


def test_no_break_hyphen_composes_into_view_text():
    text = extract_view(_no_break_hyphen_doc()).segments[0].text
    assert text == "co-op agreement with COUNTERPARTY here"


def test_undo_and_redo_work_with_no_break_hyphen_in_uncovered_run():
    data = _no_break_hyphen_doc()
    text = extract_view(data).segments[0].text
    start = text.index("COUNTERPARTY")
    new, record = apply_tokenize(data, "body/p:0", start, start + 12, "cp")
    assert extract_view(new).segments[0].text == "co-op agreement with {{cp}} here"
    restored = undo_tokenize(new, record)  # refused with OpIntegrityError pre-fix
    assert content_hash(restored) == content_hash(data)
    assert content_hash(redo_tokenize(restored, record)) == record.new_hash


def test_covered_no_break_hyphen_refuses_rather_than_degrades():
    # consuming a noBreakHyphen would silently rewrite it as a plain '-': refusal instead
    with pytest.raises(UnsupportedSpanError):
        apply_tokenize(_no_break_hyphen_doc(), "body/p:0", 0, 5, "x")
