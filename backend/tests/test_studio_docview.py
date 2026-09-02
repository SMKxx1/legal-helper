"""Document view extraction (app.studio.docview): stable addressing over the filler's traversal.

The view is the wave-B studio page's render source, so the tests pin its contract hard:

* segment text is EXACTLY the filler's normalization (run-concatenated paragraph text, existing
  ``{{tokens}}`` verbatim);
* the traversal covers body + all tables (nested) + every owned header/footer variant, in order;
* every emitted locator resolves back (``resolve_locator``) to a paragraph with the same text —
  the round-trip that makes highlight offsets trustworthy;
* the embedded ``content_hash`` is stable across load→save cycles (serializer-independent C14N)
  and changes exactly when content changes — the stale-view seam;
* malformed locators and non-docx bytes refuse with typed errors.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from conftest_studio import doc_to_bytes, rich_doc, runs_doc, single_para_doc
from docx import Document

from app.studio.docview import (
    content_hash,
    extract_view,
    load_document,
    paragraph_text,
    resolve_locator,
)
from app.studio.errors import BadDocxError, LocatorNotFoundError


def test_body_paragraphs_in_order_with_run_concatenated_text():
    view = extract_view(rich_doc())
    body = [s for s in view.segments if s.locator.startswith("body/p:")]
    assert body[0].locator == "body/p:0"
    assert (
        body[0].text == "This agreement is between ACME CORPORATION and the recipient."
    )
    assert body[1].locator == "body/p:1"
    assert body[1].text == "Signed on [EFFECTIVE DATE] by the parties."
    assert all(s.kind == "body" for s in body)


def test_existing_tokens_appear_verbatim():
    view = extract_view(rich_doc())
    cell = view.find("body/tbl:0:0:0/p:0")
    assert cell is not None
    assert cell.text == "Cell with {{existing_token}}"


def test_table_and_nested_table_locators():
    view = extract_view(rich_doc())
    locators = [s.locator for s in view.segments]
    assert "body/tbl:0:0:1/p:0" in locators
    # the nested table lives in the (1,0) cell of the outer table
    nested = [loc for loc in locators if loc.startswith("body/tbl:0:1:0/tbl:0:0:0/")]
    assert nested, f"no nested-table locators in {locators}"
    seg = view.find(nested[0])
    assert seg.text == "nested COMPANY NAME cell"


def test_header_and_footer_segments_with_kinds():
    view = extract_view(rich_doc())
    hdr = view.find("hdr:0:default/p:0")
    ftr = view.find("ftr:0:default/p:0")
    assert hdr is not None and hdr.kind == "header" and hdr.text == "Header ACME text"
    assert ftr is not None and ftr.kind == "footer" and ftr.text == "Footer fine print"


def test_linked_header_variants_are_not_emitted():
    # rich_doc defines only the default header/footer; first/even are linked -> absent.
    view = extract_view(rich_doc())
    assert not any(
        ":first/" in s.locator or ":even/" in s.locator for s in view.segments
    )
    # and extraction did not CREATE parts as a read side effect (hash unchanged by extracting)
    data = rich_doc()
    before = content_hash(data)
    extract_view(data)
    assert content_hash(data) == before


def test_every_locator_round_trips_to_the_same_text():
    data = rich_doc()
    view = extract_view(data)
    doc = load_document(data)
    assert len(view.segments) > 6
    for segment in view.segments:
        paragraph = resolve_locator(doc, segment.locator)
        assert paragraph_text(paragraph) == segment.text, segment.locator


def test_merged_table_cell_is_listed_once():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].merge(table.rows[0].cells[1])
    table.rows[0].cells[0].paragraphs[0].add_run("merged")
    table.rows[1].cells[0].paragraphs[0].add_run("a")
    table.rows[1].cells[1].paragraphs[0].add_run("b")
    view = extract_view(doc_to_bytes(doc))
    merged = [s for s in view.segments if s.text == "merged"]
    assert len(merged) == 1
    assert merged[0].locator == "body/tbl:0:0:0/p:0"


def test_content_hash_stable_across_load_save_cycle():
    data = rich_doc()
    doc = Document(BytesIO(data))
    resaved = BytesIO()
    doc.save(resaved)
    assert content_hash(resaved.getvalue()) == content_hash(data)


def test_content_hash_changes_when_content_changes():
    data = rich_doc()
    doc = Document(BytesIO(data))
    doc.paragraphs[0].runs[0].text = "Changed "
    changed = BytesIO()
    doc.save(changed)
    assert content_hash(changed.getvalue()) != content_hash(data)


def test_content_hash_covers_headers_and_footers():
    data = rich_doc()
    doc = Document(BytesIO(data))
    doc.sections[0].header.paragraphs[0].runs[0].text = "New header"
    changed = BytesIO()
    doc.save(changed)
    assert content_hash(changed.getvalue()) != content_hash(data)


def test_view_embeds_matching_content_hash_and_serializes():
    data = single_para_doc("Hello {{name}}")
    view = extract_view(data)
    assert view.content_hash == content_hash(data)
    payload = view.to_dict()
    assert payload["content_hash"] == view.content_hash
    assert payload["segments"][0] == {
        "locator": "body/p:0",
        "text": "Hello {{name}}",
        "kind": "body",
    }


def test_split_runs_text_matches_filler_normalization():
    data = runs_doc(("Dear ", {}), ("{{", {}), ("na", {}), ("me", {}), ("}}", {}))
    view = extract_view(data)
    assert view.segments[0].text == "Dear {{name}}"


@pytest.mark.parametrize(
    "locator",
    [
        "body",  # too few segments
        "body/p:99",  # paragraph out of range
        "body/p:-1",  # negative index
        "body/p:x",  # non-numeric
        "body/tbl:5:0:0/p:0",  # table out of range
        "body/row:0/p:0",  # bad table segment tag
        "body/q:0",  # bad paragraph tag
        "hdr:0:first/p:0",  # linked header (no own definition)
        "hdr:9:default/p:0",  # no such section
        "hdr:0:sideways/p:0",  # bad variant
        "attic:0/p:0",  # bad part
    ],
)
def test_resolve_locator_refuses_bad_locators(locator):
    doc = load_document(rich_doc())
    with pytest.raises(LocatorNotFoundError):
        resolve_locator(doc, locator)


def test_non_docx_bytes_refuse_typed():
    with pytest.raises(BadDocxError):
        extract_view(b"%PDF-1.7 not a docx")
    with pytest.raises(BadDocxError):
        content_hash(b"garbage")
