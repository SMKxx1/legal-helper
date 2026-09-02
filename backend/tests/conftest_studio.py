"""Shared builders for the studio (document-surgery) test suite.

Plain helpers, importable as ``from conftest_studio import ...`` (pytest's prepend import mode puts
``tests/`` on ``sys.path`` — the ``conftest_bot`` convention). Documents are built with python-docx
in memory: no fixtures on disk, no network, no DB unless a test seeds one via
:func:`seed_draft_version`.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt


def doc_to_bytes(doc: Any) -> bytes:
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def single_para_doc(text: str) -> bytes:
    """One body paragraph holding ``text`` in a single run."""
    doc = Document()
    doc.add_paragraph(text)
    return doc_to_bytes(doc)


def runs_doc(*run_specs: tuple[str, dict]) -> bytes:
    """One body paragraph whose runs are exactly ``run_specs``: (text, {bold|italic|underline|size_pt})."""
    doc = Document()
    para = doc.add_paragraph()
    for text, fmt in run_specs:
        run = para.add_run(text)
        if fmt.get("bold") is not None:
            run.bold = fmt["bold"]
        if fmt.get("italic") is not None:
            run.italic = fmt["italic"]
        if fmt.get("underline") is not None:
            run.underline = fmt["underline"]
        if fmt.get("size_pt") is not None:
            run.font.size = Pt(fmt["size_pt"])
    return doc_to_bytes(doc)


def rich_doc() -> bytes:
    """A document exercising every traversal branch: formatted multi-run body paragraphs, a table,
    a nested table, an existing ``{{token}}``, and header + footer content."""
    doc = Document()
    p0 = doc.add_paragraph()
    r = p0.add_run("This agreement is between ")
    r.bold = True
    r = p0.add_run("ACME ")
    r.italic = True
    r = p0.add_run("CORPORATION")
    r.font.size = Pt(14)
    p0.add_run(" and the recipient.")
    doc.add_paragraph("Signed on [EFFECTIVE DATE] by the parties.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Cell with {{existing_token}}")
    table.rows[0].cells[1].paragraphs[0].add_run("Plain cell")
    nested = table.rows[1].cells[0].add_table(rows=1, cols=1)
    nested.rows[0].cells[0].paragraphs[0].add_run("nested COMPANY NAME cell")
    table.rows[1].cells[1].paragraphs[0].add_run("last cell")
    doc.sections[0].header.paragraphs[0].add_run("Header ACME text")
    doc.sections[0].footer.paragraphs[0].add_run("Footer fine print")
    return doc_to_bytes(doc)


def seed_draft_version(db: Any, docx_bytes: bytes) -> str:
    """Persist a draft ``TemplateVersion`` + its ``DocumentBlob`` (content-addressed); returns the
    version id. The parent ``template`` row is not created — the throwaway SQLite test engine does
    not enforce foreign keys, and the oplog only reads the version + blob."""
    import hashlib
    import uuid

    from sqlalchemy import select

    from app.models_v2 import DocumentBlob, TemplateVersion
    from app.support_task.generator import DOCX_MIME

    # Content-addressed like the table itself (sha256 is UNIQUE): seeding two drafts from the
    # same bytes must reuse the row, exactly as oplog._store_blob does.
    sha = hashlib.sha256(docx_bytes).hexdigest()
    blob = db.execute(
        select(DocumentBlob).where(DocumentBlob.sha256 == sha)
    ).scalar_one_or_none()
    if blob is None:
        blob = DocumentBlob(
            sha256=sha,
            byte_size=len(docx_bytes),
            mime_type=DOCX_MIME,
            bytes=docx_bytes,
        )
        db.add(blob)
        db.flush()
    version = TemplateVersion(
        template_id=uuid.uuid4().hex,
        variant_code="tokenised",
        version_no=1,
        blob_id=blob.id,
        is_current=False,
    )
    db.add(version)
    db.commit()
    return version.id


def draft_bytes(db: Any, version_id: str) -> bytes:
    """The current draft .docx bytes for a version (follows the re-pointed blob)."""
    from app.models_v2 import DocumentBlob, TemplateVersion

    version = db.get(TemplateVersion, version_id)
    blob = db.get(DocumentBlob, version.blob_id)
    return blob.bytes
