"""Live checklist (app.studio.checklist) — token accounting that agrees with generation.

The checklist's one hard invariant: it counts tokens exactly the way the generation side does
(``app.support_task.generator.fill_docx``'s run-joined paragraph text and the envelope guard's
``scan_docx_tokens``), so a token the checklist reports found is a token generation WILL fill —
split-across-runs placeholders included. Also pinned: first-seen document order across body,
tables (nested), headers and footers; required-list order for ``missing_required``; unknown
tokens carry the closest known name (or ``None`` when nothing is plausibly close).
"""

from __future__ import annotations

from conftest_studio import doc_to_bytes, runs_doc, single_para_doc
from docx import Document

from app.studio.checklist import analyze, scan_token_names
from app.studio.docview import extract_view


def _full_doc() -> bytes:
    """Tokens spread across every part, including one split across three runs."""
    doc = Document()
    doc.add_paragraph("Between {{company_name}} and {{counterparty_name}}.")
    split = doc.add_paragraph()
    split.add_run("Effective {{")
    split.add_run("effective_")
    split.add_run("date}} as written.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Sig: {{signatory_name}}")
    nested = table.rows[0].cells[1].add_table(rows=1, cols=1)
    nested.rows[0].cells[0].paragraphs[0].add_run("Law: {{governing_law_state}}")
    doc.sections[0].header.paragraphs[0].add_run("Ref {{document_ref}}")
    doc.sections[0].footer.paragraphs[0].add_run("Page of {{page_count}}")
    return doc_to_bytes(doc)


def test_scan_finds_tokens_in_document_order_across_all_parts():
    assert scan_token_names(_full_doc()) == [
        "company_name",
        "counterparty_name",
        "effective_date",  # split across three runs — still counted
        "signatory_name",
        "governing_law_state",
        "document_ref",
        "page_count",
    ]


def test_scan_counts_a_duplicate_token_once_first_seen():
    doc = Document()
    doc.add_paragraph("{{company_name}} then {{effective_date}}")
    doc.add_paragraph("again {{company_name}}")
    assert scan_token_names(doc_to_bytes(doc)) == ["company_name", "effective_date"]


def test_scan_accepts_a_pre_extracted_view():
    data = single_para_doc("Hello {{name}}")
    assert scan_token_names(extract_view(data)) == ["name"]


def test_scan_ignores_bracketish_non_tokens():
    data = single_para_doc("[COMPANY NAME] and <Company> and ____ and {single}")
    assert scan_token_names(data) == []


def test_analyze_reports_found_missing_and_unknown():
    result = analyze(
        _full_doc(),
        required_tokens=["company_name", "effective_date", "term_months"],
        known_tokens=[
            "company_name",
            "counterparty_name",
            "effective_date",
            "signatory_name",
            "governing_law_state",
            "term_months",
        ],
    )
    assert result["found"][0] == "company_name"
    assert result["missing_required"] == ["term_months"]  # required-list order
    assert {u["name"] for u in result["unknown"]} == {"document_ref", "page_count"}


def test_missing_required_preserves_required_list_order():
    data = single_para_doc("only {{b}} here")
    result = analyze(
        data, required_tokens=["z", "a", "b", "m"], known_tokens=["a", "b", "m", "z"]
    )
    assert result["missing_required"] == ["z", "a", "m"]


def test_unknown_tokens_carry_closest_known_or_none():
    data = single_para_doc("{{comany_name}} and {{zzqx}}")
    result = analyze(
        data, required_tokens=[], known_tokens=["company_name", "effective_date"]
    )
    assert result["unknown"] == [
        {"name": "comany_name", "closest_known": "company_name"},
        {"name": "zzqx", "closest_known": None},
    ]


def test_empty_document_and_empty_lists():
    data = single_para_doc("no tokens at all")
    assert analyze(data, [], []) == {"found": [], "missing_required": [], "unknown": []}
    assert analyze(data, ["a"], [])["missing_required"] == ["a"]


# --------------------------------------------------------------------------- #
# Agreement with the generation side (the checklist's core invariant)
# --------------------------------------------------------------------------- #
def test_checklist_agrees_with_fill_docx_on_split_token_docs():
    from app.support_task.generator import fill_docx

    data = _full_doc()
    found = scan_token_names(data)
    values = {name: f"V_{name}" for name in found}
    filled = fill_docx(data, values)
    # every token the checklist reported found was actually filled…
    assert scan_token_names(filled) == []
    text = " ".join(s.text for s in extract_view(filled).segments)
    for name in found:
        assert f"V_{name}" in text
    # …including the one that was split across three runs
    split_para = extract_view(filled).find("body/p:1")
    assert split_para.text == "Effective V_effective_date as written."


def test_checklist_agrees_with_the_envelope_unfilled_token_guard():
    from app.bot.intents.envelope import scan_docx_tokens

    data = _full_doc()
    guard_names = [t.strip("{}").strip() for t in scan_docx_tokens(data)]
    assert scan_token_names(data) == guard_names


def test_partially_filled_doc_shows_exactly_the_remaining_tokens():
    from app.support_task.generator import fill_docx

    data = _full_doc()
    filled = fill_docx(data, {"company_name": "ACME"}, strip_unfilled=False)
    remaining = scan_token_names(filled)
    assert "company_name" not in remaining
    assert set(remaining) == set(scan_token_names(data)) - {"company_name"}


def test_split_token_in_table_cell_counted_once():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell_para = table.rows[0].cells[0].paragraphs[0]
    cell_para.add_run("{{coun")
    cell_para.add_run("terparty_name}}")
    counted = scan_token_names(doc_to_bytes(doc))
    assert counted == ["counterparty_name"]


def test_formatting_split_does_not_double_count():
    data = runs_doc(
        ("Dear ", {}),
        ("{{", {"bold": True}),
        ("counterparty", {"italic": True}),
        ("_name}}", {}),
        (" and {{counterparty_name}} again", {}),
    )
    assert scan_token_names(data) == ["counterparty_name"]
