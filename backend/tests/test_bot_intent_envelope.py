"""The ``envelope`` intent — the three entry-point matrix (PLAN §3.9, reference §3.5).

Drives :class:`app.bot.intents.envelope.EnvelopeIntent` directly with fakes (a Slack file fetcher, a
thread scanner) + the throwaway ``bot_session_factory`` — zero network, zero LLM. Covers: the ≥2-signer
confirm card + stored state (the NEW confirm-before-send, PLAN §2 #1), the unfilled-``{{token}}`` refusal,
the <2-signer signer-details button, thread-doc recovery, the no-doc / email asks, and the invariant that
the intent NEVER sends (no ``nda_envelopes`` row is written until the Confirm click).
"""

from __future__ import annotations

import base64
import json
from io import BytesIO
from typing import Any

from docx import Document

from app.bot.envelope import AttachmentRef, Envelope
from app.bot.intents import IntentContext
from app.bot.intents.envelope import (
    DEFAULT_DOC_NAME,
    DOWNLOAD_FAILED_TEXT,
    EMAIL_HAS_DOC_TEXT,
    EMAIL_NO_DOC_TEXT,
    SLACK_NO_DOC_TEXT,
    EnvelopeIntent,
    scan_docx_tokens,
)
from app.bot.models import BotCorrelation
from app.bot.router import Classification
from app.bot.thread_docs import ThreadDoc
from app.integrations.models import NdaEnvelope

pytest_plugins = ("conftest_bot",)


# --------------------------------------------------------------------------- #
# Fixtures — real .docx bytes (clean vs tokenised, incl. table + header)
# --------------------------------------------------------------------------- #
def _docx(
    paragraphs: list[str], *, table_cell: str | None = None, header: str | None = None
) -> bytes:
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    if table_cell is not None:
        doc.add_table(rows=1, cols=1).cell(0, 0).text = table_cell
    if header is not None:
        doc.sections[0].header.paragraphs[0].text = header
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


CLEAN_DOCX = _docx(["This NDA is between Amperesand and Acme Corp."])
# A tokenised template with leftovers in the body, a table cell, and a header.
TOKENISED_DOCX = _docx(
    ["Between {{amperesand_signer_name}} and Acme.", "Address: {{street_address}}"],
    table_cell="{{country}}",
    header="{{jurisdiction}}",
)


def _attach(name: str = "nda.docx") -> AttachmentRef:
    return AttachmentRef(filename=name, source_ref="F1")


def _slack_env(
    *,
    attachments: tuple[AttachmentRef, ...] = (),
    thread_ts: str = "T1",
    channel: str = "C1",
    sender_id: str = "U1",
) -> Envelope:
    return Envelope(
        channel="slack",
        event_key="slack:E1",
        slack_channel=channel,
        slack_thread_ts=thread_ts,
        sender_id=sender_id,
        verified_sender=True,
        text="send this for signature",
        attachments=attachments,
    )


def _email_env(*, attachments: tuple[AttachmentRef, ...] = ()) -> Envelope:
    return Envelope(
        channel="email",
        event_key="email:E1",
        sender_address="lawyer@example.com",
        verified_sender=True,
        text="send this for signature",
        attachments=attachments,
    )


def _cls(
    *,
    signers: tuple[str, ...] = (),
    cc: tuple[str, ...] = (),
    sequential: bool = False,
    cc_timing: str = "after",
) -> Classification:
    return Classification(
        intent="envelope",
        signer_emails=signers,
        cc_emails=cc,
        sequential=sequential,
        cc_timing=cc_timing,
    )


def _intent(
    bot_session_factory: Any, *, fetch: Any = None, scan: Any = None
) -> EnvelopeIntent:
    return EnvelopeIntent(
        session_factory=bot_session_factory,
        slack_fetch=fetch or (lambda att: CLEAN_DOCX),
        thread_scan=scan or (lambda ch, ts: None),
    )


def _buttons(blocks: Any) -> dict[str, str]:
    out: dict[str, str] = {}
    for block in blocks or ():
        if block.get("type") == "actions":
            for el in block.get("elements", []):
                if el.get("type") == "button":
                    out[el["action_id"]] = el.get("value", "")
    return out


def _correlations(bot_session_factory: Any) -> list[dict[str, Any]]:
    with bot_session_factory() as s:
        return [dict(r.payload_json or {}) for r in s.query(BotCorrelation).all()]


def _envelope_count(bot_session_factory: Any) -> int:
    with bot_session_factory() as s:
        return s.query(NdaEnvelope).count()


# --------------------------------------------------------------------------- #
# 0) The unfilled-token guard (pure)
# --------------------------------------------------------------------------- #
def test_scan_docx_tokens_finds_body_table_and_header() -> None:
    found = scan_docx_tokens(TOKENISED_DOCX)
    assert set(found) == {
        "{{amperesand_signer_name}}",
        "{{street_address}}",
        "{{country}}",
        "{{jurisdiction}}",
    }


def test_scan_docx_tokens_clean_doc_is_empty() -> None:
    assert scan_docx_tokens(CLEAN_DOCX) == []


def test_scan_docx_tokens_non_docx_is_clean() -> None:
    # A PDF/DOC python-docx can't open has no {{}} by construction → scans clean (guard never over-refuses).
    assert scan_docx_tokens(b"%PDF-1.7 not a docx") == []


# --------------------------------------------------------------------------- #
# (a) ≥2 signers + clean attached doc → confirm card + stored state, NO send
# --------------------------------------------------------------------------- #
def test_two_signers_clean_doc_posts_confirm_card(bot_session_factory) -> None:
    intent = _intent(bot_session_factory)
    env = _slack_env(attachments=(_attach(),))
    cls = _cls(
        signers=("amp@a.com", "cp@b.com"),
        cc=("cc@x.com",),
        sequential=True,
        cc_timing="before",
    )
    reply = intent(IntentContext(envelope=env, classification=cls))

    buttons = _buttons(reply.slack_blocks)
    assert "env_confirm_send" in buttons and "env_confirm_cancel" in buttons
    # Both buttons carry the SAME correlation ref; the value is only {v, kind, action, ref}.
    send_val = json.loads(buttons["env_confirm_send"])
    assert send_val["kind"] == "env_confirm" and send_val["action"] == "send"
    ref = send_val["ref"]
    assert json.loads(buttons["env_confirm_cancel"])["ref"] == ref

    # Durable state: the doc + routing + requester mapping (PLAN §3.10) — not in the button value.
    (state,) = _correlations(bot_session_factory)
    assert state["signer_emails"] == ["amp@a.com", "cp@b.com"]
    assert state["cc_emails"] == ["cc@x.com"]
    assert state["routing"] == "amp_first"  # sequential → per-signer i+1
    assert state["cc_timing"] == "before"
    assert state["requested_by"] == "U1"
    assert state["slack_channel"] == "C1"
    assert state["slack_thread_ts"] == "T1"
    assert base64.b64decode(state["doc_b64"]) == CLEAN_DOCX

    # The invariant: the intent NEVER sends — no envelope row until the Confirm click.
    assert _envelope_count(bot_session_factory) == 0


def test_two_signers_non_sequential_is_all_at_once(bot_session_factory) -> None:
    intent = _intent(bot_session_factory)
    reply = intent(
        IntentContext(
            envelope=_slack_env(attachments=(_attach(),)),
            classification=_cls(signers=("a@a.com", "b@b.com"), sequential=False),
        )
    )
    assert _buttons(reply.slack_blocks).get("env_confirm_send")
    assert _correlations(bot_session_factory)[0]["routing"] == "all_at_once"


# --------------------------------------------------------------------------- #
# (a) token guard — a tokenised .docx is refused, never sent
# --------------------------------------------------------------------------- #
def test_tokenised_doc_is_refused(bot_session_factory) -> None:
    intent = _intent(bot_session_factory, fetch=lambda att: TOKENISED_DOCX)
    reply = intent(
        IntentContext(
            envelope=_slack_env(attachments=(_attach(),)),
            classification=_cls(signers=("a@a.com", "b@b.com")),
        )
    )
    assert reply.slack_blocks is None
    assert "unfilled placeholders" in reply.text
    assert "{{amperesand_signer_name}}" in reply.text
    assert "generate" in reply.text.lower()
    # Refused before any state is stored / anything is sent.
    assert _correlations(bot_session_factory) == []
    assert _envelope_count(bot_session_factory) == 0


# --------------------------------------------------------------------------- #
# (b) <2 signers → signer-details button (opens the modal)
# --------------------------------------------------------------------------- #
def test_under_two_signers_posts_signer_details_button(bot_session_factory) -> None:
    intent = _intent(bot_session_factory)
    reply = intent(
        IntentContext(
            envelope=_slack_env(attachments=(_attach(),)),
            classification=_cls(signers=("only@a.com",)),
        )
    )
    buttons = _buttons(reply.slack_blocks)
    assert "send_docusign" in buttons
    val = json.loads(buttons["send_docusign"])
    assert val["kind"] == "send_docusign"
    # State stored (doc + context) so the modal/confirm can resolve it by ref.
    (state,) = _correlations(bot_session_factory)
    assert base64.b64decode(state["doc_b64"]) == CLEAN_DOCX
    assert val["ref"]  # the button carries only the key
    assert _envelope_count(bot_session_factory) == 0


# --------------------------------------------------------------------------- #
# (c) no attachment → thread-doc recovery / no-doc / email asks
# --------------------------------------------------------------------------- #
def test_no_attachment_recovers_thread_doc(bot_session_factory) -> None:
    doc = ThreadDoc(
        file_id="Fthread", file_name="signed-nda.docx", file_url="https://x/y"
    )
    intent = _intent(bot_session_factory, scan=lambda ch, ts: doc)
    reply = intent(
        IntentContext(
            envelope=_slack_env(),  # no attachment
            classification=_cls(signers=("a@a.com", "b@b.com")),
        )
    )
    buttons = _buttons(reply.slack_blocks)
    assert "env_use_doc" in buttons and "decline_doc" in buttons
    (state,) = _correlations(bot_session_factory)
    # Thread doc stored as a REF (bytes fetched lazily on env_use_doc) — no doc_b64 yet.
    assert state["slack_file_id"] == "Fthread"
    assert state["file_name"] == "signed-nda.docx"
    assert "doc_b64" not in state
    assert (
        json.loads(buttons["env_use_doc"])["ref"]
        == json.loads(buttons["decline_doc"])["ref"]
    )


def test_no_attachment_no_thread_doc_asks_to_attach(bot_session_factory) -> None:
    intent = _intent(bot_session_factory, scan=lambda ch, ts: None)
    reply = intent(
        IntentContext(
            envelope=_slack_env(), classification=_cls(signers=("a@a.com", "b@b.com"))
        )
    )
    assert reply.slack_blocks is None
    assert reply.text == SLACK_NO_DOC_TEXT
    assert _correlations(bot_session_factory) == []


def test_email_no_doc_asks_to_attach(bot_session_factory) -> None:
    intent = _intent(bot_session_factory)
    reply = intent(
        IntentContext(
            envelope=_email_env(), classification=_cls(signers=("a@a.com", "b@b.com"))
        )
    )
    assert reply.text == EMAIL_NO_DOC_TEXT
    assert reply.slack_blocks is None


def test_email_with_doc_points_to_slack(bot_session_factory) -> None:
    intent = _intent(bot_session_factory)
    reply = intent(
        IntentContext(
            envelope=_email_env(attachments=(_attach(),)),
            classification=_cls(signers=("a@a.com", "b@b.com")),
        )
    )
    assert reply.text == EMAIL_HAS_DOC_TEXT
    # Email never sends nor stores interactive state.
    assert _correlations(bot_session_factory) == []
    assert _envelope_count(bot_session_factory) == 0


def test_fetch_failure_is_friendly(bot_session_factory) -> None:
    def boom(att: AttachmentRef) -> bytes:
        raise RuntimeError("slack download 404")

    intent = _intent(bot_session_factory, fetch=boom)
    reply = intent(
        IntentContext(
            envelope=_slack_env(attachments=(_attach(),)),
            classification=_cls(signers=("a@a.com", "b@b.com")),
        )
    )
    assert reply.text == DOWNLOAD_FAILED_TEXT
    assert _correlations(bot_session_factory) == []


def test_confirm_card_document_name_defaults(bot_session_factory) -> None:
    # An attachment with no filename still gets a sane document name in the summary.
    intent = _intent(bot_session_factory)
    intent(
        IntentContext(
            envelope=_slack_env(
                attachments=(AttachmentRef(filename="", source_ref="F1"),)
            ),
            classification=_cls(signers=("a@a.com", "b@b.com")),
        )
    )
    assert _correlations(bot_session_factory)[0]["file_name"] == DEFAULT_DOC_NAME
