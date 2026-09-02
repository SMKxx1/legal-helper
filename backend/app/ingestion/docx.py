"""Document parsing — .docx only (plan §2.1: uploads are always ``.docx``, no PDF/OCR/LibreOffice).

Turns an uploaded ``.docx`` into a flat ``full_text`` plus a list of :class:`TextBlock` fragments
(paragraphs / table rows) carrying light structural hints (headings, paragraph style). The review
pipeline only needs ``full_text``; ``blocks`` are kept for callers that want structure (e.g. a
future heading-aware UI) without re-parsing.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

# Numbered-section heuristic, e.g. "1.", "1.1", "1.2.3", "10. Foo".
_NUMBERED_RE = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s+\S")
# Article/Section style headings, e.g. "ARTICLE 1", "Section 3", "Clause 2.1".
_LABELLED_RE = re.compile(r"^\s*(?:ARTICLE|SECTION|CLAUSE)\b", re.IGNORECASE)


@dataclass(slots=True)
class TextBlock:
    """A single structural fragment of a document (a paragraph or table row)."""

    index: int
    text: str
    is_heading: bool = False
    style: str = ""


@dataclass(slots=True)
class ParsedDocument:
    """Normalised representation of a parsed document."""

    full_text: str
    blocks: list[TextBlock] = field(default_factory=list)
    file_format: str = "docx"
    source_path: str = ""


def _looks_like_heading(text: str) -> bool:
    """Heuristically decide whether a line is a section heading."""
    stripped = text.strip()
    if not stripped:
        return False
    if _NUMBERED_RE.match(stripped) or _LABELLED_RE.match(stripped):
        return True
    # Short Title-Case / ALL-CAPS line with no trailing sentence punctuation.
    words = stripped.split()
    if len(words) <= 10 and not stripped.endswith((".", ",", ";", ":")):
        alpha = [w for w in words if any(c.isalpha() for c in w)]
        if alpha and all(w[0].isupper() for w in alpha):
            return True
    return False


def _detect_format(path: Path, file_format: str | None) -> str:
    """Kept for parity with the old dispatcher's signature (tests exercise it directly). Only
    ``docx`` is a valid outcome — anything else is a caller error."""
    if file_format:
        fmt = file_format.lower().lstrip(".")
        if fmt != "docx":
            raise ValueError(
                f"Unsupported file_format {file_format!r}; only 'docx' is supported."
            )
        return fmt
    if path.suffix.lower() != ".docx":
        raise ValueError(
            f"Cannot detect format from extension {path.suffix!r} for {path}; "
            "only .docx is supported."
        )
    return "docx"


def parse_document(path: str | Path, file_format: str | None = None) -> ParsedDocument:
    """Parse a ``.docx`` file at ``path``. Raises ``ValueError`` for a missing file, a non-.docx
    extension, or an empty extraction."""
    p = Path(path)
    if not p.exists():
        raise ValueError(f"File not found: {p}")
    if not p.is_file():
        raise ValueError(f"Not a file: {p}")
    _detect_format(p, file_format)
    return extract_docx(p)


def _distinct_cells(row):
    """Row cells with horizontally-merged repeats collapsed to one. python-docx's ``row.cells``
    returns the SAME underlying ``_tc`` element once per grid column a merged cell spans; without
    this collapse a merged cell (and any table nested in it) is processed once per spanned column —
    duplicating its clause text into ``full_text`` and producing duplicate reviewed clauses."""
    prev = None
    for cell in row.cells:
        if prev is not None and cell._tc is prev:
            continue
        prev = cell._tc
        yield cell


def _cell_text(cell) -> str:
    """All text under a table cell INCLUDING nested tables. python-docx's ``cell.text`` returns ONLY
    the cell's DIRECT paragraphs, so a clause that lives inside a nested table (a matrix/signature
    block nested in a layout table) is otherwise silently dropped — never reviewed."""
    parts: list[str] = []
    direct = (cell.text or "").strip()
    if direct:
        parts.append(direct)
    for (
        nested
    ) in cell.tables:  # recurse to arbitrary depth; document.tables is top-level only
        for row in nested.rows:
            row_text = "\t".join(
                c for c in (_cell_text(cl) for cl in _distinct_cells(row)) if c
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_docx(path: str | Path) -> ParsedDocument:
    """Extract text from a .docx file at ``path``. See :func:`_extract` for the real work."""
    from docx import Document

    p = Path(path)
    return _extract(Document(str(p)), source_path=str(p))


def extract_docx_bytes(data: bytes) -> ParsedDocument:
    """Extract text from in-memory ``.docx`` bytes (an uploaded file — the API route never writes
    it to disk first). Same extraction as :func:`extract_docx`, ``source_path`` left blank."""
    import io

    from docx import Document

    return _extract(Document(io.BytesIO(data)), source_path="")


def _extract(document, *, source_path: str) -> ParsedDocument:
    """Walk a python-docx ``Document``'s body in document order (incl. tables, including nested
    ones). Body-level paragraphs and tables are INTERLEAVED in the OOXML body, so we walk
    ``document.element.body`` children in document order (rather than emitting all paragraphs
    then all tables) — a mid-document table must not be reordered to the end of ``full_text``.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    blocks: list[TextBlock] = []
    text_lines: list[str] = []
    idx = 0

    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")

    for child in document.element.body.iterchildren():
        if child.tag == p_tag:
            para = Paragraph(child, document)
            text = (para.text or "").strip()
            if not text:
                continue
            style_name = ""
            try:
                style_name = para.style.name or ""  # type: ignore[union-attr]
            except Exception:
                style_name = ""
            is_heading = style_name.startswith("Heading") or _looks_like_heading(text)
            blocks.append(
                TextBlock(index=idx, text=text, is_heading=is_heading, style=style_name)
            )
            text_lines.append(text)
            idx += 1
        elif child.tag == tbl_tag:
            # _cell_text recurses into nested tables so their clause text is not lost;
            # _distinct_cells collapses horizontally-merged repeats so a merged cell is emitted once.
            table = Table(child, document)
            for row in table.rows:
                cells = [_cell_text(cell) for cell in _distinct_cells(row)]
                row_text = "\t".join(c for c in cells if c)
                if not row_text:
                    continue
                blocks.append(TextBlock(index=idx, text=row_text, style="Table"))
                text_lines.append(row_text)
                idx += 1

    full_text = "\n".join(text_lines)
    if not full_text.strip():
        raise ValueError(
            f"No extractable text in docx file: {source_path or '<upload>'}"
        )
    return ParsedDocument(
        full_text=full_text,
        blocks=blocks,
        file_format="docx",
        source_path=source_path,
    )
