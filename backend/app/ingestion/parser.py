"""Document parsing.

Turn an uploaded NDA file (.docx/.pdf/.txt/.md) into a uniform
:class:`ParsedDocument`: a flat ``full_text`` plus a list of :class:`TextBlock`
fragments (paragraphs / lines) carrying light structural hints (headings,
page numbers, paragraph style). Downstream code (``app.ingestion.segmenter``)
turns this into clause-level segments.

The module is intentionally tolerant: it never assumes a perfectly clean
document and degrades gracefully (e.g. pdfplumber -> pypdf fallback). It raises
``ValueError`` only for genuinely unusable input (missing file, unsupported
format, empty extraction).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

# Numbered-section heuristic, e.g. "1.", "1.1", "1.2.3", "10. Foo".
_NUMBERED_RE = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s+\S")
# Article/Section style headings, e.g. "ARTICLE 1", "Section 3", "Clause 2.1".
_LABELLED_RE = re.compile(r"^\s*(?:ARTICLE|SECTION|CLAUSE)\b", re.IGNORECASE)

_SUPPORTED_FORMATS = {"docx", "pdf", "txt", "md"}
_EXT_TO_FORMAT = {
    ".docx": "docx",
    ".pdf": "pdf",
    ".txt": "txt",
    ".md": "md",
    ".markdown": "md",
    ".text": "txt",
}


class NoTextLayerError(ValueError):
    """Raised when a PDF has no extractable text layer (e.g. a scanned image).

    Subclasses ``ValueError`` so existing ``except ValueError`` callers keep
    working, while letting the review service distinguish "needs OCR" from other
    parse failures and surface a clear, actionable state to the user.
    """


@dataclass(slots=True)
class TextBlock:
    """A single structural fragment of a document (a paragraph or line)."""

    index: int
    text: str
    is_heading: bool = False
    page: int | None = None
    style: str = ""


@dataclass(slots=True)
class ParsedDocument:
    """Normalised representation of a parsed document."""

    full_text: str
    blocks: list[TextBlock] = field(default_factory=list)
    file_format: str = ""
    source_path: str = ""


def _looks_like_heading(text: str) -> bool:
    """Heuristically decide whether a line is a section heading."""
    stripped = text.strip()
    if not stripped:
        return False
    if _NUMBERED_RE.match(stripped) or _LABELLED_RE.match(stripped):
        return True
    # Short Title-Case / ALL-CAPS line with no trailing sentence punctuation.
    words = stripped.split()
    if len(words) <= 10 and not stripped.endswith((".", ",", ";", ":")):
        alpha = [w for w in words if any(c.isalpha() for c in w)]
        if alpha and all(w[0].isupper() for w in alpha):
            return True
    return False


def _detect_format(path: Path, file_format: str | None) -> str:
    if file_format:
        fmt = file_format.lower().lstrip(".")
        if fmt not in _SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file_format {file_format!r}; "
                f"expected one of {sorted(_SUPPORTED_FORMATS)}."
            )
        return fmt
    detected = _EXT_TO_FORMAT.get(path.suffix.lower())
    if detected is None:
        raise ValueError(
            f"Cannot detect format from extension {path.suffix!r} for "
            f"{path}; pass file_format explicitly."
        )
    return detected


def parse_document(path: str | Path, file_format: str | None = None) -> ParsedDocument:
    """Parse a document at ``path`` into a :class:`ParsedDocument`.

    Auto-detects the format from the file extension when ``file_format`` is
    ``None``. Raises ``ValueError`` for a missing file or unsupported format.
    """
    p = Path(path)
    if not p.exists():
        raise ValueError(f"File not found: {p}")
    if not p.is_file():
        raise ValueError(f"Not a file: {p}")

    fmt = _detect_format(p, file_format)
    if fmt == "docx":
        return extract_docx(p)
    if fmt == "pdf":
        return extract_pdf(p)
    # txt / md
    return extract_text_file(p, file_format=fmt)


def _distinct_cells(row):
    """Row cells with horizontally-merged repeats collapsed to one. python-docx's ``row.cells``
    returns the SAME underlying ``_tc`` element once per grid column a merged cell spans; without
    this collapse a merged cell (and any table nested in it) is processed once per spanned column —
    duplicating its clause text into ``full_text`` and producing duplicate reviewed clauses."""
    prev = None
    for cell in row.cells:
        if prev is not None and cell._tc is prev:
            continue
        prev = cell._tc
        yield cell


def _cell_text(cell) -> str:
    """All text under a table cell INCLUDING nested tables. python-docx's ``cell.text`` returns ONLY
    the cell's DIRECT paragraphs, so a clause that lives inside a nested table (a matrix/signature
    block nested in a layout table) is otherwise silently dropped — never segmented, never reviewed."""
    parts: list[str] = []
    direct = (cell.text or "").strip()
    if direct:
        parts.append(direct)
    for (
        nested
    ) in cell.tables:  # recurse to arbitrary depth; document.tables is top-level only
        for row in nested.rows:
            row_text = "\t".join(
                c for c in (_cell_text(cl) for cl in _distinct_cells(row)) if c
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_docx(path: str | Path) -> ParsedDocument:
    """Extract text from a .docx file using python-docx (incl. tables, including nested ones).

    Body-level paragraphs and tables are INTERLEAVED in the OOXML body, so we walk
    ``document.element.body`` children in document order (rather than emitting all
    paragraphs then all tables). Emitting paragraphs-then-tables reordered a mid-document
    table to the end of ``full_text``, drifting clause ``start_char``/``end_char`` offsets
    and segmentation out of sync with the source."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    p = Path(path)
    document = Document(str(p))

    blocks: list[TextBlock] = []
    text_lines: list[str] = []
    idx = 0

    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")

    for child in document.element.body.iterchildren():
        if child.tag == p_tag:
            para = Paragraph(child, document)
            text = (para.text or "").strip()
            if not text:
                continue
            style_name = ""
            try:
                style_name = para.style.name or ""  # type: ignore[union-attr]
            except Exception:
                style_name = ""
            is_heading = style_name.startswith("Heading") or _looks_like_heading(text)
            blocks.append(
                TextBlock(index=idx, text=text, is_heading=is_heading, style=style_name)
            )
            text_lines.append(text)
            idx += 1
        elif child.tag == tbl_tag:
            # Pull text out of the table (cell-by-cell, row joined with tabs). ``Table`` wraps
            # only THIS top-level w:tbl; _cell_text recurses into nested tables so their clause
            # text is not lost. _distinct_cells collapses horizontally-merged repeats so a merged
            # cell (and its nested table) is emitted once.
            table = Table(child, document)
            for row in table.rows:
                cells = [_cell_text(cell) for cell in _distinct_cells(row)]
                row_text = "\t".join(c for c in cells if c)
                if not row_text:
                    continue
                blocks.append(TextBlock(index=idx, text=row_text, style="Table"))
                text_lines.append(row_text)
                idx += 1

    full_text = "\n".join(text_lines)
    if not full_text.strip():
        raise ValueError(f"No extractable text in docx file: {p}")
    return ParsedDocument(
        full_text=full_text,
        blocks=blocks,
        file_format="docx",
        source_path=str(p),
    )


def extract_pdf(path: str | Path) -> ParsedDocument:
    """Extract text from a .pdf file, preferring pdfplumber then pypdf."""
    p = Path(path)

    blocks: list[TextBlock] = []
    page_texts: list[str] = []
    idx = 0

    # Extract page-by-page so a SINGLE bad page can't discard every page extracted before it
    # (the prior whole-loop try/except reset blocks/page_texts on any page error and re-ran the
    # weaker pypdf path over the whole doc, throwing away good pdfplumber text). page_map holds
    # page_no -> text from whichever extractor yielded it; pages are assembled in order below.
    page_map: dict[int, str] = {}
    total_pages = 0

    try:
        import pdfplumber  # noqa: PLC0415
    except ImportError:
        pdfplumber = None  # type: ignore[assignment]

    if pdfplumber is not None:
        try:
            with pdfplumber.open(str(p)) as pdf:
                pages = pdf.pages
                total_pages = len(pages)
                for page_no, page in enumerate(pages, start=1):
                    try:
                        page_text = page.extract_text() or ""
                    except Exception:  # noqa: BLE001 — one page must not poison the rest
                        import logging

                        logging.getLogger(__name__).warning(
                            "pdfplumber failed on page %s of %s; will try pypdf for it",
                            page_no,
                            p,
                        )
                        continue
                    if page_text.strip():
                        page_map[page_no] = page_text
        except Exception:  # noqa: BLE001 — open()/iterate failed; let pypdf do the whole doc
            page_map = {}

    # Fill pages pdfplumber couldn't yield (or the whole doc when pdfplumber was unavailable) via
    # pypdf, page-by-page, WITHOUT re-extracting pages pdfplumber already got.
    if not page_map or len(page_map) < total_pages:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(p))
            n = len(reader.pages)
            total_pages = max(total_pages, n)
            for page_no in range(1, n + 1):
                if page_no in page_map:
                    continue
                try:
                    page_text = reader.pages[page_no - 1].extract_text() or ""
                except Exception:  # noqa: BLE001 — skip a page pypdf can't read either
                    continue
                if page_text.strip():
                    page_map[page_no] = page_text
        except Exception as exc:  # pragma: no cover - both extractors failed
            if not page_map:
                raise ValueError(
                    f"Failed to extract text from pdf file {p}: {exc}"
                ) from exc
            # pdfplumber yielded SOME pages; keep them even if pypdf blew up opening the file.

    for page_no in sorted(page_map):
        page_text = page_map[page_no]
        page_texts.append(page_text)
        for line in page_text.splitlines():
            line = line.strip()
            if not line:
                continue
            blocks.append(
                TextBlock(
                    index=idx,
                    text=line,
                    is_heading=_looks_like_heading(line),
                    page=page_no,
                )
            )
            idx += 1

    full_text = "\n".join(page_texts)
    # Scanned / image-only PDFs have no usable text layer. Many still carry a
    # *thin* text overlay (e.g. a DocuSign signature/date stamp) that extracts to
    # a few characters, so a "zero text" check is not enough — trigger fully-local
    # OCR (zero egress) whenever the text layer is too sparse per page, and keep
    # whichever source produced more text. Transparent to every caller.
    if _looks_scanned(p, full_text):
        ocr_texts = _ocr_pdf(p)
        ocr_full = "\n".join(t for t in ocr_texts if t.strip()).strip()
        if len(ocr_full) > len(full_text.strip()):
            blocks = []
            page_texts = []
            idx = 0
            for page_no, page_text in enumerate(ocr_texts, start=1):
                if not page_text.strip():
                    continue
                page_texts.append(page_text)
                for line in page_text.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    blocks.append(
                        TextBlock(
                            index=idx,
                            text=line,
                            is_heading=_looks_like_heading(line),
                            page=page_no,
                        )
                    )
                    idx += 1
            full_text = "\n".join(page_texts)
    if not full_text.strip():
        raise NoTextLayerError(f"No extractable text in pdf file: {p}")
    return ParsedDocument(
        full_text=full_text,
        blocks=blocks,
        file_format="pdf",
        source_path=str(p),
    )


#: A real contract page holds well over this many characters; a scanned page with
#: only a signature/date stamp holds far fewer. Below this average we attempt OCR.
_MIN_CHARS_PER_PAGE = 100


def _looks_scanned(p: Path, text: str) -> bool:
    """True when the extracted text layer is too sparse to be a real text PDF."""
    try:
        from app.ingestion import ocr

        if not ocr.ocr_available():
            return False
    except Exception:  # noqa: BLE001
        return False
    try:
        import fitz  # PyMuPDF

        with fitz.open(str(p)) as doc:
            pages = max(1, doc.page_count)
    except Exception:  # noqa: BLE001 — fall back to a single-page assumption
        pages = 1
    return len(text.strip()) < pages * _MIN_CHARS_PER_PAGE


def _ocr_pdf(p: Path) -> list[str]:
    """Local OCR fallback for a scanned PDF; never raises."""
    try:
        from app.ingestion import ocr

        return ocr.ocr_pdf_pages(p)
    except Exception:  # noqa: BLE001 — OCR is best-effort; degrade to NoTextLayerError
        import logging

        logging.getLogger(__name__).exception("local OCR fallback failed for %s", p)
        return []


def extract_text_file(
    path: str | Path, file_format: str | None = None
) -> ParsedDocument:
    """Read a plain-text or markdown file and split it into line blocks."""
    p = Path(path)
    # utf-8-sig strips a leading BOM (﻿) at decode time — otherwise it rides on the first line
    # and breaks first-section/heading detection, silently mislabeling the opening clause.
    raw = p.read_text(encoding="utf-8-sig", errors="replace")

    fmt = (file_format or _EXT_TO_FORMAT.get(p.suffix.lower()) or "txt").lower()
    if fmt not in {"txt", "md"}:
        fmt = "txt"

    blocks: list[TextBlock] = []
    idx = 0
    for line in raw.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        is_heading = stripped.startswith("#") or _looks_like_heading(stripped)
        # Normalise markdown ATX headings ("## Foo" -> "Foo") for the block text
        # while keeping the heading flag; full_text retains the raw content.
        display = stripped.lstrip("#").strip() if stripped.startswith("#") else stripped
        blocks.append(TextBlock(index=idx, text=display, is_heading=is_heading))
        idx += 1

    full_text = raw
    if not full_text.strip():
        raise ValueError(f"No extractable text in file: {p}")
    return ParsedDocument(
        full_text=full_text,
        blocks=blocks,
        file_format=fmt,
        source_path=str(p),
    )
