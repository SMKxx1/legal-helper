"""Reconstruct ORIGINAL vs REDLINED plaintext from a tracked-changes .docx.

Read-side counterpart to ``app/redline/docx_writer.py`` (which WRITES tracked changes). python-docx's
high-level ``Paragraph.text`` drops BOTH ``<w:ins>`` and ``<w:del>`` content, so the normal
``app.ingestion.parser.extract_docx`` path silently ignores redlines. To review only the redlined
parts we must descend to the OOXML run level and emit text per side:

  * ORIGINAL  (tracked changes REJECTED): drop text under ``<w:ins>``/``<w:moveTo>``; KEEP normal
    ``<w:t>`` and the deleted ``<w:delText>`` (it WAS in the original).
  * REDLINED  (tracked changes ACCEPTED): drop text under ``<w:del>``/``<w:moveFrom>``; KEEP normal
    and inserted ``<w:t>``; never emit ``<w:delText>``.

Both sides are assembled with the SAME paragraph-then-table ordering + ``"\n"``/``"\t"`` joining as
``extract_docx`` (reusing ``_distinct_cells`` for horizontally-merged cells), so unchanged regions are
byte-identical between the two versions and ``segment_clauses`` + ``align_clauses`` surface only the
genuinely-changed clauses.
"""

from __future__ import annotations

import io

from docx.oxml.ns import qn

from app.ingestion.parser import _distinct_cells

# Qualified OOXML tag names (resolved once).
_W_T = qn("w:t")  # normal / inserted run text
_W_DELTEXT = qn("w:delText")  # deleted run text (always under <w:del>)
_W_TAB = qn("w:tab")
_W_BR = qn("w:br")
_W_CR = qn("w:cr")
_W_INS = qn("w:ins")
_W_DEL = qn("w:del")
_W_MOVEFROM = qn("w:moveFrom")
_W_MOVETO = qn("w:moveTo")

#: Revision containers whose text is excluded from each reconstructed side.
_ORIG_DROP = (_W_INS, _W_MOVETO)  # original = changes rejected
_RED_DROP = (_W_DEL, _W_MOVEFROM)  # redlined = changes accepted

_TRACKED_TAGS = frozenset({_W_INS, _W_DEL, _W_MOVEFROM, _W_MOVETO})


def _has_ancestor(el, drop: tuple[str, ...]) -> bool:
    """True if ``el`` sits anywhere under one of the ``drop`` revision containers."""
    cur = el.getparent()
    while cur is not None:
        if cur.tag in drop:
            return True
        cur = cur.getparent()
    return False


def _node_text(element, drop: tuple[str, ...]) -> str:
    """Document-ordered text of an element subtree, excluding anything under a ``drop`` container.

    ``<w:delText>`` is treated as text (kept on the side that hasn't dropped ``<w:del>``); ``<w:tab>``
    → tab, ``<w:br>``/``<w:cr>`` → newline. Run text is read verbatim (honors ``xml:space``); the
    caller strips the assembled line."""
    out: list[str] = []
    for el in element.iter():
        tag = el.tag
        if tag in (_W_T, _W_DELTEXT):
            if _has_ancestor(el, drop):
                continue
            out.append(el.text or "")
        elif tag == _W_TAB:
            if not _has_ancestor(el, drop):
                out.append("\t")
        elif tag in (_W_BR, _W_CR):
            if not _has_ancestor(el, drop):
                out.append("\n")
    return "".join(out)


def _cell_text_side(cell, drop: tuple[str, ...]) -> str:
    """Per-side text of a table cell INCLUDING nested tables (mirror of ``parser._cell_text``)."""
    parts: list[str] = []
    for para in cell.paragraphs:
        t = _node_text(para._p, drop).strip()
        if t:
            parts.append(t)
    for (
        nested
    ) in cell.tables:  # recurse to arbitrary depth (document.tables is top-level only)
        for row in nested.rows:
            row_text = "\t".join(
                c
                for c in (_cell_text_side(cl, drop) for cl in _distinct_cells(row))
                if c
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_side(document, drop: tuple[str, ...]) -> str:
    """Assemble one reconstructed side, mirroring ``extract_docx``'s paragraph-then-table order."""
    lines: list[str] = []
    for para in document.paragraphs:
        t = _node_text(para._p, drop).strip()
        if t:
            lines.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [_cell_text_side(cell, drop) for cell in _distinct_cells(row)]
            row_text = "\t".join(c for c in cells if c)
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def extract_redline_versions(data: bytes) -> tuple[str, str]:
    """``(original_text, redlined_text)`` reconstructed from a tracked-changes .docx (see module docs).

    Raises ``ValueError`` on an unreadable docx or when either side reconstructs to empty (e.g. a
    document whose entire body is a single tracked insertion → no original)."""
    from docx import Document

    try:
        document = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001 — surface as ValueError (parity with extract_docx)
        raise ValueError(f"Could not open the .docx for redline extraction: {e}") from e

    original = _extract_side(document, _ORIG_DROP)
    redlined = _extract_side(document, _RED_DROP)
    if not original.strip():
        raise ValueError(
            "The original (changes-rejected) text is empty — nothing to compare against."
        )
    if not redlined.strip():
        raise ValueError("The redlined (changes-accepted) text is empty.")
    return original, redlined


def has_tracked_changes(data: bytes) -> bool:
    """True iff the docx body contains any tracked-change container (ins/del/moveFrom/moveTo).

    Cheap structural probe used to return a clean ``no_redlines`` error instead of running an empty
    review. Never raises — a malformed/unreadable docx returns ``False`` (the caller then surfaces a
    clear message)."""
    try:
        from docx import Document

        document = Document(io.BytesIO(data))
        return any(el.tag in _TRACKED_TAGS for el in document.element.body.iter())
    except Exception:  # noqa: BLE001 — best-effort probe
        return False
