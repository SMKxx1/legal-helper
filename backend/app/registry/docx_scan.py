"""Scan a stored .docx blob for the ``{{token}}`` placeholders it contains (PLAN §3.7).

The inverse of :mod:`app.support_task.generator`'s filler: where the filler walks the body, tables
(recursively), and section headers/footers joining run text to *replace* placeholders, this walks the
same containers joining run text to *find* them. Joining per paragraph is what makes a placeholder that
Word fragmented across several runs (``{{`` / ``counter`` / ``party_name}}``) still resolve — a naive
byte regex over the zipped blob would miss it (the XML is deflate-compressed and run-split).

Pure read-only + fail-soft: an unreadable / non-docx / ``None`` blob yields an empty set (a garbage or
not-yet-loaded template blob must never crash a usage scan or a generation guard).
"""

from __future__ import annotations

import re
from io import BytesIO

from ..telemetry import get_logger

log = get_logger("nda.registry.docx_scan")

#: A ``{{token_name}}`` placeholder — the same snake_case token shape the generator fills.
_PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")


def _container_paragraph_texts(container) -> list[str]:
    """Joined run text for every paragraph in a docx container (Document / _Cell / header / footer),
    recursing into tables so a placeholder inside a table cell is seen too."""
    texts: list[str] = []
    for paragraph in container.paragraphs:
        texts.append("".join(run.text for run in paragraph.runs))
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(_container_paragraph_texts(cell))
    return texts


def scan_docx_tokens(docx_bytes: bytes | None) -> set[str]:
    """Return the set of token NAMES (bare, no braces) referenced by ``{{…}}`` in ``docx_bytes``.

    Walks the body, all tables (recursively), and every section header/footer part, joining each
    paragraph's runs before matching so run-split placeholders are found. Returns an empty set for a
    ``None`` / empty / unreadable blob (fail-soft).
    """
    if not docx_bytes:
        return set()
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001 — python-docx should be present; degrade rather than crash
        log.warning("registry.docx_scan.no_docx", error=repr(exc))
        return set()
    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as exc:  # noqa: BLE001 — non-docx / corrupt bytes: a token scan finds nothing
        log.warning("registry.docx_scan.unreadable", error=repr(exc))
        return set()

    texts = _container_paragraph_texts(doc)
    for section in doc.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            try:
                texts.extend(_container_paragraph_texts(part))
            except Exception:  # noqa: BLE001 — a missing/linked header part is non-fatal
                pass

    found: set[str] = set()
    for text in texts:
        if "{{" in text:
            found.update(_PLACEHOLDER_RE.findall(text))
    return found


__all__ = ["scan_docx_tokens"]
