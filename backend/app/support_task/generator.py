"""Tokenised-.docx generation.

``docxtpl`` is intentionally NOT a dependency (only ``python-docx`` is available), so we do the
``{{token}}`` substitution ourselves with a run-aware two-pass replace:

1. fast path — replace inside each individual run (placeholder fully contained in one run);
   this preserves all inline formatting, which is the common case.
2. split path — if a placeholder is split across runs (Word often fragments typed text into
   several runs), re-join the paragraph's run text, replace, and collapse into the first run.
   Only paragraphs that still contain ``{{`` after pass 1 are touched, so formatting loss is
   limited to paragraphs that actually carried a split placeholder.

Replacement runs over the body, all tables (recursively), and section headers/footers.
Tokens with no supplied value are left as-is (``{{token}}``) so missing data is visible rather
than silently blanked — the caller (n8n) is responsible for sending the full token table.
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document

from app.api.errors import EngineError

# An unfilled token placeholder, plus an immediately-preceding separator (", " / spaces) to consume
# when removing it. So "{{city_zip}}, {{country}}" with no country renders "…Shenzhen" rather than a
# dangling "…Shenzhen, " — the comma is kept only when the following token actually has a value.
# All our template placeholders are {{token_name}} of this shape.
_STRIP_RE = re.compile(r"[ \t]*,?[ \t]*\{\{\s*[A-Za-z0-9_]+\s*\}\}")

# Parentheses left empty after the only token inside them was stripped (e.g. a company's
# "(... {{counterparty_company_registration_number}})" with no number -> "()"). Removed together
# with a preceding space so "incorporated in ()" renders "incorporated in".
_EMPTY_PARENS_RE = re.compile(r"[ \t]*\([ \t]*\)")

# Lenient code normalisation: callers (the Slack bot / Tally) speak lowercase/underscored codes;
# the schema's ref_* tables use CamelCase. Map to the canonical ref codes.
_JUR = {"us": "US", "sg": "SG"}
_CP = {
    "company": "Company",
    "individual": "Individual",
    "serviceprovider": "ServiceProvider",
    "service_provider": "ServiceProvider",
    "service provider": "ServiceProvider",
    "sp": "ServiceProvider",
}
_MUT = {
    "mutual": "Mutual",
    "unilateral": "Unilateral",
    "notapplicable": "NotApplicable",
    "not_applicable": "NotApplicable",
    "na": "NotApplicable",
    "n/a": "NotApplicable",
    "": "NotApplicable",
}

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def normalize_codes(
    jurisdiction: str | None, counterparty_type: str | None, mutuality: str | None
) -> tuple[str, str, str]:
    """Map lenient inputs to the canonical (jurisdiction, counterparty_type, mutuality) ref codes,
    enforcing the same invariant as the DB CHECK: mutuality applies only to Individual."""
    jur = _JUR.get((jurisdiction or "").strip().lower())
    cp = _CP.get((counterparty_type or "").strip().lower())
    if not jur or not cp:
        raise EngineError(
            400,
            "bad_request",
            f"Unknown jurisdiction/counterparty: {jurisdiction!r} / {counterparty_type!r}.",
        )
    if cp == "Individual":
        mut = _MUT.get((mutuality or "").strip().lower(), "NotApplicable")
        if (
            mut == "NotApplicable"
        ):  # Individual must be Mutual or Unilateral; default to Mutual
            mut = "Mutual"
    else:
        mut = "NotApplicable"
    return jur, cp, mut


def _build_replacements(values: dict) -> dict[str, str]:
    """Accept token names ('counterparty_name') or full placeholders ('{{counterparty_name}}')
    as keys; return a {placeholder: value} map."""
    repl: dict[str, str] = {}
    for key, val in (values or {}).items():
        if key is None:
            continue
        k = str(key).strip()
        if not k:
            continue
        placeholder = (
            k if (k.startswith("{{") and k.endswith("}}")) else "{{" + k + "}}"
        )
        repl[placeholder] = "" if val is None else str(val)
    return repl


def _fill_paragraph(paragraph, repl: dict[str, str], strip_unfilled: bool) -> bool:
    """Fill/strip placeholders in one paragraph. Returns True if an unfilled token left the
    paragraph empty (the caller may then drop the now-blank line)."""
    had_token = any("{{" in run.text for run in paragraph.runs)
    # pass 1: per-run (preserves formatting; handles single-run placeholders)
    for run in paragraph.runs:
        text = run.text
        if "{{" in text:
            for ph, val in repl.items():
                if ph in text:
                    text = text.replace(ph, val)
            if (
                strip_unfilled and "{{" in text
            ):  # drop any placeholder left unfilled (+ its separator)
                text = _STRIP_RE.sub("", text)
            run.text = text
    # clean up parentheses left empty by a stripped token, when both parens are in one run
    if strip_unfilled:
        for run in paragraph.runs:
            if (
                "(" in run.text
                and ")" in run.text
                and _EMPTY_PARENS_RE.search(run.text)
            ):
                run.text = _EMPTY_PARENS_RE.sub("", run.text)
    # pass 2: placeholder/parens split across runs (re-join, replace/strip/clean, collapse to run 0)
    joined = "".join(run.text for run in paragraph.runs)
    needs_join = ("{{" in joined and "}}" in joined) or (
        strip_unfilled and bool(_EMPTY_PARENS_RE.search(joined))
    )
    if needs_join:
        new = joined
        for ph, val in repl.items():
            if ph in new:
                new = new.replace(ph, val)
        if strip_unfilled:
            new = _STRIP_RE.sub("", new)
            new = _EMPTY_PARENS_RE.sub("", new)
        if new != joined and paragraph.runs:
            paragraph.runs[0].text = new
            for run in paragraph.runs[1:]:
                run.text = ""
    # a token-only line (e.g. "{{country}}") that got stripped is now blank -> signal removal
    return bool(
        strip_unfilled
        and had_token
        and "".join(run.text for run in paragraph.runs).strip() == ""
    )


def _fill_container(container, repl: dict[str, str], strip_unfilled: bool) -> None:
    """Fill paragraphs + tables of a docx container (Document, _Cell, header, or footer)."""
    paragraphs = list(container.paragraphs)
    emptied = [p for p in paragraphs if _fill_paragraph(p, repl, strip_unfilled)]
    # Drop lines that an unfilled token emptied, so a token-only line doesn't become a blank line.
    # Guard: never remove ALL paragraphs (a table cell must keep >=1, and we don't want to gut a
    # paragraph that was only blank for other reasons — those have had_token False and aren't here).
    if emptied and len(emptied) < len(paragraphs):
        for paragraph in emptied:
            el = paragraph._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                _fill_container(cell, repl, strip_unfilled)


def fill_docx(docx_bytes: bytes, values: dict, *, strip_unfilled: bool = True) -> bytes:
    """Return ``docx_bytes`` with ``{{token}}`` placeholders replaced from ``values``.

    With ``strip_unfilled`` (default), any placeholder with no supplied value is removed along with
    an immediately-preceding separator, so optional/unentered tokens leave no ``{{token}}`` or
    dangling comma in the output.
    """
    repl = _build_replacements(values)
    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as exc:  # noqa: BLE001 — surface a clean 422 rather than a 500
        raise EngineError(
            422, "bad_template", "Source file is not a readable .docx."
        ) from exc

    _fill_container(doc, repl, strip_unfilled)
    for section in doc.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            try:
                _fill_container(part, repl, strip_unfilled)
            except Exception:  # noqa: BLE001 — a missing/linked header part is non-fatal
                pass

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def resolve_template_docx(
    db,
    jurisdiction_code: str,
    counterparty_type_code: str,
    mutuality_code: str,
    *,
    org_id: str | None = None,
    variant: str = "tokenised",
) -> tuple[bytes, object]:
    """Fetch the current ``variant`` template .docx bytes for the (jurisdiction, counterparty,
    mutuality) combo from the normalized schema. Raises EngineError when missing."""
    from sqlalchemy import select

    from app.models_v2 import DocumentBlob, Template, TemplateVersion

    stmt = (
        select(DocumentBlob, Template)
        .join(TemplateVersion, TemplateVersion.blob_id == DocumentBlob.id)
        .join(Template, Template.id == TemplateVersion.template_id)
        .where(
            Template.jurisdiction_code == jurisdiction_code,
            Template.counterparty_type_code == counterparty_type_code,
            Template.mutuality_code == mutuality_code,
            TemplateVersion.variant_code == variant,
            TemplateVersion.is_current.is_(True),
        )
        .limit(1)
    )
    if org_id:
        stmt = stmt.where(Template.org_id == org_id)

    row = db.execute(stmt).first()
    if row is None:
        raise EngineError(
            404,
            "template_not_found",
            f"No current '{variant}' template for "
            f"{jurisdiction_code}/{counterparty_type_code}/{mutuality_code} "
            "(check the template metadata + that the .docx blob was loaded).",
        )
    blob, template = row
    if not blob.bytes:
        raise EngineError(
            409,
            "template_blob_missing",
            "Template version found but its .docx bytes are not loaded "
            "(run the Seed NDA Templates upload).",
        )
    return blob.bytes, template
