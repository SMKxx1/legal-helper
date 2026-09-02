"""Redlined .docx export.

Builds a valid Word document representing the reviewed *incoming* NDA with true
Word **tracked changes** for every suggestion the reviewer has ACCEPTED. For an
accepted issue with suggested language we emit a ``<w:del>`` run carrying the
original incoming text and a ``<w:ins>`` run carrying the suggested replacement,
so Word shows it as a tracked edit (reject -> original, accept -> suggestion).

Tracked changes are not exposed by python-docx's high-level API, so we drop down
to raw OOXML (``docx.oxml.OxmlElement`` + ``qn``) for the ``w:del`` / ``w:ins``
wrappers. Everything else (headings, body paragraphs, the italic rationale note)
uses the friendly python-docx API.

The renderer is defensive: a malformed issue never aborts the export — each
clause is wrapped in try/except and skipped on error.
"""

from __future__ import annotations

import contextlib
import re
from collections.abc import Iterable
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument  # the class, for type annotations
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

_AUTHOR = "NDA Review"


def apply_source_page_size(document: Any, review: Any) -> None:
    """Match the DOCX section page size to the source document's first page.

    Best-effort (guarded): reads the source PDF's first-page dimensions and sets
    the section width/height so the exported DOCX keeps the original paper size
    (e.g. US Letter vs A4, portrait vs landscape). Never raises.
    """
    try:
        from pathlib import Path

        from docx.shared import Emu

        from ..ingestion.pdf_layout import extract_pdf_layout

        src = getattr(review, "display_pdf_path", None) or (
            getattr(review, "incoming_path", None)
            if (getattr(review, "incoming_format", "") or "") == "pdf"
            else None
        )
        if not src or not Path(src).exists():
            return
        layout = extract_pdf_layout(src)
        if not layout.pages:
            return
        w_pt, h_pt = layout.pages[0].width, layout.pages[0].height
        section = document.sections[0]
        section.page_width = Emu(int(w_pt * 12700))  # 1 pt = 12700 EMU
        section.page_height = Emu(int(h_pt * 12700))
    except Exception:  # pragma: no cover - never break export over page size
        pass


def _tracked_date() -> str:
    """OOXML xsd:dateTime stamped on tracked changes, computed at export time.

    (Previously captured once at import, so every export shared the
    process-start timestamp.)
    """
    return datetime.now().replace(microsecond=0).isoformat()


# --------------------------------------------------------------------------- #
# Raw-OOXML tracked-change helpers
# --------------------------------------------------------------------------- #
# Characters lxml/python-docx reject (raise on assignment): the C0 controls (minus tab/LF/CR),
# unpaired surrogates U+D800–U+DFFF, and the non-characters U+FFFE/U+FFFF. Any of these in a model
# span/suggestion must be stripped or the run fails to build and the whole clause redline is dropped.
_XML_INVALID = re.compile("[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff￾￿]")


def _xml_safe(text: str) -> str:
    """Strip XML-illegal control characters so a stray control byte in a model-generated span or
    suggestion can't make the run element fail to build and silently drop the entire redline."""
    return _XML_INVALID.sub("", text or "")


def _make_run_element(text: str, *, text_tag: str) -> Any:
    """Build a ``<w:r>`` element wrapping *text* in ``text_tag`` (w:t / w:delText).

    ``xml:space="preserve"`` keeps leading/trailing whitespace intact.
    """
    run = OxmlElement("w:r")
    # Split on UNIVERSAL newlines (\r\n, \r, \n) and emit <w:br/> between segments — a bare newline
    # in w:t renders as a single space in Word, so a multi-line suggestion/deletion would otherwise
    # run together; splitting only on \n would also leave a stray \r as a literal &#13; artifact.
    for i, segment in enumerate(re.split(r"\r\n|[\r\n]", _xml_safe(text))):
        if i:
            run.append(OxmlElement("w:br"))
        text_el = OxmlElement(text_tag)
        text_el.set(qn("xml:space"), "preserve")
        text_el.text = segment
        run.append(text_el)
    return run


def add_tracked_deletion(
    paragraph: Paragraph,
    text: str,
    author: str = _AUTHOR,
    date_str: str | None = None,
) -> Any:
    """Append a tracked *deletion* (``<w:del>``) of *text* to *paragraph*.

    Deleted text must live in ``<w:delText>`` rather than ``<w:t>`` per OOXML.
    """
    date_str = date_str or _tracked_date()
    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), str(_next_change_id()))
    del_el.set(qn("w:author"), author)
    del_el.set(qn("w:date"), date_str)
    del_el.append(_make_run_element(text, text_tag="w:delText"))
    paragraph._p.append(del_el)
    return del_el


def add_tracked_insertion(
    paragraph: Paragraph,
    text: str,
    author: str = _AUTHOR,
    date_str: str | None = None,
) -> Any:
    """Append a tracked *insertion* (``<w:ins>``) of *text* to *paragraph*."""
    date_str = date_str or _tracked_date()
    ins_el = OxmlElement("w:ins")
    ins_el.set(qn("w:id"), str(_next_change_id()))
    ins_el.set(qn("w:author"), author)
    ins_el.set(qn("w:date"), date_str)
    ins_el.append(_make_run_element(text, text_tag="w:t"))
    paragraph._p.append(ins_el)
    return ins_el


_change_counter = {"n": 0}


def _next_change_id() -> int:
    """Return a monotonically increasing revision id for w:del / w:ins."""
    _change_counter["n"] += 1
    return _change_counter["n"]


# --------------------------------------------------------------------------- #
# Small attribute helpers (tolerant of ORM objects or SimpleNamespace fakes)
# --------------------------------------------------------------------------- #
def _attr(obj: Any, name: str, default: str = "") -> str:
    value = getattr(obj, name, default)
    if value is None:
        return default
    return str(value)


def _clause_label(issue: Any) -> str:
    number = _attr(issue, "clause_number").strip()
    heading = _attr(issue, "clause_heading").strip()
    if number and heading:
        return f"{number} {heading}"
    return number or heading or "Clause"


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #
def build_redlined_docx(
    review: Any,
    issues: Iterable[Any],
    out_path: str | Path,
) -> Path:
    """Render *review* + its *issues* to a tracked-changes .docx at *out_path*.

    For issues with ``status == "accepted"`` and non-empty ``suggested_language``
    the original ``incoming_text`` is rendered as a tracked deletion followed by
    the suggested language as a tracked insertion. Other issues render their
    incoming text as a plain run. Each clause is followed by an italic note
    summarising the severity / title / rationale.

    Never raises on odd per-clause input: clause rendering is guarded and the
    write always completes. Returns ``Path(out_path)``.
    """
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    apply_source_page_size(document, review)  # preserve original paper dimensions
    tracked_date = _tracked_date()  # one consistent timestamp for this export

    # --- Title + subtitle ------------------------------------------------- #
    # The title + provider/model are untrusted (model-generated counterparty text / uploaded filename
    # stems), so they MUST go through _xml_safe like body spans — they render BEFORE the guarded
    # clause loop, so a stray control char/surrogate here would abort the WHOLE export, not skip one clause.
    title = _xml_safe(_attr(review, "title", "Untitled review") or "Untitled review")
    document.add_heading(f"{title} — NDA Review Redline", level=0)

    provider = _attr(review, "provider")
    model = _attr(review, "model")
    backend = _xml_safe(
        " / ".join(part for part in (provider, model) if part) or "deterministic"
    )
    subtitle = document.add_paragraph()
    sub_run = subtitle.add_run(
        f"Analysis backend: {backend}  •  "
        f"Generated by NDA Review on {tracked_date}. "
        "Accepted suggestions appear as Word tracked changes."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(9)
    with contextlib.suppress(Exception):
        sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # --- Clause-by-clause body -------------------------------------------- #
    issue_list = list(issues or [])
    if not issue_list:
        document.add_paragraph("No issues were flagged for this review.")

    for issue in issue_list:
        mark = len(document.paragraphs)
        try:
            _render_clause(document, issue, tracked_date)
        except Exception as exc:  # noqa: BLE001 - never let one clause abort export
            # Roll back the partial paragraphs this clause added (an orphaned heading, a dangling
            # w:del with no matching w:ins) so the document stays well-formed, then leave a skip note.
            for p in document.paragraphs[mark:]:
                p._p.getparent().remove(p._p)
            try:
                note = document.add_paragraph()
                # _xml_safe the message too: str(exc) may echo the offending model text incl. a
                # control char, which would make add_run raise and (via the except below) drop the
                # clause with NO trace — the recovery note itself must never be the failure point.
                run = note.add_run(_xml_safe(f"[skipped clause: {exc}]"))
                run.italic = True
                run.font.size = Pt(8)
            except Exception:
                # Even the error note failed; just move on.
                pass
            continue

    document.save(str(out))
    return out


def _render_clause(
    document: DocxDocument, issue: Any, date_str: str | None = None
) -> None:
    """Render one issue: heading, body (tracked or plain), rationale note."""
    # Clause heading/body/note are untrusted model text too — sanitize them like the tracked runs
    # (add_tracked_* go through _make_run_element/_xml_safe, but add_heading/add_run do NOT), else a
    # control char here raises and the per-clause guard drops the ENTIRE well-formed finding.
    document.add_heading(_xml_safe(_clause_label(issue)), level=2)

    status = _attr(issue, "status").strip().lower()
    incoming_text = _attr(issue, "incoming_text")
    suggested = _attr(issue, "suggested_language")

    body = document.add_paragraph()
    if status == "accepted" and suggested.strip():
        # Tracked replacement: delete the original, insert the suggestion.
        if incoming_text.strip():
            add_tracked_deletion(body, incoming_text, date_str=date_str)
        add_tracked_insertion(body, suggested, date_str=date_str)
    else:
        # Plain rendering of the incoming text (or a placeholder if empty).
        body.add_run(
            _xml_safe(incoming_text) if incoming_text.strip() else "(no incoming text)"
        )

    # --- Lightweight reviewer note (stand-in for a Word comment) ---------- #
    severity = _attr(issue, "severity", "low").strip() or "low"
    issue_title = _attr(issue, "title").strip()
    rationale = _attr(issue, "rationale").strip()
    note_text = f"[{severity.upper()}] {issue_title}"
    if rationale:
        note_text += f" — {rationale}"

    note = document.add_paragraph()
    note_run = note.add_run(_xml_safe(note_text))
    note_run.italic = True
    note_run.font.size = Pt(9)
    with contextlib.suppress(Exception):
        note_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
