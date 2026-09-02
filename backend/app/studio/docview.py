"""Document view extraction — the stable, addressable text model of a template .docx.

``extract_view`` renders a .docx into an ordered list of paragraph segments, each carrying:

- ``locator`` — a deterministic path to the paragraph (grammar below);
- ``text``    — the paragraph's **normalized run-concatenated text**: exactly
  ``"".join(run.text for run in paragraph.runs)``, the SAME join the generation filler's split
  pass uses (``app.support_task.generator._fill_paragraph``), so character offsets agreed with a
  view are the offsets the surgery in ``tokenize_ops`` operates on. Existing ``{{tokens}}``
  appear as-is. (Like the filler, text inside hyperlink wrappers is not part of the run list and
  is therefore not addressable — the filler cannot fill there either.)
- ``kind``    — which document part the paragraph lives in: ``body`` | ``header`` | ``footer``
  (table membership is visible in the locator itself).
- ``parts``   — the paragraph broken into ordered render PARTS (see :class:`ViewPart`): formatted
  run slices and atomic ``{{token}}`` parts. The invariant that keeps offsets sound is
  ``"".join(part.text for part in seg.parts) == seg.text`` — a token part's ``text`` is the
  verbatim ``{{...}}`` match (braces, inner spaces and all), so parts are a pure re-slicing of the
  offset basis, never a re-writing of it.
- ``style`` / ``heading`` / ``align`` — paragraph-level presentation (style name, heading level
  0=none, alignment ``""``/``center``/``right``/``justify``) for faithful rendering. Presentation
  extraction is fail-soft: it can never change ``text``/``locator``/``content_hash`` semantics.

Locator grammar (segments joined by ``/``):

    part      := "body" | "hdr:<section>:<variant>" | "ftr:<section>:<variant>"
    variant   := "default" | "first" | "even"
    table     := "tbl:<table>:<row>:<col>"        (repeated once per nesting level)
    paragraph := "p:<index>"
    locator   := part ("/" table)* "/" paragraph

e.g. ``body/p:3``, ``body/tbl:0:1:2/p:0``, ``body/tbl:0:0:0/tbl:0:1:1/p:2``, ``hdr:0:default/p:1``.

The traversal mirrors the filler exactly: body paragraphs, then all tables (recursively, row-major
by cell), then each section's header/footer variants in the filler's order. Two deliberate,
documented refinements over the filler's write-time loop (which is idempotent so never needed
them): a merged table cell is listed once, at the first (row, col) where it appears, and a
header/footer that is *linked to previous* (no definition of its own) is skipped — its content is
already listed under the section that owns it, and merely reading a linked part would create one.

**Stale-view detection**: the view embeds ``content_hash`` — a SHA-256 over the **canonicalized
(C14N) XML** of the parts the studio can touch (``word/document.xml`` + every
``word/header*.xml`` / ``word/footer*.xml``), NOT over the raw zip bytes. Canonical XML is stable
across load→save cycles and serializer differences (Word vs lxml), so the hash changes exactly
when document content/formatting changes. A locator + char-range from a view whose
``content_hash`` still matches the document addresses exactly one span in the .docx.
"""

from __future__ import annotations

import hashlib
import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from lxml import etree

from .errors import BadDocxError, LocatorNotFoundError

#: A ``{{token}}`` placeholder — the exact shape the generation filler strips/fills
#: (``app.support_task.generator._STRIP_RE`` / the envelope guard's ``_TOKEN_RE``).
TOKEN_RE = re.compile(r"\{\{\s*[A-Za-z0-9_]+\s*\}\}")
#: Same shape, capturing the bare token name.
TOKEN_NAME_RE = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")

#: (prefix, variant) -> the python-docx Section attribute, in the filler's traversal order.
_HDR_FTR_ATTRS: tuple[tuple[str, str, str], ...] = (
    ("hdr", "default", "header"),
    ("hdr", "first", "first_page_header"),
    ("hdr", "even", "even_page_header"),
    ("ftr", "default", "footer"),
    ("ftr", "first", "first_page_footer"),
    ("ftr", "even", "even_page_footer"),
)


@dataclass(frozen=True)
class ViewPart:
    """One rendered part of a paragraph: a formatted run slice, or an atomic ``{{token}}``.

    ``text`` is ALWAYS the exact characters the part occupies in the segment's offset basis
    (``ViewSegment.text``). For a token part that is the verbatim ``{{...}}`` match — so
    ``len(part.text)`` is the part's underlying length (the ``data-plen`` the client sums),
    regardless of what friendly label the view renders inside the chip.
    """

    text: str
    is_token: bool = False
    name: str = ""  # token parts only: the bare token name
    bold: bool = False
    italic: bool = False
    underline: bool = False

    @property
    def plen(self) -> int:
        """The part's length in the segment's offset basis (== ``len(text)``)."""
        return len(self.text)


@dataclass(frozen=True)
class ViewSegment:
    """One addressable paragraph of the document view."""

    locator: str
    text: str
    kind: str  # body | header | footer
    parts: tuple[ViewPart, ...] = ()
    style: str = ""  # paragraph style name (e.g. "Heading 1"), display only
    heading: int = 0  # 1..6 for heading styles, 0 otherwise
    align: str = ""  # "" | center | right | justify

    def to_dict(self) -> dict[str, str]:
        # The serialized contract stays the addressable core (locator/text/kind); parts and
        # presentation metadata are a render-layer concern consumed via the objects directly.
        return {"locator": self.locator, "text": self.text, "kind": self.kind}


@dataclass(frozen=True)
class DocumentView:
    """The full extracted view: ordered segments + the content hash it was built against."""

    content_hash: str
    segments: tuple[ViewSegment, ...]

    def find(self, locator: str) -> ViewSegment | None:
        for seg in self.segments:
            if seg.locator == locator:
                return seg
        return None

    def to_dict(self) -> dict[str, Any]:
        return {
            "content_hash": self.content_hash,
            "segments": [s.to_dict() for s in self.segments],
        }


def paragraph_text(paragraph: Any) -> str:
    """THE normalization: the filler's run-concatenated paragraph text (offsets agree with it)."""
    return "".join(run.text for run in paragraph.runs)


def load_document(docx_bytes: bytes) -> Any:
    """Open .docx bytes as a python-docx Document, or refuse with a typed error."""
    from docx import Document

    try:
        return Document(BytesIO(docx_bytes))
    except Exception as exc:  # noqa: BLE001 — surface a clean typed refusal, never a 500
        raise BadDocxError() from exc


def content_hash(docx_bytes: bytes) -> str:
    """SHA-256 over the canonicalized XML of every part the studio can touch.

    C14N (rather than raw zip-entry bytes) makes the hash stable across serializers: a Word-saved
    and an lxml-saved copy of the same content hash identically, which is what keeps the oplog's
    prior/new hash chain sound across apply → undo → redo cycles.
    """
    try:
        with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
            names = sorted(
                name
                for name in zf.namelist()
                if name == "word/document.xml"
                or (
                    name.startswith(("word/header", "word/footer"))
                    and name.endswith(".xml")
                )
            )
            if "word/document.xml" not in names:
                raise BadDocxError("Zip archive has no word/document.xml part.")
            digest = hashlib.sha256()
            for name in names:
                digest.update(name.encode("utf-8"))
                digest.update(b"\0")
                digest.update(
                    etree.tostring(etree.fromstring(zf.read(name)), method="c14n")
                )
                digest.update(b"\0")
            return digest.hexdigest()
    except BadDocxError:
        raise
    except Exception as exc:  # noqa: BLE001 — bad zip / bad xml => typed refusal
        raise BadDocxError() from exc


# --------------------------------------------------------------------------- #
# Extraction
# --------------------------------------------------------------------------- #
_HEADING_STYLE_RE = re.compile(r"^heading\s+(\d+)$", re.IGNORECASE)


def _run_flags(run: Any) -> tuple[bool, bool, bool]:
    """(bold, italic, underline) direct formatting of a run — fail-soft, display only."""
    try:
        return bool(run.bold), bool(run.italic), bool(run.underline)
    except Exception:  # noqa: BLE001 — presentation can never break extraction
        return False, False, False


def _para_presentation(paragraph: Any) -> tuple[str, int, str]:
    """(style_name, heading_level, alignment) of a paragraph — fail-soft, display only."""
    style = ""
    try:
        style = paragraph.style.name or ""
    except Exception:  # noqa: BLE001
        style = ""
    heading = 0
    m = _HEADING_STYLE_RE.match(style.strip())
    if m:
        heading = min(max(int(m.group(1)), 1), 6)
    elif style.strip().lower() == "title":
        heading = 1
    align = ""
    try:
        al = paragraph.alignment
        name = getattr(al, "name", "").lower() if al is not None else ""
        if name in ("center", "right", "justify"):
            align = name
    except Exception:  # noqa: BLE001
        align = ""
    return style, heading, align


def _extract_parts(paragraph: Any) -> tuple[ViewPart, ...]:
    """Slice a paragraph into ordered :class:`ViewPart`\\ s.

    ``{{token}}`` matches (which may straddle runs) become single token parts carrying the
    verbatim matched text; everything else becomes run parts split at run boundaries (adjacent
    parts with identical formatting merge). The invariant: the concatenation of every part's
    ``text`` is EXACTLY :func:`paragraph_text` — parts re-slice the offset basis, never alter it.
    """
    text = paragraph_text(paragraph)
    if not text:
        return ()
    run_spans: list[tuple[int, int, tuple[bool, bool, bool]]] = []
    pos = 0
    for run in paragraph.runs:
        length = len(run.text)
        if length:
            run_spans.append((pos, pos + length, _run_flags(run)))
            pos += length

    parts: list[ViewPart] = []

    def emit_runs(a: int, b: int) -> None:
        for r_start, r_end, flags in run_spans:
            lo, hi = max(a, r_start), min(b, r_end)
            if lo >= hi:
                continue
            chunk = text[lo:hi]
            prev = parts[-1] if parts else None
            if (
                prev is not None
                and not prev.is_token
                and (prev.bold, prev.italic, prev.underline) == flags
            ):
                parts[-1] = ViewPart(
                    text=prev.text + chunk,
                    bold=flags[0],
                    italic=flags[1],
                    underline=flags[2],
                )
            else:
                parts.append(
                    ViewPart(
                        text=chunk, bold=flags[0], italic=flags[1], underline=flags[2]
                    )
                )

    pos = 0
    for match in TOKEN_RE.finditer(text):
        if match.start() > pos:
            emit_runs(pos, match.start())
        name_m = TOKEN_NAME_RE.match(match.group(0))
        parts.append(
            ViewPart(
                text=match.group(0),
                is_token=True,
                name=name_m.group(1) if name_m else match.group(0),
            )
        )
        pos = match.end()
    if pos < len(text):
        emit_runs(pos, len(text))
    return tuple(parts)


def _walk_container(
    container: Any, prefix: str, kind: str, out: list[ViewSegment]
) -> None:
    """Paragraphs then tables (recursively) — the filler's ``_fill_container`` traversal."""
    for p_idx, paragraph in enumerate(container.paragraphs):
        style, heading, align = _para_presentation(paragraph)
        out.append(
            ViewSegment(
                locator=f"{prefix}/p:{p_idx}",
                text=paragraph_text(paragraph),
                kind=kind,
                parts=_extract_parts(paragraph),
                style=style,
                heading=heading,
                align=align,
            )
        )
    for t_idx, table in enumerate(container.tables):
        seen_tc: list[Any] = []  # keeps lxml proxies alive so identity stays stable
        seen_ids: set[int] = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                tc = cell._tc
                if id(tc) in seen_ids:  # merged cell already listed at its first slot
                    continue
                seen_tc.append(tc)
                seen_ids.add(id(tc))
                _walk_container(
                    cell, f"{prefix}/tbl:{t_idx}:{r_idx}:{c_idx}", kind, out
                )


def extract_view(docx_bytes: bytes) -> DocumentView:
    """Extract the stable, addressable document view (see module docstring for the contract)."""
    doc = load_document(docx_bytes)
    segments: list[ViewSegment] = []
    _walk_container(doc, "body", "body", segments)
    for s_idx, section in enumerate(doc.sections):
        for prefix, variant, attr in _HDR_FTR_ATTRS:
            kind = "header" if prefix == "hdr" else "footer"
            try:
                part = getattr(section, attr)
                if part.is_linked_to_previous:
                    continue  # no definition of its own; reading one would CREATE it
                _walk_container(part, f"{prefix}:{s_idx}:{variant}", kind, segments)
            except Exception:  # noqa: BLE001 — a missing/broken part is non-fatal (filler parity)
                continue
    return DocumentView(content_hash=content_hash(docx_bytes), segments=tuple(segments))


# --------------------------------------------------------------------------- #
# Resolution
# --------------------------------------------------------------------------- #
def resolve_locator(doc: Any, locator: str) -> Any:
    """Resolve a locator to its python-docx Paragraph, touching ONLY the addressed part.

    (Short-circuiting matters: python-docx creates a header/footer part as a side effect of
    reading a linked one, so body edits must never traverse headers.)

    Raises :class:`LocatorNotFoundError` for anything that does not address exactly one existing
    paragraph of ``doc``.
    """
    segs = locator.split("/")
    if len(segs) < 2:
        raise LocatorNotFoundError(locator, "too few segments")
    try:
        container = _resolve_part(doc, segs[0], locator)
        for seg in segs[1:-1]:
            tag, *idx = seg.split(":")
            if tag != "tbl" or len(idx) != 3:
                raise LocatorNotFoundError(locator, f"bad table segment {seg!r}")
            t, r, c = (_index(i, locator) for i in idx)
            container = container.tables[t].rows[r].cells[c]
        tag, *idx = segs[-1].split(":")
        if tag != "p" or len(idx) != 1:
            raise LocatorNotFoundError(locator, f"bad paragraph segment {segs[-1]!r}")
        return container.paragraphs[_index(idx[0], locator)]
    except LocatorNotFoundError:
        raise
    except IndexError as exc:
        raise LocatorNotFoundError(locator, "index out of range") from exc


def _index(raw: str, locator: str) -> int:
    try:
        value = int(raw)
    except ValueError as exc:
        raise LocatorNotFoundError(locator, f"non-numeric index {raw!r}") from exc
    if value < 0:  # negative indexing would silently address the wrong element
        raise LocatorNotFoundError(locator, f"negative index {raw!r}")
    return value


def _resolve_part(doc: Any, part_seg: str, locator: str) -> Any:
    if part_seg == "body":
        return doc
    tag, *rest = part_seg.split(":")
    if tag not in ("hdr", "ftr") or len(rest) != 2:
        raise LocatorNotFoundError(locator, f"bad part segment {part_seg!r}")
    s_idx = _index(rest[0], locator)
    attr = next((a for p, v, a in _HDR_FTR_ATTRS if p == tag and v == rest[1]), None)
    if attr is None:
        raise LocatorNotFoundError(locator, f"bad header/footer variant {rest[1]!r}")
    try:
        part = getattr(doc.sections[s_idx], attr)
    except IndexError as exc:
        raise LocatorNotFoundError(locator, "no such section") from exc
    if part.is_linked_to_previous:
        # The view never emits locators for linked parts; resolving one would create a part.
        raise LocatorNotFoundError(locator, "header/footer is linked to previous")
    return part
