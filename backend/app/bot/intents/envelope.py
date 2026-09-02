"""The ``envelope`` intent + its DocuSign interactivity chains (PLAN §3.9, reference §3.5/§3.7).

A behavioral port of the envelope half of the n8n ``NDA: Envelope Review`` sub-workflow and the
``send_docusign`` / ``env_use_doc`` / ``decline_doc`` / ``nda_docusign`` branches of ``NDA: Interactivity``
— with the DELIBERATE behavior change PLAN §2 #1 makes explicit: **the direct ≥2-signer path now gains
a human confirm card before the DocuSign send** (today it sends immediately), closing the spoofed-email
→ outbound-envelope hole (PLAN §6). The modal is a COLLECTOR, not the guard; the Confirm & send click is.

Three entry points (PLAN §3.9), all decided from the already-hardened
:class:`~app.bot.router.Classification` + the inbound :class:`~app.bot.envelope.Envelope`:

* **(a) ≥2 signer_emails + an attached doc** — fetch the doc, run the **unfilled-``{{token}}`` guard**
  (a tokenised leftover → the ported refusal), then post the NEW confirm card summarising
  signers/routing/CC with *Confirm & send* / *Cancel*. The actual DocuSign send happens on the click.
* **(b) <2 signers** — confirm the document + an *Enter signing details* button (the ported
  ``send_docusign`` contract) → the ``nda_docusign`` MODAL (Amperesand signer, counterparty signer,
  signing order, CC emails, CC timing) → on ``view_submission`` build the SAME confirm card path.
* **(c) no attached doc** — thread-doc recovery (:mod:`app.bot.thread_docs`; Slack scan, else the ported
  email "attach the document" ask) → *Yes, use it* (``env_use_doc``) / *No, attach a file*
  (``decline_doc``) → back into (a) or (b) for the recovered document.

State across every chain lives in a durable ``bot_correlation`` row (PLAN §3.9): the Slack button VALUES
carry only ``{v:1, kind, ref}`` (well under Slack's 2000-char limit — "typed payloads carry only keys"),
and the ``ref`` keys the row that holds the document bytes + routing + the requester mapping.

On a confirmed send: :meth:`app.integrations.docusign.DocuSignClient.create_and_send_envelope` +
:func:`app.integrations.models.save_envelope_attempt` (recording ``requested_by`` + channel/thread from
the envelope — the PLAN §3.10 requester mapping the P4 watcher DMs from), the ported success/failure
replies, and a friendly "e-signature isn't set up" when the DOCUSIGN capability is disabled (capabilities
fail soft, PLAN §6). Failed sends persist too (``status='failed'``) so the audit trail is honest.

Every collaborator is injected (a fake DocuSign sender, a fake Slack fetcher/scanner, a fake modal
opener), so the whole matrix runs with **zero network** (PLAN house rules). The interactivity handlers are
registered onto the shared :class:`~app.bot.interactivity.InteractivityRegistry` via
:func:`register_envelope` (called from ``default_interactivity_registry``) — no change to the dispatcher.
"""

from __future__ import annotations

import base64
import re
from collections.abc import Callable
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from io import BytesIO
from typing import TYPE_CHECKING, Any, Literal

from pydantic import Field
from sqlalchemy import select

from ...telemetry import get_logger
from ..blockkit import (
    ACTION_DECLINE_DOC,
    ACTION_ENV_CONFIRM_CANCEL,
    ACTION_ENV_CONFIRM_SEND,
    ACTION_ENV_USE_DOC,
    ACTION_SEND_DOCUSIGN,
    ENVELOPE_CONFIRM_FALLBACK_TEXT,
    KIND_DECLINE_DOC,
    KIND_ENV_CONFIRM,
    KIND_ENV_USE_DOC,
    KIND_NDA_DOCUSIGN_MODAL,
    KIND_SEND_DOCUSIGN,
    MODAL_ACTION_AMP,
    MODAL_ACTION_CC,
    MODAL_ACTION_CC_SEQ,
    MODAL_ACTION_CP,
    MODAL_ACTION_SEQ,
    MODAL_BLOCK_AMP,
    MODAL_BLOCK_CC,
    MODAL_BLOCK_CC_SEQ,
    MODAL_BLOCK_CP,
    MODAL_BLOCK_SEQ,
    MODAL_CALLBACK_ID,
    SEND_DOCUSIGN_FALLBACK_TEXT,
    THREAD_DOC_FALLBACK_TEXT,
    docusign_modal_view,
    envelope_confirm_blocks,
    signer_details_button_blocks,
    thread_doc_confirm_blocks,
)
from ..channels.protocol import Reply
from ..envelope import AttachmentRef, Envelope
from ..interactivity import (
    ButtonPayload,
    Interaction,
    InteractivityDeps,
    InteractivityRegistry,
)
from ..models import BotCorrelation
from ..thread_docs import ThreadDoc, ThreadScanner
from . import IntentContext, IntentReply
from .review import HttpSlackFileFetcher, SlackFileFetcher

if TYPE_CHECKING:
    from app.capabilities import CapabilityRegistry
    from app.config import Settings
    from app.integrations.docusign import EnvelopeResult

log = get_logger("nda.bot.intent.envelope")

# --------------------------------------------------------------------------- #
# Ported copy (reference §3.5/§3.7) + constants
# --------------------------------------------------------------------------- #
DEFAULT_DOC_NAME = "NDA.docx"
#: Doc-like suffixes used to PICK the envelope document off the inbound attachments.
_DOC_SUFFIXES = frozenset({".docx", ".doc", ".pdf"})

#: bot_correlation kind for the envelope confirm/collect state (PLAN §3.9). Rows are LOADED by their
#: unique ``key`` (kind-agnostic) so the generate flow can seed one under its own kind and hand us the
#: ref via a ``send_docusign`` button (task deliverable 3).
ENV_CORRELATION_KIND = "env_confirm"
#: How long a pending confirm/collect state lives before the worker sweep reaps it.
ENV_CORRELATION_TTL_HOURS = 24.0

#: Email path (no interactive buttons): the ported "attach the document" ask (reference §3.5).
EMAIL_NO_DOC_TEXT = (
    "To send an NDA for signature, attach the finished .docx and send your request "
    "from Slack — the confirm-before-send step needs Slack's buttons."
)
#: Email path with a doc: e-signature sends are interactive → point the user at Slack.
EMAIL_HAS_DOC_TEXT = (
    "Thanks — to send this NDA for signature I need the confirm-before-send step, which "
    "lives in Slack. Please start the signature request there and I'll walk you through it."
)
#: Slack, no attachment and no recoverable thread doc (reference §3.5 "No-Doc Reject").
SLACK_NO_DOC_TEXT = (
    "I couldn't find a document to send. Attach the finished NDA (`.docx`) and send your "
    "request again."
)
#: The attachment couldn't be fetched / re-fetched (Slack download failed, spool gone).
DOWNLOAD_FAILED_TEXT = "I couldn't open that document. Please re-attach the finished NDA (`.docx`) and try again."
#: A recovered/stored document is no longer reachable at send time.
DOC_GONE_TEXT = "I couldn't retrieve that document any more. Please re-attach the NDA and try again."
#: The confirm/collect state expired or was already consumed (stale button).
EXPIRED_STATE_TEXT = (
    "Sorry — that signature request has expired. Please start again by mentioning me "
    "with your envelope request."
)
#: A second Confirm click after a successful send (idempotent — no second envelope).
ALREADY_SENT_TEXT = (
    "That NDA has already been sent to DocuSign — I won't send it again."
)
#: The Cancel button (reference: confirm card cancel).
CANCEL_TEXT = "No problem — I won't send that NDA for signature."
#: The ported decline reply (reference §3.7 "Decline Reply").
DECLINE_TEXT = "Ok — attach the file you want to use and send your request again."
#: The modal couldn't be opened (no Slack client wired / views.open failed).
MODAL_UNAVAILABLE_TEXT = "Sorry — I couldn't open the signing-details form just now. Please try again in a moment."
#: Modal submit with unreadable signer emails (post-ack, so we can't return inline modal errors).
MODAL_INVALID_TEXT = (
    "I couldn't read valid signer email addresses from that form. Mention me and start the "
    "envelope request again."
)
#: DOCUSIGN capability disabled/unhealthy — capabilities fail soft (PLAN §6).
DOCUSIGN_UNAVAILABLE_TEXT = (
    "e-signature isn't set up right now, so I can't send this NDA to DocuSign. "
    "Let the team know and try again later."
)
#: Ported send outcomes (reference §3.7 "Env Confirm Reply" / "Modal Confirm Reply").
SEND_SUCCESS_TEXT = ":incoming_envelope: Sent to DocuSign for signature. I'll let you know when it's completed."


def _send_failure_text(detail: str) -> str:
    tail = f" ({detail})" if detail else ""
    return (
        ":warning: DocuSign didn't accept that envelope"
        f"{tail}. Please check the signing details and try again."
    )


def _token_refusal_text(tokens: list[str]) -> str:
    """Ported "Build Template Reject" (reference §3.5): name the leftover placeholders, suggest generate."""
    shown = ", ".join(tokens[:6]) + ("…" if len(tokens) > 6 else "")
    return (
        "That document still has unfilled placeholders "
        f"({shown}) — it looks like a template, not a finished NDA. Run *generate* to fill "
        "it in first, then send the completed document for signature."
    )


# --------------------------------------------------------------------------- #
# Unfilled-{{token}} guard (port of "Detect Unfilled Tokens (env)", reference §3.5)
# --------------------------------------------------------------------------- #
_TOKEN_RE = re.compile(r"\{\{\s*[A-Za-z0-9_]+\s*\}\}")


def _container_text(container: Any) -> list[str]:
    """All paragraph text of a docx container (Document / _Cell), recursing into tables — the SAME
    traversal the filler uses (``app.support_task.generator._fill_container``), in read-only reverse."""
    parts = [p.text for p in container.paragraphs]
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(_container_text(cell))
    return parts


def scan_docx_tokens(docx_bytes: bytes) -> list[str]:
    """Return the unique ``{{token}}`` placeholders still present in a .docx (body + tables +
    headers/footers), preserving first-seen order (reference §3.5 unfilled-token guard).

    A non-.docx (a PDF/DOC that python-docx can't open) has no ``{{}}`` placeholders by construction, so
    it scans CLEAN (``[]``) — the guard only ever refuses a genuinely tokenised .docx. ``paragraph.text``
    joins a placeholder split across runs, so split placeholders are caught too.
    """
    from docx import Document

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception:  # noqa: BLE001 — not a readable .docx => nothing to guard against
        return []
    parts = _container_text(doc)
    for section in doc.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            try:
                parts.extend(_container_text(part))
            except Exception:  # noqa: BLE001 — a missing/linked header part is non-fatal
                pass
    seen: list[str] = []
    for match in _TOKEN_RE.findall("\n".join(parts)):
        if match not in seen:
            seen.append(match)
    return seen


# --------------------------------------------------------------------------- #
# Small pure helpers
# --------------------------------------------------------------------------- #
_EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
_CC_SPLIT_RE = re.compile(r"[\s,;]+")

_ROUTING_LABELS = {
    "all_at_once": "everyone at once",
    "amp_first": "Amperesand signs first",
    "cp_first": "counterparty signs first",
}
_VALID_ROUTINGS = frozenset(_ROUTING_LABELS)
_VALID_CC_TIMINGS = frozenset({"before", "after"})


def _clean_emails(values: Any) -> list[str]:
    """Strip + de-dupe email-shaped strings (drops blanks / non-``@`` junk), preserving order."""
    out: list[str] = []
    for value in values or []:
        email = str(value or "").strip()
        if email and "@" in email and email not in out:
            out.append(email)
    return out


def _valid_email(value: str) -> bool:
    return bool(_EMAIL_RE.match(value.strip()))


def _parse_cc_emails(raw: str) -> list[str]:
    """Split the modal CC field on newline/comma/semicolon/space, keep ``@`` tokens (reference §3.7)."""
    out: list[str] = []
    for token in _CC_SPLIT_RE.split(raw or ""):
        token = token.strip()
        if token and "@" in token and token not in out:
            out.append(token)
    return out


def _routing_from_classification(cls: Any) -> str:
    """The direct-path routing (reference §2.9): sequential → per-signer ``i+1`` (== ``amp_first``);
    otherwise parallel (``all_at_once``). The 3-way order is only collected in the modal."""
    return "amp_first" if getattr(cls, "sequential", False) else "all_at_once"


def _normalize_routing(value: str) -> str:
    return value if value in _VALID_ROUTINGS else "all_at_once"


def _normalize_cc_timing(value: str) -> str:
    return value if value in _VALID_CC_TIMINGS else "after"


def build_summary(
    signers: list[str],
    routing: str,
    cc: list[str],
    cc_timing: str,
    file_name: str,
) -> str:
    """The confirm card's mrkdwn summary: document, ordered signers (Amperesand = index 0), CC + timing."""
    lines = [f"*Document:* {file_name or DEFAULT_DOC_NAME}"]
    lines.append(
        f"*Signers* (order: {_ROUTING_LABELS.get(routing, 'everyone at once')}):"
    )
    for i, email in enumerate(signers):
        who = "Amperesand" if i == 0 else "counterparty"
        lines.append(f" • {email} — {who}")
    if cc:
        when = "before signing" if cc_timing == "before" else "after signing"
        lines.append(f"*CC* ({when}): " + ", ".join(cc))
    return "\n".join(lines)


def _pick_doc_attachment(
    attachments: tuple[AttachmentRef, ...],
) -> AttachmentRef | None:
    """The first ``.docx`` / ``.doc`` / ``.pdf`` attachment (reference §3.5 "Attach File (env)"), else
    the first of any kind, else ``None``."""
    if not attachments:
        return None
    for att in attachments:
        name = (att.filename or "").lower()
        if any(name.endswith(suffix) for suffix in _DOC_SUFFIXES):
            return att
    return attachments[0]


def _requester(env: Envelope) -> str:
    """The verified principal for the requester mapping (PLAN §3.10): Slack user id / email address."""
    return env.sender_id or env.sender_address or ""


def _dig(data: Any, *keys: str) -> Any:
    cur = data
    for key in keys:
        if not isinstance(cur, dict):
            return None
        cur = cur.get(key)
    return cur


# --------------------------------------------------------------------------- #
# bot_correlation store (the durable state behind the ref keys)
# --------------------------------------------------------------------------- #
def _now() -> datetime:
    return datetime.now(UTC)


def _store_correlation(session_factory: Any, payload: dict[str, Any]) -> str:
    """Persist envelope state, returning a fresh opaque ``ref`` key (the button value's only secret)."""
    from ...auth.security import new_token

    key = new_token(18)
    with session_factory() as session:
        session.add(
            BotCorrelation(
                key=key,
                kind=ENV_CORRELATION_KIND,
                payload_json=payload,
                expires_at=_now() + timedelta(hours=ENV_CORRELATION_TTL_HOURS),
            )
        )
        session.commit()
    return key


def stash_generated_document(
    session_factory: Any, origin: Envelope, docx_bytes: bytes, file_name: str = ""
) -> str:
    """Public seam for the GENERATE flow (PLAN §3.9 / generate_completion): stash a freshly generated
    .docx as envelope state and return the ``ref`` for :func:`blockkit.send_docusign_value`, so the
    "Send via DocuSign? Yes" button enters the exact same modal → confirm → send chain as an
    attached document. Signers/cc start empty (the modal collects them); routing defaults ported."""
    import base64 as _b64

    payload = _base_payload(
        origin,
        file_name=file_name or DEFAULT_DOC_NAME,
        doc_b64=_b64.b64encode(docx_bytes).decode("ascii"),
        signers=[],
        cc=[],
        routing="all_at_once",
        cc_timing="after",
    )
    return _store_correlation(session_factory, payload)


def _load_correlation(session_factory: Any, ref: str) -> dict[str, Any] | None:
    """Read envelope state by its unique ``ref`` key (kind-agnostic), or ``None`` if absent/expired."""
    if not ref or session_factory is None:
        return None
    try:
        with session_factory() as session:
            row = session.execute(
                select(BotCorrelation).where(BotCorrelation.key == ref)
            ).scalar_one_or_none()
            if row is None:
                return None
            if _expired(row.expires_at):
                return None
            return dict(row.payload_json or {})
    except Exception as exc:  # noqa: BLE001 — a state read must never crash the ack path
        log.warning("bot.envelope.correlation_read_failed", ref=ref, error=repr(exc))
        return None


def _update_correlation(
    session_factory: Any, ref: str, updates: dict[str, Any]
) -> None:
    """Merge ``updates`` into the stored payload (a NEW dict so SQLAlchemy tracks the JSON change)."""
    if not ref or session_factory is None:
        return
    try:
        with session_factory() as session:
            row = session.execute(
                select(BotCorrelation).where(BotCorrelation.key == ref)
            ).scalar_one_or_none()
            if row is None:
                return
            row.payload_json = {**(row.payload_json or {}), **updates}
            session.commit()
    except Exception as exc:  # noqa: BLE001
        log.warning("bot.envelope.correlation_update_failed", ref=ref, error=repr(exc))


def _expired(expires_at: datetime | None) -> bool:
    if expires_at is None:
        return False
    exp = expires_at if expires_at.tzinfo else expires_at.replace(tzinfo=UTC)
    return exp < _now()


def _base_payload(
    env: Envelope,
    *,
    file_name: str,
    doc_b64: str | None = None,
    slack_file_id: str = "",
    file_url: str = "",
    signers: list[str],
    cc: list[str],
    routing: str,
    cc_timing: str,
) -> dict[str, Any]:
    """The envelope-state row: the document (inline b64 OR a Slack file ref) + routing + the requester
    mapping (channel/thread/requested_by — PLAN §3.10)."""
    payload: dict[str, Any] = {
        "file_name": file_name or DEFAULT_DOC_NAME,
        "channel": env.channel,
        "slack_channel": env.slack_channel,
        "slack_thread_ts": env.slack_thread_ts,
        "email_message_id": env.email_message_id,
        "requested_by": _requester(env),
        "signer_emails": signers,
        "cc_emails": cc,
        "routing": routing,
        "cc_timing": cc_timing,
    }
    if doc_b64 is not None:
        payload["doc_b64"] = doc_b64
    if slack_file_id:
        payload["slack_file_id"] = slack_file_id
    if file_url:
        payload["file_url"] = file_url
    return payload


def _doc_bytes_from_payload(
    payload: dict[str, Any], *, slack_fetch: SlackFileFetcher
) -> bytes:
    """Resolve the stored document to bytes: inline base64 (the common case), else a lazy Slack fetch of
    the stored file ref (the thread-doc recovery path — bytes are pulled only on the ``env_use_doc``
    click)."""
    b64 = payload.get("doc_b64")
    if b64:
        return base64.b64decode(b64)
    ref = str(payload.get("slack_file_id") or payload.get("file_url") or "")
    if ref:
        att = AttachmentRef(
            filename=str(payload.get("file_name") or DEFAULT_DOC_NAME), source_ref=ref
        )
        return slack_fetch(att)
    raise ValueError("no document in correlation state")


# =========================================================================== #
# The envelope INTENT handler (IntentRegistry) — the three entry points
# =========================================================================== #
#: ``extract bytes for a Slack attachment`` — reused from the review intent (injectable, no network).
TextExtractor = Callable[[str | None, bytes], str]


class EnvelopeIntent:
    """The ``envelope`` intent handler (reference §3.5). Callable ``(ctx) -> IntentReply``.

    Runs AFTER the pipeline's fail-closed allowlist/approvals gate (envelope is a gated intent). Fetches
    the document (attachment → thread recovery), applies the unfilled-token guard, stashes durable state
    in ``bot_correlation``, and returns the Slack confirm card / signer button / thread-doc confirm — or
    the ported text ask on email / no-doc. Every collaborator is an injected constructor dep (session
    factory, Slack fetch, thread scanner) so the whole matrix is tested with fakes + zero network.
    """

    def __init__(
        self,
        *,
        session_factory: Any | None = None,
        settings: Settings | None = None,
        slack_fetch: SlackFileFetcher | None = None,
        thread_scan: ThreadScanner | None = None,
    ) -> None:
        self._session_factory = session_factory
        self._settings = settings
        self._slack_fetch = slack_fetch
        self._thread_scan = thread_scan

    # -- lazy production defaults ------------------------------------------
    def _get_settings(self) -> Settings:
        if self._settings is not None:
            return self._settings
        from app.config import get_settings

        return get_settings()

    def _get_session_factory(self) -> Any:
        if self._session_factory is not None:
            return self._session_factory
        from app.db import SessionLocal

        return SessionLocal

    def _get_slack_fetch(self) -> SlackFileFetcher:
        if self._slack_fetch is not None:
            return self._slack_fetch
        return HttpSlackFileFetcher(self._get_settings().slack_bot_token)

    def _get_thread_scan(self) -> ThreadScanner:
        if self._thread_scan is not None:
            return self._thread_scan
        from ..thread_docs import HttpSlackThreadScanner

        return HttpSlackThreadScanner(self._get_settings().slack_bot_token)

    # -- entry point -------------------------------------------------------
    def __call__(self, ctx: IntentContext) -> IntentReply:
        env = ctx.envelope
        cls = ctx.classification
        att = _pick_doc_attachment(env.attachments)

        if env.channel != "slack":
            # Email can't click a confirm card; the human-confirm invariant (PLAN §6) lives in Slack.
            log.info(
                "bot.envelope.email_path",
                event_key=env.event_key,
                has_doc=att is not None,
            )
            return IntentReply(
                text=EMAIL_HAS_DOC_TEXT if att is not None else EMAIL_NO_DOC_TEXT
            )

        if att is not None:
            try:
                data = self._load_bytes(env, att)
            except Exception as exc:  # noqa: BLE001 — a fetch failure is a friendly reply, not a crash
                log.warning(
                    "bot.envelope.fetch_failed",
                    event_key=env.event_key,
                    filename=att.filename,
                    error=repr(exc),
                )
                return IntentReply(text=DOWNLOAD_FAILED_TEXT)
            return self._with_document(env, cls, data, att.filename or DEFAULT_DOC_NAME)

        # (c) no attachment → thread-doc recovery.
        thread_doc = self._recover_thread_doc(env)
        if thread_doc is None:
            log.info("bot.envelope.no_doc", event_key=env.event_key)
            return IntentReply(text=SLACK_NO_DOC_TEXT)
        return self._offer_thread_doc(env, cls, thread_doc)

    # -- (a)/(b): we have the document bytes -------------------------------
    def _with_document(
        self, env: Envelope, cls: Any, data: bytes, file_name: str
    ) -> IntentReply:
        # Unfilled-token guard FIRST (both paths — a tokenised doc never reaches DocuSign).
        leftovers = scan_docx_tokens(data)
        if leftovers:
            log.info(
                "bot.envelope.tokenised_refused",
                event_key=env.event_key,
                tokens=len(leftovers),
            )
            return IntentReply(text=_token_refusal_text(leftovers))

        signers = _clean_emails(cls.signer_emails)
        cc = _clean_emails(cls.cc_emails)
        routing = _routing_from_classification(cls)
        cc_timing = _normalize_cc_timing(cls.cc_timing or "after")
        doc_b64 = base64.b64encode(data).decode("ascii")
        payload = _base_payload(
            env,
            file_name=file_name,
            doc_b64=doc_b64,
            signers=signers,
            cc=cc,
            routing=routing,
            cc_timing=cc_timing,
        )
        ref = _store_correlation(self._get_session_factory(), payload)

        if len(signers) >= 2:
            # (a) confirm card — the NEW human confirm before send (PLAN §2 #1).
            log.info(
                "bot.envelope.confirm_card",
                event_key=env.event_key,
                ref=ref,
                signers=len(signers),
            )
            summary = build_summary(signers, routing, cc, cc_timing, file_name)
            return IntentReply(
                slack_blocks=tuple(envelope_confirm_blocks(summary, ref)),
                fallback_text=ENVELOPE_CONFIRM_FALLBACK_TEXT,
            )
        # (b) collect signer details via the modal.
        log.info("bot.envelope.signer_button", event_key=env.event_key, ref=ref)
        return IntentReply(
            slack_blocks=tuple(signer_details_button_blocks(file_name, ref)),
            fallback_text=SEND_DOCUSIGN_FALLBACK_TEXT,
        )

    # -- (c): no attachment → recover from the thread ----------------------
    def _recover_thread_doc(self, env: Envelope) -> ThreadDoc | None:
        if not env.slack_thread_ts:
            return None
        try:
            return self._get_thread_scan()(env.slack_channel, env.slack_thread_ts)
        except Exception as exc:  # noqa: BLE001 — recovery is best-effort
            log.warning(
                "bot.envelope.thread_scan_failed",
                event_key=env.event_key,
                error=repr(exc),
            )
            return None

    def _offer_thread_doc(
        self, env: Envelope, cls: Any, thread_doc: ThreadDoc
    ) -> IntentReply:
        payload = _base_payload(
            env,
            file_name=thread_doc.file_name,
            slack_file_id=thread_doc.file_id,
            file_url=thread_doc.file_url,
            signers=_clean_emails(cls.signer_emails),
            cc=_clean_emails(cls.cc_emails),
            routing=_routing_from_classification(cls),
            cc_timing=_normalize_cc_timing(cls.cc_timing or "after"),
        )
        ref = _store_correlation(self._get_session_factory(), payload)
        log.info(
            "bot.envelope.thread_doc_offer",
            event_key=env.event_key,
            ref=ref,
            file_name=thread_doc.file_name,
        )
        return IntentReply(
            slack_blocks=tuple(thread_doc_confirm_blocks(thread_doc.file_name, ref)),
            fallback_text=THREAD_DOC_FALLBACK_TEXT,
        )

    def _load_bytes(self, env: Envelope, att: AttachmentRef) -> bytes:
        """Resolve attachment bytes from its channel handle (Slack file) — mirrors the review intent."""
        return self._get_slack_fetch()(att)


# =========================================================================== #
# Typed button-value payloads (validated by the dispatcher before a handler runs)
# =========================================================================== #
class SendDocusignPayload(ButtonPayload):
    """The ``send_docusign`` button (reference §3.5 / task deliverable 3): open the signing modal.

    ``ref`` keys the ``bot_correlation`` row that holds the document (this intent's <2-signer path, OR a
    generated .docx the generate flow stashed before emitting this exact button value)."""

    kind: Literal["send_docusign"] = "send_docusign"
    ref: str = Field(min_length=1)


class EnvConfirmPayload(ButtonPayload):
    """The confirm card's *Confirm & send* / *Cancel* buttons. ``action`` (``send``/``cancel``) is the
    producer's authoritative decision; the ``action_id`` is the fallback."""

    kind: Literal["env_confirm"] = "env_confirm"
    ref: str = Field(min_length=1)
    action: str = ""


class EnvUseDocPayload(ButtonPayload):
    """The *Yes, use it* thread-doc confirm (reference §3.5 ``env_use_doc``)."""

    kind: Literal["env_use_doc"] = "env_use_doc"
    ref: str = Field(min_length=1)


class DeclineDocPayload(ButtonPayload):
    """The *No, attach a file* decline (reference §3.7 ``decline_doc``)."""

    kind: Literal["decline_doc"] = "decline_doc"
    ref: str = ""


# =========================================================================== #
# Envelope interactivity handlers + registration
# =========================================================================== #
#: A DocuSign sender: keyword-only ``create_and_send_envelope`` shape → :class:`EnvelopeResult`.
EnvelopeSender = Callable[..., "EnvelopeResult"]
#: Open a Slack modal: ``(trigger_id, view) -> None`` (``views.open``). Injectable for tests.
OpenView = Callable[[str, dict], None]


@dataclass(frozen=True)
class EnvelopeDeps:
    """Envelope-specific collaborators, bound at :func:`register_envelope` time (the default registry
    binds NONE and each is resolved lazily from the per-dispatch :class:`InteractivityDeps` + settings).

    Tests bind fakes here (a capturing modal opener, a fake DocuSign sender, a fake Slack fetcher) and
    pass an ordinary ``InteractivityDeps`` (session factory + reply service) at dispatch — so the whole
    chain runs with zero network.
    """

    sender: EnvelopeSender | None = None
    slack_fetch: SlackFileFetcher | None = None
    open_view: OpenView | None = None
    docusign_registry: CapabilityRegistry | None = None


@dataclass
class _ModalInput:
    amp_email: str
    cp_email: str
    order: str
    cc_emails: list[str]
    cc_timing: str


class EnvelopeInteractivity:
    """The ``send_docusign`` / ``nda_docusign`` / ``env_confirm`` / ``env_use_doc`` / ``decline_doc``
    kind handlers (reference §3.7). Registered onto the shared :class:`InteractivityRegistry`.
    """

    def __init__(self, deps: EnvelopeDeps | None = None) -> None:
        self._deps = deps or EnvelopeDeps()

    # -- registration ------------------------------------------------------
    def register(self, registry: InteractivityRegistry) -> None:
        registry.register_action(ACTION_SEND_DOCUSIGN, KIND_SEND_DOCUSIGN)
        registry.register_kind(
            KIND_SEND_DOCUSIGN,
            self._handle_send_docusign,
            value_model=SendDocusignPayload,
        )
        registry.register_callback(MODAL_CALLBACK_ID, KIND_NDA_DOCUSIGN_MODAL)
        registry.register_kind(KIND_NDA_DOCUSIGN_MODAL, self._handle_modal_submit)
        registry.register_action(ACTION_ENV_CONFIRM_SEND, KIND_ENV_CONFIRM)
        registry.register_action(ACTION_ENV_CONFIRM_CANCEL, KIND_ENV_CONFIRM)
        registry.register_kind(
            KIND_ENV_CONFIRM, self._handle_confirm, value_model=EnvConfirmPayload
        )
        registry.register_action(ACTION_ENV_USE_DOC, KIND_ENV_USE_DOC)
        registry.register_kind(
            KIND_ENV_USE_DOC, self._handle_use_doc, value_model=EnvUseDocPayload
        )
        registry.register_action(ACTION_DECLINE_DOC, KIND_DECLINE_DOC)
        registry.register_kind(
            KIND_DECLINE_DOC, self._handle_decline, value_model=DeclineDocPayload
        )

    # -- (b) open the signing modal (send_docusign) ------------------------
    def _handle_send_docusign(
        self, interaction: Interaction, deps: InteractivityDeps
    ) -> None:
        payload = interaction.payload
        if not isinstance(payload, SendDocusignPayload):
            return
        env = _interaction_envelope(interaction, deps)
        state = _load_correlation(deps.session_factory, payload.ref)
        if state is None:
            _deliver_text(env, EXPIRED_STATE_TEXT, deps)
            return
        open_view = self._get_open_view(deps)
        if open_view is None or not interaction.trigger_id:
            log.warning("bot.envelope.modal.no_opener", ref=payload.ref)
            _deliver_text(env, MODAL_UNAVAILABLE_TEXT, deps)
            return
        view = docusign_modal_view(
            payload.ref, file_name=str(state.get("file_name") or "")
        )
        try:
            open_view(interaction.trigger_id, view)
            log.info("bot.envelope.modal.opened", ref=payload.ref)
        except Exception as exc:  # noqa: BLE001 — a views.open failure is friendly, not a crash
            log.warning(
                "bot.envelope.modal.open_failed", ref=payload.ref, error=repr(exc)
            )
            _deliver_text(env, MODAL_UNAVAILABLE_TEXT, deps)

    # -- (b) modal submit → build the SAME confirm card --------------------
    def _handle_modal_submit(
        self, interaction: Interaction, deps: InteractivityDeps
    ) -> None:
        ref = str(_dig(interaction.raw, "view", "private_metadata") or "")
        state = _load_correlation(deps.session_factory, ref)
        if not ref or state is None:
            log.warning("bot.envelope.modal.no_state", ref=ref)
            return  # the modal already closed; no thread to recover a reply into
        env = _stored_envelope(state, deps)
        parsed = _parse_modal(interaction.state_values)
        if not (_valid_email(parsed.amp_email) and _valid_email(parsed.cp_email)):
            log.info("bot.envelope.modal.invalid_emails", ref=ref)
            _deliver_text(env, MODAL_INVALID_TEXT, deps)
            return
        signers = [parsed.amp_email.strip(), parsed.cp_email.strip()]
        updates = {
            "signer_emails": signers,
            "cc_emails": parsed.cc_emails,
            "routing": parsed.order,
            "cc_timing": parsed.cc_timing,
        }
        _update_correlation(deps.session_factory, ref, updates)
        summary = build_summary(
            signers,
            parsed.order,
            parsed.cc_emails,
            parsed.cc_timing,
            str(state.get("file_name") or ""),
        )
        log.info("bot.envelope.modal.collected", ref=ref, order=parsed.order)
        _post_blocks(
            env,
            envelope_confirm_blocks(summary, ref),
            ENVELOPE_CONFIRM_FALLBACK_TEXT,
            deps,
        )

    # -- (a)/(b) the Confirm & send / Cancel click -------------------------
    def _handle_confirm(
        self, interaction: Interaction, deps: InteractivityDeps
    ) -> None:
        payload = interaction.payload
        if not isinstance(payload, EnvConfirmPayload):
            return
        env = _interaction_envelope(interaction, deps)
        action = _confirm_action(interaction, payload)
        if action == "cancel":
            _update_correlation(deps.session_factory, payload.ref, {"cancelled": True})
            _deliver_text(env, CANCEL_TEXT, deps)
            return
        state = _load_correlation(deps.session_factory, payload.ref)
        if state is None or state.get("cancelled"):
            _deliver_text(env, EXPIRED_STATE_TEXT, deps)
            return
        if state.get("sent_envelope_id"):
            _deliver_text(env, ALREADY_SENT_TEXT, deps)
            return
        self._send(env, payload.ref, state, deps)

    # -- (c) use the recovered thread doc → (a)/(b) ------------------------
    def _handle_use_doc(
        self, interaction: Interaction, deps: InteractivityDeps
    ) -> None:
        payload = interaction.payload
        if not isinstance(payload, EnvUseDocPayload):
            return
        env = _interaction_envelope(interaction, deps)
        state = _load_correlation(deps.session_factory, payload.ref)
        if state is None:
            _deliver_text(env, EXPIRED_STATE_TEXT, deps)
            return
        try:
            docx_bytes = _doc_bytes_from_payload(
                state, slack_fetch=self._get_slack_fetch(deps)
            )
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "bot.envelope.use_doc.fetch_failed", ref=payload.ref, error=repr(exc)
            )
            _deliver_text(env, DOC_GONE_TEXT, deps)
            return
        leftovers = scan_docx_tokens(docx_bytes)
        if leftovers:
            _deliver_text(env, _token_refusal_text(leftovers), deps)
            return
        # Inline the bytes so confirm/modal don't re-fetch a possibly-deleted Slack file.
        _update_correlation(
            deps.session_factory,
            payload.ref,
            {"doc_b64": base64.b64encode(docx_bytes).decode("ascii")},
        )
        signers = _clean_emails(state.get("signer_emails"))
        file_name = str(state.get("file_name") or DEFAULT_DOC_NAME)
        if len(signers) >= 2:
            summary = build_summary(
                signers,
                _normalize_routing(str(state.get("routing") or "")),
                _clean_emails(state.get("cc_emails")),
                _normalize_cc_timing(str(state.get("cc_timing") or "after")),
                file_name,
            )
            _post_blocks(
                env,
                envelope_confirm_blocks(summary, payload.ref),
                ENVELOPE_CONFIRM_FALLBACK_TEXT,
                deps,
            )
        else:
            _post_blocks(
                env,
                signer_details_button_blocks(file_name, payload.ref),
                SEND_DOCUSIGN_FALLBACK_TEXT,
                deps,
            )

    def _handle_decline(
        self, interaction: Interaction, deps: InteractivityDeps
    ) -> None:
        env = _interaction_envelope(interaction, deps)
        _deliver_text(env, DECLINE_TEXT, deps)

    # -- the actual send ---------------------------------------------------
    def _send(
        self,
        env: Envelope | None,
        ref: str,
        state: dict[str, Any],
        deps: InteractivityDeps,
    ) -> None:
        signers = _clean_emails(state.get("signer_emails"))
        if len(signers) < 1:
            _deliver_text(env, MODAL_INVALID_TEXT, deps)
            return
        cc = _clean_emails(state.get("cc_emails"))
        routing = _normalize_routing(str(state.get("routing") or ""))
        cc_timing = _normalize_cc_timing(str(state.get("cc_timing") or "after"))
        file_name = str(state.get("file_name") or DEFAULT_DOC_NAME)
        try:
            docx_bytes = _doc_bytes_from_payload(
                state, slack_fetch=self._get_slack_fetch(deps)
            )
        except Exception as exc:  # noqa: BLE001
            log.warning("bot.envelope.send.doc_gone", ref=ref, error=repr(exc))
            _deliver_text(env, DOC_GONE_TEXT, deps)
            return

        from ...integrations.docusign import (
            DocuSignError,
            DocuSignRetryableError,
            DocuSignTerminalError,
            DocuSignUnavailable,
        )

        try:
            sender = self._get_sender(deps)
        except DocuSignUnavailable:
            log.info("bot.envelope.send.capability_off", ref=ref)
            _deliver_text(env, DOCUSIGN_UNAVAILABLE_TEXT, deps)
            return

        try:
            result = sender(
                docx_bytes=docx_bytes,
                filename=file_name,
                signers=signers,
                routing=routing,
                cc=cc,
                cc_timing=cc_timing,
            )
        except (DocuSignTerminalError, DocuSignRetryableError, DocuSignError) as exc:
            detail = getattr(exc, "error_code", None) or ""
            log.warning("bot.envelope.send.failed", ref=ref, error=repr(exc))
            self._persist(
                deps,
                state,
                status="failed",
                envelope_id=None,
                idempotency_key=_failed_idempotency_key(
                    docx_bytes, signers, routing, cc, cc_timing
                ),
                signers=signers,
                cc=cc,
                routing=routing,
            )
            _deliver_text(env, _send_failure_text(str(detail)), deps)
            return

        _update_correlation(
            deps.session_factory, ref, {"sent_envelope_id": result.envelope_id}
        )
        self._persist(
            deps,
            state,
            status="sent",
            envelope_id=result.envelope_id,
            idempotency_key=result.idempotency_key,
            signers=signers,
            cc=cc,
            routing=routing,
        )
        log.info(
            "bot.envelope.send.sent",
            ref=ref,
            envelope_id=result.envelope_id,
            requested_by=state.get("requested_by", ""),
        )
        _deliver_text(env, SEND_SUCCESS_TEXT, deps)

    def _persist(
        self,
        deps: InteractivityDeps,
        state: dict[str, Any],
        *,
        status: str,
        envelope_id: str | None,
        idempotency_key: str,
        signers: list[str],
        cc: list[str],
        routing: str,
    ) -> None:
        """Persist the attempt with its requester mapping (PLAN §3.10) — idempotent on the key. Best
        effort: a persistence failure never costs the send (mirrors the review intent's save)."""
        if deps.session_factory is None:
            return
        from ...integrations.models import save_envelope_attempt

        try:
            with deps.session_factory() as db:
                save_envelope_attempt(
                    db,
                    idempotency_key=idempotency_key,
                    status=status,
                    channel=str(state.get("channel") or "slack"),
                    routing=routing,
                    requested_by=str(state.get("requested_by") or ""),
                    signer_emails=signers,
                    cc_emails=cc,
                    envelope_id=envelope_id,
                    slack_channel=str(state.get("slack_channel") or ""),
                    slack_thread_ts=str(state.get("slack_thread_ts") or ""),
                    email_message_id=str(state.get("email_message_id") or ""),
                )
        except Exception:  # noqa: BLE001 — audit persistence is best-effort
            log.exception("bot.envelope.persist_failed", envelope_id=envelope_id)

    # -- collaborator resolution (fake in tests, production lazily) ---------
    def _get_slack_fetch(self, deps: InteractivityDeps) -> SlackFileFetcher:
        if self._deps.slack_fetch is not None:
            return self._deps.slack_fetch
        settings = deps.settings or _get_settings()
        return HttpSlackFileFetcher(settings.slack_bot_token)

    def _get_open_view(self, deps: InteractivityDeps) -> OpenView | None:
        if self._deps.open_view is not None:
            return self._deps.open_view
        # Resolve the Slack client behind the reply sink's post_blocks (no new dep on the deps object).
        pb = deps.post_blocks
        sink = getattr(pb, "__self__", None)
        client = getattr(sink, "_client", None)
        if client is None or not hasattr(client, "views_open"):
            return None

        def _open(trigger_id: str, view: dict) -> None:
            client.views_open(trigger_id=trigger_id, view=view)

        return _open

    def _get_sender(self, deps: InteractivityDeps) -> EnvelopeSender:
        if self._deps.sender is not None:
            return self._deps.sender
        from ...capabilities import build_registry
        from ...integrations.docusign import build_docusign_client

        settings = deps.settings or _get_settings()
        registry = self._deps.docusign_registry or build_registry(settings)
        client = build_docusign_client(
            settings, registry
        )  # raises DocuSignUnavailable if off
        return client.create_and_send_envelope


# --------------------------------------------------------------------------- #
# Interactivity helpers (reply/blocks delivery, envelope building, modal parse)
# --------------------------------------------------------------------------- #
def _get_settings() -> Settings:
    from app.config import get_settings

    return get_settings()


def _interaction_envelope(
    interaction: Interaction, deps: InteractivityDeps
) -> Envelope | None:
    """A Slack reply envelope addressed at the interaction's OWN thread (the card/button's message)."""
    channel = interaction.channel_id
    if not channel:
        return None
    thread = interaction.thread_ts or interaction.message_ts
    from_email = deps.settings.nda_bot_from_email if deps.settings else ""
    return Envelope(
        channel="slack",
        event_key=f"slack:int:env:{channel}:{thread or 'root'}",
        slack_channel=channel,
        slack_thread_ts=thread or "",
        verified_sender=True,
        from_email=from_email,
    )


def _stored_envelope(state: dict[str, Any], deps: InteractivityDeps) -> Envelope | None:
    """A Slack reply envelope from the STORED origin (modal submit has no channel context of its own)."""
    channel = str(state.get("slack_channel") or "")
    if not channel:
        return None
    thread = str(state.get("slack_thread_ts") or "")
    from_email = deps.settings.nda_bot_from_email if deps.settings else ""
    return Envelope(
        channel="slack",
        event_key=f"slack:int:env:{channel}:{thread or 'root'}",
        slack_channel=channel,
        slack_thread_ts=thread,
        verified_sender=True,
        from_email=from_email,
    )


def _deliver_text(env: Envelope | None, text: str, deps: InteractivityDeps) -> None:
    if env is None or deps.service is None:
        return
    try:
        deps.service.deliver(env, Reply(text=text))
    except Exception as exc:  # noqa: BLE001 — delivery is fail-soft
        log.warning("bot.envelope.deliver_failed", error=repr(exc))


def _post_blocks(
    env: Envelope | None,
    blocks: list[dict],
    fallback: str,
    deps: InteractivityDeps,
) -> None:
    if env is None or deps.post_blocks is None:
        return
    try:
        deps.post_blocks(env, blocks, fallback)
    except Exception as exc:  # noqa: BLE001 — delivery is fail-soft
        log.warning("bot.envelope.post_blocks_failed", error=repr(exc))


def _confirm_action(interaction: Interaction, payload: EnvConfirmPayload) -> str:
    a = (payload.action or "").strip().lower()
    if a in ("send", "cancel"):
        return a
    return "cancel" if interaction.action_id == ACTION_ENV_CONFIRM_CANCEL else "send"


def _modal_value(state: dict[str, Any], block_id: str, action_id: str) -> str:
    """Read one value out of a Slack view ``state.values`` (plain_text_input or static_select)."""
    block = state.get(block_id) or {}
    element = block.get(action_id) if isinstance(block, dict) else None
    if not isinstance(element, dict):
        return ""
    if "value" in element:
        return str(element.get("value") or "")
    option = element.get("selected_option")
    if isinstance(option, dict):
        return str(option.get("value") or "")
    return ""


def _parse_modal(state_values: dict[str, Any]) -> _ModalInput:
    """Parse the ``nda_docusign`` modal submit (reference §3.7 "Parse Modal Submit")."""
    return _ModalInput(
        amp_email=_modal_value(state_values, MODAL_BLOCK_AMP, MODAL_ACTION_AMP),
        cp_email=_modal_value(state_values, MODAL_BLOCK_CP, MODAL_ACTION_CP),
        order=_normalize_routing(
            _modal_value(state_values, MODAL_BLOCK_SEQ, MODAL_ACTION_SEQ)
        ),
        cc_emails=_parse_cc_emails(
            _modal_value(state_values, MODAL_BLOCK_CC, MODAL_ACTION_CC)
        ),
        cc_timing=_normalize_cc_timing(
            _modal_value(state_values, MODAL_BLOCK_CC_SEQ, MODAL_ACTION_CC_SEQ)
        ),
    )


def _failed_idempotency_key(
    docx_bytes: bytes,
    signers: list[str],
    routing: str,
    cc: list[str],
    cc_timing: str,
) -> str:
    """Compute the ported idempotency key for a FAILED attempt (the sender returns none on failure), so
    the audit row still keys the same doc+recipients (reference §2.9)."""
    from ...integrations.docusign import (
        build_recipients,
        derive_idempotency_key,
        normalize_routing,
    )

    docx_b64 = base64.b64encode(docx_bytes).decode("ascii")
    recipients = build_recipients(signers, normalize_routing(routing), cc, cc_timing)
    return derive_idempotency_key(docx_b64, recipients)


def register_envelope(
    registry: InteractivityRegistry, *, deps: EnvelopeDeps | None = None
) -> None:
    """Register the envelope interactivity kinds onto ``registry`` (called from
    ``default_interactivity_registry``; tests pass ``deps`` to bind fakes)."""
    EnvelopeInteractivity(deps).register(registry)


__all__ = [
    "EnvelopeIntent",
    "EnvelopeInteractivity",
    "EnvelopeDeps",
    "EnvelopeSender",
    "OpenView",
    "SendDocusignPayload",
    "EnvConfirmPayload",
    "EnvUseDocPayload",
    "DeclineDocPayload",
    "register_envelope",
    "scan_docx_tokens",
    "build_summary",
    "ENV_CORRELATION_KIND",
    "EMAIL_NO_DOC_TEXT",
    "EMAIL_HAS_DOC_TEXT",
    "SLACK_NO_DOC_TEXT",
    "DECLINE_TEXT",
    "CANCEL_TEXT",
    "SEND_SUCCESS_TEXT",
    "DOCUSIGN_UNAVAILABLE_TEXT",
    "ALREADY_SENT_TEXT",
    "EXPIRED_STATE_TEXT",
]
