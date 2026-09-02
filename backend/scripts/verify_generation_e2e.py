"""End-to-end verification of the REAL NDA generation pipeline — no mocks, no network, no LLM.

The 1400+ unit tests each exercise one seam (``test_fill_docx`` fills in isolation;
``test_bot_flow_completion`` drives the flow but STUBS the template resolve + fill). This harness
closes the gap the ``/verify`` philosophy targets: drive the ACTUAL assembled pipeline against a REAL
tokenised .docx and inspect the REAL output artifact.

Chain exercised (every step production code, on a throwaway SQLite DB):

  build a real tokenised .docx  →  seed it as the current US/Company template (models_v2 rows)
  →  run_generation (the Tally webhook's generation seam) with a realistic token table + routing
  →  real normalize_codes → resolve_template_docx → fill_docx → deliver to a captured reply sink
  →  open the delivered .docx and ASSERT the values substituted and NO {{token}} remains.

Run:  backend/.venv/bin/python -m scripts.verify_generation_e2e   (from backend/, or via the module
path). Exit 0 = the generation pipeline produced a correct filled document end to end.
"""

from __future__ import annotations

import io
import sys
import tempfile
from pathlib import Path

from docx import Document


def _make_tokenised_template() -> bytes:
    """A minimal but realistic tokenised NDA .docx: body + a table cell + a header, each carrying
    {{tokens}}, some deliberately SPLIT across runs (bold mid-token) to exercise the run-aware fill."""
    doc = Document()
    doc.add_heading("Non-Disclosure Agreement", level=1)
    doc.add_paragraph(
        "This Agreement is between Amperesand (signed by {{amperesand_signer_name}}) "
        "and {{counterparty_name}}."
    )
    # A split-across-runs token: "{{" + bold "purpose" + "}}" — the filler must rejoin runs.
    p = doc.add_paragraph("Purpose: ")
    p.add_run("{{")
    p.add_run("purpose").bold = True
    p.add_run("}}")
    doc.add_paragraph("Effective date: {{effective_date}}.")
    # A token inside a table cell (the filler recurses into tables).
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Counterparty registration no."
    table.rows[0].cells[1].text = "{{counterparty_company_registration_number}}"
    # A token in the header (the filler recurses into headers/footers).
    doc.sections[0].header.paragraphs[0].text = "NDA — {{counterparty_name}}"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _seed_template(db, docx_bytes: bytes) -> None:
    """Attach the tokenised .docx as the CURRENT version of the EXISTING US/Company/NotApplicable
    template. This mirrors production exactly: the 8 template metadata rows are seeded by migrations
    (``init_db``/``seed_catalog``) with NULL blob bytes; loading a real .docx = add a DocumentBlob +
    a current TemplateVersion pointing at it (what the template studio / seeder does)."""
    import hashlib
    import uuid

    from sqlalchemy import select

    from app.models_v2 import DocumentBlob, Template, TemplateVersion

    tpl = db.execute(
        select(Template).where(
            Template.jurisdiction_code == "US",
            Template.counterparty_type_code == "Company",
            Template.mutuality_code == "NotApplicable",
        )
    ).scalar_one()

    blob = DocumentBlob(
        id=uuid.uuid4().hex,
        sha256=hashlib.sha256(docx_bytes).hexdigest(),
        bytes=docx_bytes,
        byte_size=len(docx_bytes),
        mime_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    db.add(blob)

    # The current tokenised version row is pre-seeded with NULL blob_id (metadata-owned-by-migrations);
    # "loading the .docx" = point it at the blob — exactly the production gap the studio/seeder closes.
    ver = db.execute(
        select(TemplateVersion).where(
            TemplateVersion.template_id == tpl.id,
            TemplateVersion.variant_code == "tokenised",
            TemplateVersion.is_current.is_(True),
        )
    ).scalar_one()
    ver.blob_id = blob.id
    db.commit()


def main() -> int:
    # Point the app at a throwaway SQLite file BEFORE importing anything that builds the engine.
    tmpdir = tempfile.mkdtemp(prefix="nda-verify-")
    db_path = Path(tmpdir) / "verify.db"
    import os

    os.environ["DATABASE_URL"] = f"sqlite:///{db_path}"
    os.environ["APP_ENV"] = "test"

    from app.bot.flows import run_generation
    from app.db import SessionLocal, init_db

    init_db()  # create_all (all models) + seed the default org (the template FK target)

    template_bytes = _make_tokenised_template()
    with SessionLocal() as db:
        _seed_template(db, template_bytes)

    # Capture the delivered document from the REAL generation flow — the exact seam the Tally webhook
    # drives (``run_generation``: normalize_codes → resolve_template_docx → fill_docx → deliver). The
    # token table + routing selectors + origin here are what ``tally.map_submission`` would produce.
    captured: dict[str, object] = {}

    class _Sink:
        def deliver(self, envelope, reply):
            captured["envelope"] = envelope
            captured["reply"] = reply
            return "ok"

    result = run_generation(
        values={
            "amperesand_signer_name": "Alice Tan",
            "counterparty_name": "Acme Corporation",
            "purpose": "evaluating a commercial partnership",
            "counterparty_company_registration_number": "UEN-2026-42",
        },
        jurisdiction="US",
        counterparty_type="company",
        mutuality="",
        origin_context={
            "channel": "slack",
            "slack_channel": "C_VERIFY",
            "slack_thread_ts": "1700.1",
        },
        ref="verify-tally",
        service=_Sink(),
        session_factory=SessionLocal,
    )

    assert result.ok, f"completion not ok: {result.reason}"
    assert result.delivered, "document was not delivered"
    reply = captured.get("reply")
    assert reply is not None and reply.attachments, "no document attachment delivered"
    filled_bytes = reply.attachments[0].content

    # Open the REAL produced .docx and verify substitution + no leftover tokens.
    filled = Document(io.BytesIO(filled_bytes))
    full_text = "\n".join(p.text for p in filled.paragraphs)
    for table in filled.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    for section in filled.sections:
        full_text += "\n" + section.header.paragraphs[0].text

    checks = {
        "amperesand signer filled": "Alice Tan" in full_text,
        "counterparty filled": "Acme Corporation" in full_text,
        "split-run purpose token filled": "evaluating a commercial partnership"
        in full_text,
        "table-cell token filled": "UEN-2026-42" in full_text,
        "no leftover {{ tokens": "{{" not in full_text and "}}" not in full_text,
        "delivered to origin Slack thread": getattr(
            captured["envelope"], "slack_thread_ts", None
        )
        == "1700.1",
    }

    print("\n=== NDA generation end-to-end verification ===")
    ok = True
    for name, passed in checks.items():
        print(f"  [{'PASS' if passed else 'FAIL'}] {name}")
        ok = ok and passed
    print(
        f"  filled document: {len(filled_bytes)} bytes, {len(filled.paragraphs)} paragraphs"
    )
    print("=" * 46)
    if ok:
        print("RESULT: PASS — the real pipeline produced a correct filled NDA.\n")
        return 0
    print("RESULT: FAIL — see the checks above.\n")
    return 1


if __name__ == "__main__":
    sys.exit(main())
