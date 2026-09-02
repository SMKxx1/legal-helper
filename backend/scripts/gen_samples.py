"""Generate the synthetic sample `.docx` files under `samples/` (plan §6 Phase 2 item 5).

Three documents, each built to exercise a specific reviewer/coverage behavior, with no real
company text: a mutual NDA missing its governing-law clause, an MSA with an uncapped liability
clause, and a one-page letter that isn't a contract at all (a negative case — no findings, no
missing-clause noise). Run with ``python -m scripts.gen_samples`` from ``backend/``.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

_REPO_ROOT = Path(__file__).resolve().parents[2]
_SAMPLES_DIR = _REPO_ROOT / "samples"


def _heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def _clause(doc: Document, heading: str, body: str) -> None:
    doc.add_paragraph(heading, style="Heading 2")
    doc.add_paragraph(body)


def build_nda_missing_governing_law() -> Document:
    """A mutual NDA that is otherwise standard but never states a governing law or dispute
    forum — the coverage agent's required-checklist net should catch this."""
    doc = Document()
    _heading(doc, "MUTUAL NON-DISCLOSURE AGREEMENT")
    doc.add_paragraph(
        "This Mutual Non-Disclosure Agreement (this \"Agreement\") is entered into between "
        "Northwind Robotics, Inc. (\"Northwind\") and Fictive Supply Co. (\"Fictive\"), "
        "collectively the \"Parties,\" for the purpose of evaluating a potential business "
        "relationship (the \"Purpose\")."
    )
    _clause(
        doc,
        "1. Confidential Information",
        "\"Confidential Information\" means any non-public information disclosed by one Party "
        "(the \"Disclosing Party\") to the other (the \"Receiving Party\"), whether in writing, "
        "orally, or by inspection, that is marked confidential or that a reasonable person would "
        "understand to be confidential given the nature of the information and the circumstances "
        "of disclosure. Confidential Information does not include information that: (a) is or "
        "becomes publicly available through no fault of the Receiving Party; (b) was already "
        "known to the Receiving Party without an obligation of confidentiality; (c) is "
        "independently developed by the Receiving Party without use of the Confidential "
        "Information; or (d) is rightfully received from a third party without restriction.",
    )
    _clause(
        doc,
        "2. Obligations",
        "Each Receiving Party shall use the other Party's Confidential Information solely for "
        "the Purpose, shall protect it using at least the same degree of care it uses to protect "
        "its own confidential information (and in no event less than reasonable care), and shall "
        "not disclose it to any third party without the Disclosing Party's prior written consent, "
        "except to employees and advisors who have a need to know and are bound by confidentiality "
        "obligations at least as protective as those in this Agreement.",
    )
    _clause(
        doc,
        "3. Term",
        "This Agreement remains in effect for two (2) years from the date first written above. "
        "The confidentiality obligations in Section 2 survive termination of this Agreement for a "
        "period of three (3) years, except that obligations with respect to trade secrets survive "
        "for as long as the information remains a trade secret under applicable law.",
    )
    _clause(
        doc,
        "4. Return or Destruction",
        "Upon the Disclosing Party's written request, the Receiving Party shall promptly return "
        "or destroy all Confidential Information in its possession and certify such destruction "
        "in writing, except for one archival copy retained solely to demonstrate compliance with "
        "this Agreement.",
    )
    _clause(
        doc,
        "5. No License",
        "Nothing in this Agreement grants either Party any license or other right, by "
        "implication or otherwise, to the other Party's Confidential Information, patents, "
        "copyrights, trademarks, or other intellectual property.",
    )
    # Deliberately NO governing-law / dispute-resolution clause anywhere in the document.
    _clause(
        doc,
        "6. Miscellaneous",
        "This Agreement constitutes the entire agreement between the Parties regarding its "
        "subject matter and supersedes all prior discussions on that subject. Neither Party may "
        "assign this Agreement without the other Party's prior written consent, except to an "
        "affiliate or successor in a merger or acquisition. This Agreement may be signed in "
        "counterparts.",
    )
    return doc


def build_msa_uncapped_liability() -> Document:
    """A master services agreement whose limitation-of-liability clause is deliberately
    one-sided and uncapped on the customer's side — the reviewer's high-severity finding case."""
    doc = Document()
    _heading(doc, "MASTER SERVICES AGREEMENT")
    doc.add_paragraph(
        "This Master Services Agreement (this \"Agreement\") is entered into between Acme "
        "Fulfillment Corp. (\"Customer\") and Widgets & Sundries Ltd. (\"Provider\") for the "
        "provision of logistics services described in one or more Statements of Work."
    )
    _clause(
        doc,
        "1. Services",
        "Provider shall perform the services described in each Statement of Work executed under "
        "this Agreement (the \"Services\") in a professional and workmanlike manner consistent "
        "with generally accepted industry standards.",
    )
    _clause(
        doc,
        "2. Fees and Payment",
        "Customer shall pay Provider's fees as set out in the applicable Statement of Work "
        "within thirty (30) days of receipt of a correct invoice. Amounts not paid when due "
        "accrue interest at one percent (1%) per month or the maximum rate permitted by law, "
        "whichever is lower.",
    )
    _clause(
        doc,
        "3. Confidentiality",
        "Each party shall protect the other's confidential information using reasonable care "
        "and shall not disclose it to any third party except as needed to perform this Agreement "
        "or as required by law.",
    )
    _clause(
        doc,
        "4. Term and Termination",
        "This Agreement commences on the Effective Date and continues for one (1) year, "
        "automatically renewing for successive one-year terms unless either party gives at least "
        "sixty (60) days' notice of non-renewal. Either party may terminate for the other's "
        "uncured material breach on thirty (30) days' written notice.",
    )
    _clause(
        doc,
        "5. Indemnification",
        "Customer shall indemnify, defend, and hold harmless Provider from and against any and "
        "all claims, damages, and expenses arising out of Customer's use of the Services, "
        "regardless of the theory of liability asserted.",
    )
    _clause(
        doc,
        "6. Limitation of Liability",
        "PROVIDER'S TOTAL LIABILITY ARISING OUT OF OR RELATED TO THIS AGREEMENT SHALL NOT "
        "EXCEED THE FEES PAID BY CUSTOMER IN THE TWELVE (12) MONTHS PRECEDING THE CLAIM. "
        "CUSTOMER'S LIABILITY UNDER THIS AGREEMENT IS UNCAPPED AND UNLIMITED, AND CUSTOMER "
        "SHALL BE LIABLE FOR ALL DIRECT, INDIRECT, INCIDENTAL, SPECIAL, AND CONSEQUENTIAL "
        "DAMAGES ARISING FROM ITS BREACH OF THIS AGREEMENT, WITHOUT LIMITATION.",
    )
    _clause(
        doc,
        "7. Governing Law",
        "This Agreement is governed by the laws of the State of Delaware, without regard to its "
        "conflict-of-laws principles. The parties consent to the exclusive jurisdiction of the "
        "state and federal courts located in Delaware.",
    )
    _clause(
        doc,
        "8. Assignment",
        "Neither party may assign this Agreement without the other's prior written consent, not "
        "to be unreasonably withheld, except to an affiliate or in connection with a merger, "
        "acquisition, or sale of substantially all assets.",
    )
    return doc


def build_cover_letter_not_a_contract() -> Document:
    """A one-page cover letter — a negative case with no contractual clauses at all, so a good
    classifier should label it something other than a contract type and the reviewer should
    return no findings (nothing to find harm in)."""
    doc = Document()
    doc.add_paragraph("Riverside Analytics LLC")
    doc.add_paragraph("400 Harbor View Road, Suite 210")
    doc.add_paragraph("Springfield, ST 00000")
    doc.add_paragraph("")
    doc.add_paragraph("August 3, 2026")
    doc.add_paragraph("")
    doc.add_paragraph("Dear Ms. Alvarez,")
    doc.add_paragraph(
        "Thank you for taking the time to meet with our team last week to discuss Riverside "
        "Analytics' upcoming data platform migration. We were glad to walk through your current "
        "reporting workflows and appreciate the context your team provided on the Q3 roadmap."
    )
    doc.add_paragraph(
        "As a follow-up to that conversation, we are preparing a short written summary of the "
        "options we discussed and expect to share it with you by the end of next week. In the "
        "meantime, please let us know if there is any additional information we can provide, or "
        "if you would like to schedule a further call with our engineering lead."
    )
    doc.add_paragraph(
        "We enjoyed the discussion and look forward to continuing the conversation."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Warm regards,")
    doc.add_paragraph("Jordan Lee")
    doc.add_paragraph("Head of Partnerships, Riverside Analytics LLC")
    return doc


def main() -> None:
    _SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    targets = {
        "nda_missing_governing_law.docx": build_nda_missing_governing_law,
        "msa_uncapped_liability.docx": build_msa_uncapped_liability,
        "cover_letter_not_a_contract.docx": build_cover_letter_not_a_contract,
    }
    for filename, builder in targets.items():
        path = _SAMPLES_DIR / filename
        builder().save(str(path))
        print(f"wrote {path.relative_to(_REPO_ROOT)}")


if __name__ == "__main__":
    main()
